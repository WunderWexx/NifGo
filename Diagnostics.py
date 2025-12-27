"""
After the program has run, the output needs to be checked for mistakes both human and software. This process should be automated as much as possible.
The ExternalDiagnostics class checks mistakes after the program has finished. This may take a while.
The InlineDiagnostics class checks things per document as it is being filled with data. This class should perform as fast as possible.
"""
import pandas as pd
import re
import Utilities as util
from docx import Document
from math import ceil
from statistics import fmean
import itertools
import os

#Variables
genes_by_phenotype_pattern = {
            r'^[U,R,N,I,P]M$': [
                ['COMT', 'CYP1A2', 'CYP2B6', 'CYP2C9', 'CYP2C19', 'CYP2D6', 'CYP3A4', 'DPYD', 'G6PD', 'MTHFR677',
                 'NUDT15', 'TPMT', 'UGT1A1', 'VKORC1'], 'NM'],  # Zoals NM
            r'^[I,N,D,P]F$': [['ABCG2', 'SLCO1B1', 'ABCB1'], 'NF'],  # Zoals NF
            r'^[a-z]{1,12}$|^non-expresser$|^intermediate-expresser$': [['CYP3A5'], 'non-expresser'], # Zoals een string tot max 12 kleine letters
            r'^normaal$|^risico$': [['HLA-B*1502', 'HLA-B*5701', 'HLA-A*3101'], 'normaal'],  # normaal of risico
            r'^[U,R,N,I,P]M$|^Deficient$': [['G6PD'], 'NM']  # Zoals NM of Deficient
            }

genes_by_genotype_pattern = {
            r'\D\D/\D\D': ['CACNA1S', 'CFTR', 'RYR1', "VDR"],  # Zoals WT/WT
            r'\d\d\d\D\D': ['SLCO1B1', "ABCG2"],  # Zoals 521TC
            r'\d\d\d\d\D\D': ['VKORC1'],  # Zoals 1639GG
            r'^AS: (\d|\d\.\d)$': ['DPYD'],  # Zoals AS: 2 of AS: 1.5
            r'\D\D\D/\D\D\D': ['COMT'],  # Zoals Met/Met
            r'(Null|Present)/(Null|Present)': ['GSTM1'],  # Zoals Null/Present
            r'(\D\.=|Null)/(\D\.=|Null)': ['MTRNR1'],  # Zoals m.=/Null
            r'((Null|\D|Val)/(Null|\D|Val)|[A-Z])': ["ABCB1", "ACE", "ADIPOQ", "ADRA2A", "ALDH2", "AMDHD1",
                                                     "BCO1",
                                                     "BDNF", "CYP1A1", "CYP2R1", "CYP17A1", "CYP24A1",
                                                     'DHCR7 / NADSYN1', "DRD2", "F2", "F5", "FTO", "G6PD", "GC",
                                                     "GCK, YKT6",
                                                     "GSTP1",
                                                     "IFNL3/IL28B", "IGF1", "LDLR", "LOC105447645; FUT2", "MAO-B",
                                                     "MC4R",
                                                     "MnSOD",
                                                     "MTHFR1298", "MTHFR677", "MTNR1B", "NBPF3", "NQ01", "OPRM1",
                                                     "PON1",
                                                     "Sult1E1", "TCF7L2", "TMEM165; CLOCK", "TNFa",
                                                     "UCP2"],  # Zoals A/A of Null/A of Null/Val
            r'(\D|\D\D|\D\d)/(\D|\D\D|\D\d)': ['BChE'],  # Zoals U/U of F2/F2 of Sc/Sc
            r'(Null|\*\D)/(Null|\*\D)': ['GSTP1', 'GSTT1'],  # Zoals *A/*A of Null/*A
            r'^negatief$|^positief$': ['HLA-A*3101', 'HLA-B*1502', 'HLA-B*5701'],  # negatief of positief
            r'(\*.*|Null)/(\*.*|Null)': ['other'],  # Zoals *1/*1 of zelfs *4.001/*7A+1B, en ook *1/Null
        }

class ExternalDiagnostics:
    def __init__(self):
        # Import complete dataframe
        self.complete_df = pd.read_excel("Output/Dataframes/complete.xlsx")

        # Get list of all genes to be reported
        self.farmacogenetic_genes = []
        most_recent_farmacogenetisch = util.get_most_recent_template('Farmacogenetisch')
        doc = Document(f'Input/Templates/{most_recent_farmacogenetisch}')
        table1 = doc.tables[1]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            self.farmacogenetic_genes.append(gene)

        self.nutrigenomic_genes = []
        most_recent_nutrigenomics = util.get_most_recent_template('Nutrigenomics')
        doc = Document(f'Input/Templates/{most_recent_nutrigenomics}')
        table1 = doc.tables[1]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            self.nutrigenomic_genes.append(gene)
        table2 = doc.tables[2]
        for row in table2.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            self.nutrigenomic_genes.append(gene)

        self.all_genes = self.farmacogenetic_genes + self.nutrigenomic_genes

        # Get a list of all sample_id's
        self.unique_sample_ids = self.complete_df['sample_id'].unique().tolist()

        # Create a diagnostics file
        with open('Output/Diagnostics/diagnostics.txt','w') as diag:
            diag.write('Diagnostics:\n\n')
        diag.close()

    def check_phenotype_shape(self):
        self.phenotype_pattern_by_genes = {}
        self.genes_by_phenotype_pattern = genes_by_phenotype_pattern
        # create a dict with the genes as keys and the regex as the values
        for regex, values in self.genes_by_phenotype_pattern.items():
            genes = values[0]  # Extract the gene list
            for gene in genes:
                self.phenotype_pattern_by_genes[gene] = regex
        # create a dict with non-matching phenotypes
        farmacogenetic_subset = self.complete_df[self.complete_df['gene'].isin(self.farmacogenetic_genes)]

        def phenotype_filter(row):
            gene = row['gene']
            phenotype = row['phenotype']
            if gene in self.phenotype_pattern_by_genes:
                pattern = self.phenotype_pattern_by_genes[gene]
                # Return True if there's NO match, because these are the rows we want to keep
                return not bool(re.fullmatch(pattern, phenotype))
            return True

        faulty_phenotypes = farmacogenetic_subset[farmacogenetic_subset.apply(phenotype_filter, axis='columns')]
        faulty_phenotypes = faulty_phenotypes.iloc[:,1:]
        # Write that dict to the diagnostics.txt file
        with open('Output/Diagnostics/diagnostics.txt','a') as diag:
            diag.write('Faulty Phenotypes:\n')
            if not faulty_phenotypes.empty:
                faulty_phenotypes.to_csv(diag, mode='a', index=False, sep='\t', header=True)
            else:
                diag.write('No faulty phenotypes')
            diag.write('\n\n')
            diag.close()

    def check_genotype_shape(self):
        self.genotype_pattern_by_genes = {}
        self.genes_by_genotype_pattern = genes_by_genotype_pattern
        # create a dict with the genes as keys and the regex as the values
        for regex, genes in self.genes_by_genotype_pattern.items():
            for gene in genes:
                self.genotype_pattern_by_genes[gene] = regex
        # create a dict with non-matching phenotypes
        def genotype_filter(row):
            gene = row['gene']
            genotype = row['genotype']
            if gene in self.genotype_pattern_by_genes:
                pattern = self.genotype_pattern_by_genes[gene]
                # Return True if there's NO match, because these are the rows we want to keep
                return not bool(re.fullmatch(pattern, genotype))
            else:
                pattern = self.genotype_pattern_by_genes['other']
                return not bool(re.fullmatch(pattern, genotype)) # This means there is an empty genotype cell somewhere. Maybe we should make it more robust to that...

        faulty_genotypes = self.complete_df[self.complete_df.apply(genotype_filter, axis='columns')]
        faulty_genotypes = faulty_genotypes.iloc[:,1:]
        faulty_phenotypes = faulty_genotypes.iloc[:, 1:]
        # Write that dict to the diagnostics.txt file
        with open('Output/Diagnostics/diagnostics.txt', 'a') as diag:
            diag.write('Faulty Genotypes:\n')
            faulty_genotypes.to_csv(diag, mode='a', index=False, sep='\t', header=True)
            diag.write('\n')
            diag.close()

    def check_customerdata_available_to_reports(self, customerdata_df):
        """
        Checks if all reports have customer data available to themin the customer data file.
        """
        samples_missing_data = []
        customerdata_samples =  customerdata_df.iloc[:, 0].tolist()
        for sample_id in self.unique_sample_ids:
            if sample_id not in customerdata_samples:
                samples_missing_data.append(sample_id)
        with open('Output/Diagnostics/diagnostics.txt', 'a') as diag:
            diag.write('Samples missing customer data:\n')
            if samples_missing_data:
                for sample in samples_missing_data:
                    diag.write(f"{sample}\n")
            else:
                diag.write('All samples have customer data available.')
            diag.write('\n\n')
            diag.close()

    def check_data_available_per_customerdata(self, customerdata_df):
        """
        Checks if all available customer data has associated data.
        """
        samples_missing_data = []
        customerdata_samples = customerdata_df.iloc[:, 0].tolist()
        for sample_id in customerdata_samples:
            if sample_id not in self.unique_sample_ids:
                samples_missing_data.append(sample_id)
        with open('Output/Diagnostics/diagnostics.txt', 'a') as diag:
            diag.write('Customer samples without associated data:\n')
            if samples_missing_data:
                for sample in samples_missing_data:
                    diag.write(f"{sample}\n")
            else:
                diag.write('All customer samples have associated data')
            diag.write('\n\n')
            diag.close()

    def check_batch_size(self):
        sample_amount = len(self.unique_sample_ids)
        reports_amount = len(util.get_reports())
        reports_per_sample = 3
        expected_reports_amount = sample_amount * reports_per_sample
        batch_size = 24
        billable_batches = ceil(sample_amount / batch_size)
        with open('Output/Diagnostics/diagnostics.txt', 'a') as diag:
            diag.write('Batch data:\n')
            diag.write(f'From {sample_amount} samples, {reports_amount} reports were produced.\n'
                       f'With a batch size of 24, this comes to {billable_batches} billable batches.\n')
            if reports_amount != expected_reports_amount:
                diag.write(f'WARNING!\n'
                           f'{reports_amount} reports have been generated out of the {expected_reports_amount} reports expected!')
            diag.write('\n\n')
            diag.close()

    def check_deviation_percentage(self):
        # Read raw lines and filter only valid ones
        file_path = "Output/Diagnostics/deviations.txt"
        valid_rows = []
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                parts = line.strip().split('\t')
                if len(parts) == 3 and not parts[0].endswith('_2'):
                    valid_rows.append(parts)

        # Convert to DataFrame
        df = pd.DataFrame(valid_rows, columns=['sample_id', 'gene', 'deviation'])

        # Count deviations per gene
        gene_counts = df['gene'].value_counts().sort_values(ascending=False)

        # Write results to diagnostics.txt
        cleaned_sample_ids = [sid for sid in self.unique_sample_ids if '_2' not in sid]
        with open('Output/Diagnostics/diagnostics.txt', 'a') as diag:
            diag.write('Deviancy percentage per gene:\n')
            diag.write('Gene\tNumber of deviations\tPercentage of total\n')
            for deviant_counts_index in range(len(gene_counts)):
                percentage_deviant = (gene_counts.iloc[deviant_counts_index] / len(cleaned_sample_ids)) * 100
                gene = gene_counts.index[deviant_counts_index]
                deviations_number = gene_counts.iloc[deviant_counts_index]
                diag.write(f'{gene}\t{deviations_number}\t{percentage_deviant:.2f}%\n')
            diag.write('\n\n')

        # Remove file
        if os.path.exists(file_path):
            os.remove(file_path)

    @staticmethod
    def check_sex(genotype_df):
        df = genotype_df
        # Make sex_metrics a tidy categorical (saves memory & speeds comparisons)
        if "sex_metrics" in df.columns:
            df["sex_metrics"] = (
                df["sex_metrics"]
                .astype("string")
                .str.strip()
                .astype(pd.CategoricalDtype(categories=["male", "female"], ordered=False))
            )

        # Keep only rows with a definite male/female call:
        if "sex_metrics" in df.columns:
            df = df[df["sex_metrics"].isin(["male", "female"])].reset_index(drop=True)

        # identify sample columns
        non_sample = {"probeset_id", "dbSNP_RS_ID", "sex_metrics"}
        sample_cols = [c for c in df.columns if c not in non_sample]

        # Work with a clean string view of the genotype calls
        samples = df[sample_cols].astype("string")
        sex = df["sex_metrics"].astype("string").str.strip().str.lower()

        # row masks by locus type (as you described)
        male_rows = sex.eq("male")
        female_rows = sex.eq("female")

        # vectorised pattern maps (per column)
        # Single-letter call: A or C or G or T (hemizygous style)
        single_letter = samples.apply(lambda s: s.str.fullmatch(r"[ACGT]"))

        # Diploid call: A/G, A/A, etc.
        diploid = samples.apply(lambda s: s.str.fullmatch(r"[ACGT]/[ACGT]"))

        # "Not present": anything that doesn't contain a base call (incl. NaN/empty/NoCall-style)
        contains_base = samples.apply(lambda s: s.str.contains(r"[ACGT]", na=False))
        not_present = ~contains_base

        n_male = int(male_rows.sum())
        n_female = int(female_rows.sum())

        # Set alpha for certainty threshold
        alpha = 0.05
        percentage = 1 - alpha

        #  per-sample proportions over the relevant rows
        # For a sample to be male:
        #   - "male" loci are single-letter
        #   - AND "female" loci are single-letter (hemizygous-looking on X)
        male_ok = pd.Series(False, index=samples.columns)
        male_loci_ok_rate = ((single_letter | not_present).loc[male_rows].sum() / n_male)
        female_single_rate = (single_letter.loc[female_rows].sum() / n_female)
        male_ok = (male_loci_ok_rate >= percentage) & (female_single_rate >= percentage)

        # For a sample to be female:
        #   - "male" loci are NOT present
        #   - AND "female" loci are diploid
        female_ok = pd.Series(False, index=samples.columns)
        male_not_present_rate = (not_present.loc[male_rows].sum() / n_male)
        female_diploid_rate = (diploid.loc[female_rows].sum() / n_female)
        female_ok = (male_not_present_rate >= percentage) & (female_diploid_rate >= percentage)

        # Final call
        call = pd.Series("unknown", index=samples.columns)
        call.loc[male_ok] = "male"
        call.loc[female_ok] = "female"

        #  Expose the metrics so borderline cases can be audited
        sex_check_df = pd.DataFrame({
            "sample_id": samples.columns,
        })

        # Build dataframe
        remove_cel_suffix = lambda x: x.replace('.CEL_call_code', '')
        sex_check_df['sample_id'] = sex_check_df['sample_id'].apply(remove_cel_suffix)

        sex_check_df["male_loci_present_rate"] = (single_letter.loc[male_rows].sum() / n_male).values
        sex_check_df["female_single_rate"] = (single_letter.loc[female_rows].sum() / n_female).values
        sex_check_df["male_not_present_rate"] = (not_present.loc[male_rows].sum() / n_male).values
        sex_check_df['female_diploid_rate'] = (diploid.loc[female_rows].sum() / n_female).values

        sex_check_df["sex_call"] = call.values

        # Sort by sample for later comparison
        sex_check_df = sex_check_df.sort_values(['sample_id']).reset_index(drop=True)

        return sex_check_df

    @staticmethod
    def compare_sex(customer_data_df, sex_check_df):
        # put your actual column names here
        id_col_a = "sample_id"  # in sex_check_df
        gender_col_a = "sex_call"

        id_col_b = "sample_id"  # in customer_data_df
        gender_col_b = "sex"

        # Normalise genders to the same codes
        def normalise_gender_a(s: pd.Series) -> pd.Series:
            return s.astype(str).str.strip().str.lower().map({
                "male": "M",
                "female": "F"
            })

        def normalise_gender_b(s: pd.Series) -> pd.Series:
            return s.astype(str).str.strip().map({
                "Hr.": "M",
                "Mw.": "F"
            })

        # Merge only on shared IDs
        merged = sex_check_df.merge(
            customer_data_df,
            how="inner",  # keep only IDs present in BOTH
            left_on=id_col_a,
            right_on=id_col_b,
            suffixes=("_df1", "_df2")
        )

        # util.store_dataframe(merged, 'merged')

        # Compare
        merged["checked_sex"] = normalise_gender_a(merged[gender_col_a])
        merged["indicated_sex"] = normalise_gender_b(merged[gender_col_b])

        # Keep only mismatches
        discrepancies = merged[
            merged["checked_sex"] != merged["indicated_sex"]
            ]

        # util.store_dataframe(discrepancies, 'discrepancies')

        # Note sex discrepancies
        with open('Output/Diagnostics/diagnostics.txt', 'a') as diag:
            diag.write("Sex discrepancies:\n")
            if discrepancies.empty:
                diag.write("No discrepancies found.\n")
            else:
                for row in discrepancies.to_dict('records'):
                    if row['checked_sex'] == 'M':
                        male_present = float(row['male_loci_present_rate'])
                        female_single = float(row['female_single_rate'])
                        certainty = fmean([male_present, female_single]) * 100.0
                    else:
                        female_diploid = float(row['female_diploid_rate'])
                        certainty = female_diploid * 100.0
                    diag.write(
                        f"{row['sample_id']} indicated as {'male' if row['indicated_sex'] == 'M' else 'female'}, "
                        f"should be {'male' if row['checked_sex'] == 'M' else 'female'}. {certainty:.2f}% certainty.\n"
                    )
            diag.write('\n\n')
            diag.close()

class InlineDiagnostics:
    def __init__(self):
        self.genes_by_phenotype_pattern = genes_by_phenotype_pattern

        self.genes_by_genotype_pattern = genes_by_genotype_pattern

    @staticmethod
    def generate_combination_deviant_regex(deviant_list):
        """
        Given a list of banned SNPs, return a regex pattern that matches
        genotypes which contain two banned alleles (in any order),
        and negates them so the result only matches normal genotypes.
        """
        # Escape the deviant symbols to neutralize regex metacharacters
        escaped_list = [re.escape(allele) for allele in deviant_list]

        pairs = itertools.combinations_with_replacement(escaped_list, 2)

        genotype_forms = set()
        for a, b in pairs:
            genotype_forms.add(f"{a}/{b}")
            genotype_forms.add(f"{b}/{a}")

        disallowed_regex = '|'.join(genotype_forms)

        # Wrap in negative lookahead
        pattern = rf'^(?!({disallowed_regex})$).*$'

        return pattern

    @staticmethod
    def generate_double_cobination_exclusion_regex(good_list, bad_list):
        """
        Constructs a regex pattern that matches genotypes consisting **only**
        of 'good' alleles and rejects any genotype that includes even a single
        'bad' allele. Genotypes are of the form 'allele1/allele2'.

        Parameters:
        - good_list: list of SNP strings deemed acceptable
        - bad_list: list of SNP strings deemed deviant

        Returns:
        - regex pattern string that matches valid genotypes
        """
        # Escape metacharacters in each allele
        good = [re.escape(snp) for snp in good_list]
        bad = [re.escape(snp) for snp in bad_list]

        # Generate all disallowed combinations:
        # - any pair where at least one allele is from the bad list
        disallowed = set()

        # Bad/Bad combinations
        disallowed.update(f"{a}/{b}" for a, b in itertools.combinations_with_replacement(bad, 2))
        disallowed.update(f"{b}/{a}" for a, b in itertools.combinations(bad, 2))  # Ensure all orders

        # Good/Bad and Bad/Good combinations
        for g in good:
            for b in bad:
                disallowed.add(f"{g}/{b}")
                disallowed.add(f"{b}/{g}")

        # Join disallowed into a regex alternation group
        disallowed_regex = '|'.join(disallowed)

        # Negative lookahead: only match if the entire genotype is NOT in the disallowed list
        pattern = rf'^(?!({disallowed_regex})$).*$'

        return pattern

    # Any combination of two slows, or a fast and a slow should be red
    NAT1_fast_snp = ['*4', '*18', '*20', '*21', '*23', '*24', '*25', '*27', '*29']
    NAT1_slow_snp = ['*11', '*14', '*15', '*17', '*18', '*19', '*22']
    NAT1_pattern = generate_double_cobination_exclusion_regex(NAT1_fast_snp, NAT1_slow_snp)

    NAT2_fast_snp = ['*4', '*18']
    NAT2_slow_snp = ['*5', '*6', '*7', '*10', '*12', '*14', '*17']
    NAT2_pattern = generate_double_cobination_exclusion_regex(NAT2_fast_snp, NAT2_slow_snp)

    genes_by_normal_genotype = {
        'ABCB1': 'C/C',
        'ACE': 'A/A',
        'ADIPOQ': 'G/G',
        'ADRA2A': 'G/G',
        'ALDH2': 'G/G',
        'AMDHD1': 'C/C',
        'BChE': r'^(?!K/K$|A/K$|U/A$|U/F1$|U/F2$|U/H$|U/J$|U/Sc$|K/H$|K/J$|K/Sc$|A/A$|A/F1$|A/F2$|F1/F1$|F1/F2$|F2/F2$|H/H$|H/J$|H/Sc$|J/J$|J/Sc$|Sc/Sc$|K/F2$|K/K\+A$).*$',
        'BCO1': 'A/A',
        'BDNF': r'C/C|Val/Val',
        'CACNA1S': 'WT/WT',
        'CFTR': 'WT/WT',
        'CYP1A1': 'T/T',
        'CYP1B1': r'^(?!.*\*(2|3)).*$',
        'CYP2A6': r'^(?!.*(?:\*4|\*12|\*27|34|\*47|\*53)).*$',
        'CYP2C8': r'^(?!.*\*(2|3|4)).*$',
        'CYP2E1': r'^(?!.*\*(5B)).*$',
        'CYP2F1': r'^(?!.*\*(2|3|4|6)).*$',
        'CYP2R1': r'A/A',
        'CYP4F2': r'^(?!.*\*3).*$',
        'CYP17A1': 'A/A',
        'CYP24A1': 'T/T',
        'DHCR7 / NADSYN1': 'G/G',
        'DRD2': 'C/C',
        'F2': 'G/G',
        'F5': 'C/C',
        'FTO': 'G/G',
        'GC': 'T/T',
        'GCK, YKT6': 'G/G',
        'GSTP1': r'\*A/\*A',
        'GSTT1': r'Null/\*A|\*A/\*A|\*A/Null',
        'GSTM1': r'Null/Present|Present/Null|Present/Present',
        'HLA-B*3101': 'WT/WT',
        'IFNL3/IL28B': r'^(?!T/C|T/T$).*$',
        'IGF1': 'G/G',
        'LDLR': 'G/G',
        'LOC105447645; FUT2': 'A/A',
        'MAO-B': 'T/T|T',
        'MC4R': 'T/T',
        'MnSOD': r'A/A|Val/Val',
        'MTHFR1298': 'A/A',
        'MTRNR1': r'm.=/Null|m.=/m.=|Null/m.=',
        'MTNR1B': 'C/C',
        'NAT1': NAT1_pattern,
        'NAT2': NAT2_pattern,
        'NBPF3': 'T/T',
        'NQ01': 'G/G',
        'OPRM1': 'A/A',
        'PON1': 'T/T',
        'RYR1': r'^(?!.*\u200B).*$',
        'Sult1A1': r'\*1/\*1',
        'Sult1E1': 'C/C',
        'TCF7L2': 'C/C',
        'TMEM165; CLOCK': 'A/A',
        'TNFa': 'G/G',
        'UCP2': 'G/G',
        'VDR': 'WT/WT'
    }
    def is_fenotype_deviation(self, phenotype, gene):
        for pattern in self.genes_by_phenotype_pattern.keys():
            if re.fullmatch(pattern, phenotype): # Als hier een error op is, is er waarschijnlijk geen fenotype ingevuld in de unknowns
                if gene in self.genes_by_phenotype_pattern[pattern][0]:
                    return not phenotype == self.genes_by_phenotype_pattern[pattern][1]

    def is_genotype_deviation(self, genotype, gene_to_check):
        regex_pattern = self.genes_by_normal_genotype[gene_to_check]
        return not re.fullmatch(regex_pattern, genotype)
    """
    Als deze functie een error geeft, is er waarschijnlijk iets verkeerds in de unknown file gezet.
    Iets wat leeg is wat niet leeg zou moeten zijn, of een andere rare input.
    """

    def is_customer_data_present(self,document_table_row):
        for cell_number in [1,3,5]:
            customer_data = document_table_row.cells[cell_number].text
            if not re.search(r'\w', customer_data):
                return False
        return True