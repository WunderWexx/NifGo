# This is where unknown handling is done, including both unknown detection and replacement.

import pandas as pd
import PySimpleGUI as sg
from docx import Document
import Utilities as util


class HandlingUnknowns:
    def __init__(self, dataframe):
        self.dataframe = dataframe
        self.unknown_signs = ['ERROR', 'MISSING', 'Not_PM', 'Not_NM', 'Not_IM', 'Not_RM', 'Not_Determined',
                              'Not_UM', 'EM', 'unknown', '---', '']
        # Get list of all genes to be reported
        self.farmacogenetic_genes = []
        most_recent_farmacogenetisch = util.get_most_recent_template('Farmacogenetisch')
        doc = Document(f'Input/Templates/{most_recent_farmacogenetisch}')
        table1 = doc.tables[1]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            self.farmacogenetic_genes.append(gene)

        self.nutrigenomic_genes = []
        most_recent_nutrigenomics = util.get_most_recent_template('Nutrigenomics')
        doc = Document(f'Input/Templates/{most_recent_nutrigenomics}')
        table1 = doc.tables[1]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            self.nutrigenomic_genes.append(gene)
        table2 = doc.tables[2]
        for row in table2.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            self.nutrigenomic_genes.append(gene)

        self.complete_genes = self.farmacogenetic_genes + self.nutrigenomic_genes

    def detect_unknowns(self):
        # Detect unknown phenotypes and genotypes
        unknowns_df = self.dataframe[
            self.dataframe['phenotype'].isin(self.unknown_signs) |
            self.dataframe['genotype'].isin(self.unknown_signs)
            ].copy()

        # Detect missing genes
        # Create a mapping of sample_id to existing genes
        genes_by_sample = self.dataframe.groupby('sample_id')['gene'].apply(set)

        # Identify missing genes for each sample
        missing_genes_records = [
            {'sample_id': sample_id, 'gene': gene, 'phenotype': 'MISSING', 'genotype': 'MISSING'}
            for sample_id, genes_present in genes_by_sample.items()
            for gene in set(self.complete_genes) - genes_present
        ]

        # Create a DataFrame of missing genes
        missing_df = pd.DataFrame(missing_genes_records)

        # Combine unknowns and missing data
        unknowns_df = pd.concat([unknowns_df, missing_df], ignore_index=True)

        # Delete all rows from the dataframe where the phenotype does not have to be determined, but causes unknown
        most_recent_template = util.get_most_recent_template('Nutrigenomics')
        doc = Document(f'Input/Templates/{most_recent_template}')
        non_phenotype_genes = []
        for table_number in range(1,3):
            table = doc.tables[table_number]
            for row in table.rows[1:]:
                gene = row.cells[0].paragraphs[0].text
                non_phenotype_genes.append(gene)

        to_remove_mask = (
                unknowns_df['gene'].isin(non_phenotype_genes) &
                unknowns_df['phenotype'].isin(self.unknown_signs) &
                ~unknowns_df['genotype'].isin(self.unknown_signs)
        )
        unknowns_df = unknowns_df[~to_remove_mask]

        # Save unknowns to Excel
        unknowns_df.to_excel('Output/Diagnostics/unknowns.xlsx')

    def ask_for_unknowns_df(self):
        file_is_present = sg.popup_yes_no("Heeft u het bestand met de gecorrigeerde unknowns?")
        if file_is_present == 'Yes':
            correcties_path = sg.popup_get_file(
                "Selecteer alstublieft het .xlsx (Excel) bestand met de unknown correcties",
                title="Invoer unknown correctie")
            correcties_df = pd.read_excel(correcties_path)
            return correcties_df
        else:
            return None

    def correct_unknowns(self):
        # correct_unknowns_df is the hand-corrected version of unknowns_df
        # faulty_df is the copy of complete.xlsx that is being edited
        correct_unknowns_df = self.ask_for_unknowns_df()

        if isinstance(correct_unknowns_df, pd.DataFrame):
            faulty_df = self.dataframe

            # Check if all sample_id / gene combinations are unique
            duplicates = correct_unknowns_df.duplicated(subset=['sample_id', 'gene'], keep=False)
            duplicates_df = correct_unknowns_df[duplicates].sort_values(by=['sample_id', 'gene']).drop_duplicates()
            # If not, notify user and terminate program
            if not duplicates_df.empty:
                non_duplicates_df = correct_unknowns_df[~duplicates]
                correct_unknowns_df = pd.concat([non_duplicates_df, duplicates_df], ignore_index=True)
                correct_unknowns_df.to_excel('Output/Dataframes/duplicate_unknowns.xlsx')
                print('Niet-unieke combinaties van sample_id en gene zijn aangetroffen in de unknowns.xlsx.\n'
                      'Verwijder de duplicaten alvorens het programma opnieuw te draaien.')
                util.printEntire(duplicates_df)
                exit('non-unique multi-index')

            # Create a unique key for efficient matching
            correct_unknowns_df['key'] = correct_unknowns_df['sample_id'].astype(str) + "_" + correct_unknowns_df[
                'gene']
            faulty_df['key'] = faulty_df['sample_id'].astype(str) + "_" + faulty_df['gene']

            # Separate records into updates and new additions
            to_update = correct_unknowns_df[correct_unknowns_df['key'].isin(faulty_df['key'])]
            to_add = correct_unknowns_df[~correct_unknowns_df['key'].isin(faulty_df['key'])]

            # Create a mapping of sample_id and gene to phenotype and genotype from to_update
            update_mapping = to_update.set_index(['sample_id', 'gene'])[['phenotype', 'genotype']]

            # Update the phenotype and genotype columns in faulty_df based on the mapping
            faulty_df.set_index(['sample_id', 'gene'], inplace=True)
            faulty_df.update(update_mapping)
            faulty_df.reset_index(inplace=True)

            # Add new records to the dataframe
            if not to_add.empty:
                to_add = to_add[['sample_id','gene','phenotype','genotype']]  # Drop the temporary key column
                faulty_df = pd.concat([faulty_df, to_add], ignore_index=True)

            # Remove the temporary key column from the main dataframe
            faulty_df = faulty_df.drop(columns=['key'])

            # Update the class attribute
            self.dataframe = faulty_df
