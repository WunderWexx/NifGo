# This is where the pharmacist info sheet is generated

# Imports
from docx import Document
import Utilities as util
from pathlib import Path
from Diagnostics import InlineDiagnostics

class InfoSheet():
    def __init__(self, sample_id, dataframe, customer_data=None):
        # Benodigde data vaststellen
        self.sample_id = sample_id
        self.infosheet_df = dataframe[dataframe['sample_id'] == sample_id]

        # Als beschikbaar wordt de klantdata vastgesteld
        self.fullname = ''
        self.birthdate = ''

        if customer_data is not None:
            self.customer_data_presence = True
            customer_row = customer_data.loc[customer_data['sample_id'] == sample_id]

            if not customer_row.empty:
                initials = customer_row.get('initials', '').iloc[0] or ''
                lastname = customer_row.get('lastname', '').iloc[0]  # Assumes lastname is always present
                self.fullname = f"{initials} {lastname}".strip()

                birthdate = customer_row.get('birthdate', '').iloc[0]
                self.birthdate = '' if birthdate == '20237-01-01' else (birthdate or '')

    def report_generation(self):
        # Load the Word document
        most_recent_template = util.get_most_recent_template('InfoSheet')
        doc = Document(f'Input/Templates/{most_recent_template}')

        # Fill customer data table
        name_cell = doc.tables[0].rows[0].cells[1]
        bdate_cell = doc.tables[0].rows[0].cells[3]
        sample_code_cell = doc.tables[0].rows[0].cells[5]
        util.change_table_cell(name_cell, change_text=f'{self.fullname}', font_size=10)
        util.change_table_cell(bdate_cell, change_text=f'{self.birthdate}', font_size=10)
        util.change_table_cell(sample_code_cell, change_text=f'{self.sample_id}', font_size=10)
        if self.customer_data_presence:
            if not InlineDiagnostics().is_customer_data_present(doc.tables[0].rows[0]):
                print(f'WARNING: {self.sample_id} is missing customer data!')

        #Fill results table
        table1 = doc.tables[1]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            phenotype_cell = row.cells[1]
            result_cell = row.cells[2]
            result = self.infosheet_df.loc[self.infosheet_df['gene'] == gene, 'genotype'].values
            result = result[0] if result.size > 0 else ''
            util.change_table_cell(result_cell, change_text=f' {result}', font_size=9)
            phenotype = self.infosheet_df.loc[self.infosheet_df['gene'] == gene, 'phenotype'].values
            phenotype = phenotype[0] if phenotype.size > 0 else ''
            util.change_table_cell(phenotype_cell, change_text=f' {phenotype}', font_size=9)
            if InlineDiagnostics().is_fenotype_deviation(phenotype, gene):
                util.change_table_row(row, font_color='FF0000')

        #Save the document
        document_name = 'InfoSheet' + "_" + self.sample_id + ".docx"
        path = Path("Output/Reports/")
        path.mkdir(parents=True, exist_ok=True)
        doc.save(f"Output\\Reports\\{document_name}")