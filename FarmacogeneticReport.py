"""This is where the farmacogenetic report is generated from a template"""

from docx import Document
import Utilities as util
from pathlib import Path

class farmacogenetic_report:
    def __init__(self, sample_id, dataframe):
        self.sample_id = sample_id
        self.farmacogenetic_df = dataframe[dataframe['sample_id'] == sample_id]

    def report_generation(self):
        # Load the Word document
        doc = Document('Input/Templates/Farmacogenetisch-2024-12.docx')

        #Fill customer data table with sample code
        sample_code_cell = doc.tables[0].rows[0].cells[5]
        util.change_table_cell(sample_code_cell, bold=True, change_text=f'{self.sample_id}', font_size=10)

        #Fill results table 1
        table1 = doc.tables[1]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            phenotype_cell = row.cells[1]
            result_cell = row.cells[2]
            result = self.farmacogenetic_df.loc[self.farmacogenetic_df['gene'] == gene, 'genotype'].values
            result = result[0] if result.size > 0 else ''
            util.change_table_cell(result_cell, change_text=f' {result}', font_size=9)
            phenotype = self.farmacogenetic_df.loc[self.farmacogenetic_df['gene'] == gene, 'phenotype'].values
            phenotype = phenotype[0] if phenotype.size > 0 else ''
            util.change_table_cell(phenotype_cell, change_text=f' {phenotype}', font_size=9)

        #Save the document
        document_name = 'FarmacogeneticReport' + "_" + self.sample_id + ".docx"
        path = Path("Output/Test/")
        path.mkdir(parents=True, exist_ok=True)
        doc.save(f"Output\\Test\\{document_name}")