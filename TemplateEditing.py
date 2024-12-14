from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load the Word document
doc = Document('Miscellaneous/Pangenomix/Nutrigenomics_2025.docx')

# Function to format the cell
def format_cell(cell, text):
    # Set text content
    cell.text = text

    # Format the first paragraph in the cell
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align left horizontally

    # Set vertical alignment to center
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

    # Make text bold, white, and size 9
    run = paragraph.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    run.font.size = Pt(9)  # Set font size to 9z

# Iterate through all tables and cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            format_cell(cell, "Sample Text")

# Function to find and replace text in text boxes
def replace_text_in_shapes(doc, replacements):
    for shape in doc.element.body.xpath('.//w:sdtContent'):
        for paragraph in shape.xpath('.//w:p'):
            for run in paragraph.xpath('.//w:r'):
                text = run.xpath('./w:t')
                if text:
                    text_value = text[0].text
                    # Replace text if it's a match
                    for old_text, new_text in replacements.items():
                        if old_text in text_value:
                            text[0].text = text_value.replace(old_text, new_text)

# Define replacements for the text boxes
replacements = {
    "Naam: . . .": "Naam: John Doe",
    "Geb. datum: . . .": "Geb. datum: 1985-10-23",
    "Code: . . .": "Code: CUST12345"
}

# Replace text in the document
replace_text_in_shapes(doc, replacements)

# Save the updated document
doc.save('Miscellaneous/Pangenomix/Nutrigenomics_2025_tested.docx')

#----------------------------------------------------------------------------

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Load the document
doc_path = "Miscellaneous/Pangenomix/Nutrigenomics_2025_edited.docx"
doc = Document(doc_path)

# Create a paragraph with a hyperlink to the bookmark
def add_internal_hyperlink(paragraph, text, bookmark_name):
    """
    Add an internal hyperlink to a paragraph.
    :param paragraph: The paragraph to which the hyperlink will be added.
    :param text: The display text for the hyperlink.
    :param bookmark_name: The name of the bookmark to link to.
    """
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    # Create the run element (the text to display)
    run = OxmlElement('w:r')
    text_element = OxmlElement('w:t')
    text_element.text = text
    run.append(text_element)

    # Add the run to the hyperlink
    hyperlink.append(run)

    # Append the hyperlink to the paragraph
    paragraph._element.append(hyperlink)

# Add a paragraph and an internal hyperlink
paragraph = doc.add_paragraph()
add_internal_hyperlink(paragraph, "Go to BChE section", "BChE") # A hyperlink with name BChE should first manually be made in the source document.

# Save the document
output_path = "Miscellaneous/Pangenomix/Nutrigenomics_2025_edited_updated.docx"
doc.save(output_path)

print(f"Document saved with an internal hyperlink at: {output_path}")