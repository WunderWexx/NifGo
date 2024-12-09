# This is where the nutrinomics report is generated

# Imports
from docx.shared import Pt, Cm
from WordDocument import WordEditing as wd
import ELT
from datetime import date

class NutrinomicsReport(wd):
    def __init__(self, sample_id, dataframe):
        super().__init__(sample_id, dataframe)
        self.nutri_markers = ELT.Extract().nutrimarkers()

    def nutrinomics_IA(self):

        nutrinomics_genes = [
            'ABCB1',
            'ABCG2',
            'ACE',
            'ADIPOQ',
            'ADRA2A',
            'ALDH2',
            'AMDHD1',
            'BCO1',
            'BChE',
            'BDNF-AS; BDNF',
            'COMT',
            'CYP17A',
            'CYP24A1',
            'CYP2R1',
            'DHCR7',
            'DRD2',
            'F2',
            'F5',
            'FTO',
            'GC',
            'GCK, YKT6',
            'GSTP1',
            'HLA-B*3101',
            'IFNL3/IL28B',
            'IGF1',
            'LDLR',
            'LOC105447645; FUT2',
            'MAO-B',
            'MC4R',
            'MTHFR 1298A>C',
            'MTHFRC677T',
            'MTNR1B',
            'MnSOD',
            'NADSYN1',
            'NBPF3',
            'NQ01',
            'OPRM1',
            'PON1',
            'SLCO1B1',
            'Sult1A1',
            'Sult1E1',
            'TCF7L2',
            'TMEM165; CLOCK',
            'TNFa',
            'UCP2',
            'VDR',
            'VKORC1'
        ]

        nutri_markers = self.nutri_markers[self.nutri_markers['gene_name'].isin(nutrinomics_genes)]

        df = self.dataframe[self.dataframe['sample_id'] == self.sample_id]
        df = df[df['gene'].isin(nutrinomics_genes)]
        df.drop(['sample_id','phenotype'], axis='columns', inplace=True)
        df.rename(columns={'gene': 'Gen', 'genotype': 'Uitslag'}, inplace=True)

        dbSNP_list = []
        for gen in nutrinomics_genes:
            index = nutri_markers[nutri_markers['gene_name'] == gen].index.values
            dbSNP_list.append(nutri_markers.loc[index, 'dbSNP_RS_ID'].item())

        df.insert(1, 'dbSNP', dbSNP_list)
        return df

    def inleiding(self):
        """
        Generates block 1 of all standard text.
        :return:
        """
        self.document.add_picture("Input/Icons/NifGo_logo.png")
        self.heading("\nFarmacogenetisch Rapport")
        self.colour_bar()
        run = self.document.add_paragraph().add_run(
            "Dit farmacogenetisch rapport geeft een analyse van het DNA en identificeert de relevante genetische variaties en hun effecten op de veiligheid en werkzaamheid van medicijnen. U bent getest op die genen, die de werkzaamheid van uw medicatie kunnen beïnvloeden. Het DNA is geïsoleerd uit speeksel.\n\n"
            "Dit rapport mag niet worden gebruikt om medicatie te veranderen, zonder begeleiding van een arts of apotheker. Raadpleeg altijd de (huis)arts en/of apotheker en wijzig nooit zelf de voorgeschreven medicijnen.\n\n"
            "NIFGO adviseert nadrukkelijk om de uitslag van het onderzoek te laten vastleggen in het medische dossier van de (huis)arts. Vraag ook de apotheker om deze gegevens te registreren in het informatiesysteem van de apotheek. Apothekers registreren deze 18 belangrijkste genen in hun informatiesysteem. Na registratie bent u er zeker van, dat de apotheker controleert of de voorgeschreven medicijnen wel passen bij de uitslag van dit onderzoek.\n\n"
            "Het rapport is persoonlijk en is opgesteld op basis van de huidige kennis en stand van zaken. Daarbij wordt nadrukkelijk het voorbehoud gemaakt, dat het DNA-materiaal van de hier genoemde persoon is."
        )
        run.font.name = "Calibri"
        run.font.size = Pt(12)

    def id_table(self):
        """
        Generates the table containing the customer data.
        :return:
        """
        self.colour_bar()
        customer_table = self.document.add_table(rows=1, cols=6)  # style = "Table Grid"
        customerDatatypes = ["Naam", "Geb. datum", "Code"]
        for type, column in zip(customerDatatypes, range(0, 5, 2)):
            paragraph = customer_table.cell(0, column).paragraphs[0]
            run = paragraph.add_run(type)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(12)
        paragraph = customer_table.cell(0, 5).paragraphs[0]
        run = paragraph.add_run(self.sample_id)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        customer_table.allow_autofit = False
        width_dict = {0: 1.30, 1: 6, 2: 2.5, 3: 2.5, 4: 1.25, 5: 2.78}
        for i in range(5):
            for cell in customer_table.columns[i].cells:
                cell.width = Cm(width_dict[i])
        self.colour_bar()

    def inhoudsopgave(self):
        """
        Maakt de inhoudsopgave
        :return:
        """
        self.heading("\nInhoudsopgave")
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run("- Overzicht uitslag\n\n"
                                "- Toelichting uitslag\n\n"
                                "- Variaties waarop is getest\n\n")
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

    def logo_titel_header(self):
        self.document.add_picture("Input/Icons/niFGo_logo.png")
        self.heading("\nFarmacogenetisch Rapport NUTRI {}".format(date.today()))
        self.heading("Bijlage nutrigenomics {}".format(self.sample_id))
        self.colour_bar()
        self.heading("De uitslag nutrigenomics:",lined=True)

    def table(self, style = "Table Grid"):
        """
        Creates a Word table from a dataframe
        :param: style: The style of the table, default is Table Grid.
        :return: A Word table in the specified document
        """

        df = self.nutrinomics_IA()

        # add a table to the end and create a reference variable
        # extra row is so we can add the header row
        t = self.document.add_table(df.shape[0] + 1, df.shape[1])
        t.style = style

        # add the header rows.
        for j in range(df.shape[-1]):
            header_cells = t.rows[0].cells
            paragraph = header_cells[j].paragraphs[0]
            run = paragraph.add_run(df.columns[j])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                normal_cells = t.rows[i+1].cells
                paragraph = normal_cells[j].paragraphs[0]
                run = paragraph.add_run(str(df.values[i, j]))
                run.font.name = 'Calibri'
                run.font.size = Pt(12)

        t.allow_autofit = False
        width_dict = {0:5.49, 1:3.75 ,2:1.59} #{0:4, 1:3.5 ,2:2 ,3:10} when frequencies are included
        for i in range(len(width_dict)):
            for cell in t.columns[i].cells:
                cell.width = Cm(width_dict[i])

    def Toelichting(self):
        self.document.add_page_break()
        self.heading("Toelichting\n", lined=True)

        def first_word_bold(text,chosen_font='Calibri', chosen_size=11):
            split_text = text.split(sep=" ", maxsplit=1)
            bold_word = split_text[0]
            rest = split_text[1]
            paragraph = self.document.add_paragraph()
            self.linepacing(paragraph,spacing=1)
            run = paragraph.add_run("{} ".format(bold_word))
            run.bold = True
            run.font.name = chosen_font
            run.font.size = Pt(chosen_size)
            run = paragraph.add_run(rest)
            run.bold = False    #Niet nodig om dit te declareren, maar wel fijn voor de leesbaarheid
            run.font.name = chosen_font
            run.font.size = Pt(chosen_size)

        first_word_bold("ABCB1 is betrokken bij de transportfunctie van cellen. Het enzym, P-glycoproteïne, fungeert als een soort pomp in onze cellen, die bepaalde stoffen uit de cel kan pompen, voordat ze zich kunnen ophopen in het lichaam en schade kunnen veroorzaken. Langs deze weg is ABCB1 dus ook betrokken bij de manier, waarop ons lichaam medicijnen opneemt, verdeelt, afbreekt en uitscheidt. Variaties in het gen ABCB1 kunnen dus invloed hebben op hoe we reageren op bepaalde medicijnen, zoals antidepressiva, antistollingsmedicijnen, maagbeschermers, antibiotica, anti-epileptica en cholesterolverlagers.  Het ABCB1 gen speelt ook een belangrijke rol in de regulering van de bloed-hersenbarrière. Deze barrière controleert, welke stoffen vanuit de bloedbaan de hersenen kunnen bereiken. Variaties in ABCB1 hebben dus ook invloed op de werking van deze barrière. In dit onderzoek is de variatie rs1045642 getest. T geeft de aanwezigheid van de variatie aan. De meest voorkomende testuitslag bij mensen van West-Europese afkomst is C/T (heterozygoot). De variatie rs1045642 is geassocieerd met een lagere expressie van P-glycoproteïne. Afhankelijk van de gebruikte software wordt als uitslag ook een A in plaats van een T aangegeven, als de variatie rs1045642 is aangetoond. De meest voorkomende testuitslag is dan dus A/G in plaats van C/T. Voor het bepalen van het fenotype maakt het echter niet uit. De geteste variatie rs1045642 is van invloed op de werking van Digoxine, Dabigatran, Amlodipine en Cannabis. Van Sint-Janskruid is bekend dat gebruik ervan de functie van ABCB1 stimuleert.")
        first_word_bold("ACE codeert voor de aanmaak van angiotensine II. angiotensine II is een hormoon, dat vernauwing van de bloedvaten veroorzaakt. Het ACE-enzym heeft twee functies: namelijk ten eerste bevordert het de afgifte van een hormoon genaamd aldosteron, dat de balans van natrium en kalium in de nieren reguleert. Ten tweede breekt het enzym bradykinine af, een stof die de bloedvaten verwijdt en de bloeddruk kan verlagen. In dit onderzoek is de variatie rs4341 getest. C geeft de aanwezigheid van de variatie aan. C duidt de typering Deletie (D) aan. Deze typering heeft invloed op de productie van het ACE-enzym. De testuitslag C wordt geassocieerd met een lagere ACE-enzymactiviteit. Studies hebben aangetoond dat mensen met een testuitslag C een hogere bloeddruk hebben dan mensen met een testuitslag G.")
        first_word_bold("ADIPOQ bepaalt hoe ons lichaam omgaat met suiker en vet. Het adiponectine enzym is van invloed op de insulinegevoeligheid. Een lage adiponectine concentratie wordt geassocieerd met een verminderde insulinegevoeligheid. Daarnaast stimuleert adiponectine de verbranding van vetzuren in onze spierweefsels, wat bijdraagt aan een gezond vetmetabolisme. Ook regelt adiponectine onze eetlust en is van invloed op ons verzadigingsgevoel en kan op die manier van invloed zijn op het lichaamsgewicht en de voedselopname. Adiponectine heeft ook ontstekingsremmende eigenschappen. In dit onderzoek is de variatie rs17300539 getest. A geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een lagere expressie van adiponectine. Een lagere expressie verhoogt het risico op insulineresistentie. De testuitslag G/G is het meest voorkomend (86%) onder de West-Europese bevolking.")
        first_word_bold("ADRA2A is van invloed op hoe ons zenuwstelsel reageert op stress. Deze receptor komt voor in verschillende weefsels en organen, waaronder de hersenen, het hart en de bloedvaten. In de hart en de bloedvaten kan activering van de receptor leiden tot vernauwing van de bloedvaten en verhoging van de bloeddruk. In de hersenen reguleert de α2A-adrenerge receptor de afgifte van neurotransmitters, zoals noradrenaline en dopamine. Dit heeft invloed op onze stemming, aandacht en cognitieve functies. De α2A-adrenerge receptor is ook aanwezig op de bètacellen van de alvleesklier, waar insuline wordt geproduceerd. De activering van deze receptor kan de afgifte van insuline beïnvloeden en daarmee de regulering van de bloedsuikerspiegel. In dit onderzoek is de variatie rs10885122 getest. T geeft de aanwezigheid van de variatie aan. De testuitslag G/G is het meest voorkomend onder de West-Europese bevolking.")
        first_word_bold("ALDH2 speelt een belangrijke rol bij de afbraak van ethanol, de alcohol, die we drinken. Variaties in het ALDH2-gen kunnen leiden tot verminderde of afwezige activiteit van het enzym, wat kan leiden tot alcoholintolerantie. In dit onderzoek is de variatie rs671 getest. A geeft de aanwezigheid van de variatie aan. De testuitslag G/G is het meest voorkomend onder de West-Europese bevolking. De testuitslagen A/G en A/A geven een verminderde werking van het ALDH2-enzym aan en kunnen dus van invloed zijn op hoe goed iemand alcohol kan verwerken in vergelijking met de testuitslag G/G.")
        first_word_bold("AMDHD1 is betrokken bij het afbreken van bepaalde aminozuren, specifiek in de route waarbij L-histidine wordt omgezet in L-glutamaat. In dit onderzoek is de variatie rs10745742 getest. Deze variatie wordt in verband gebracht met veranderingen in de plasmawaarden van calcidiol (25-hydroxyvitamine D) en dus met een verminderde activiteit van AMDHD1. De testuitslag C is het wildtype. De testuitslag T geeft de variatie aan. De verminderde activiteit van AMDHD1 (testuitslag C/T en T/T) kan resulteren in een verhoogde concentratie van metabolieten, wat mogelijk invloed heeft op de effectiviteit en bijwerkingen van medicijnen. Het kennen van de verminderde activiteit van AMDHD1 is van belang bij het gebruik van antidepressiva, antipsychotica en opioïden.")
        first_word_bold("BChE is belangrijk bij het snel afbreken van bepaalde neurotransmitters, zoals acetylcholine. Acetylcholine is essentieel voor de communicatie tussen zenuwcellen en spieren. Daarnaast is BChE ook betrokken bij het ontgiften van lichaamsvreemde stoffen, waaronder dus ook medicijnen. Een verminderde activiteit van BChE kan bij gebruik van spierverslappers en anesthetica leiden tot sterke bijwerkingen. Een verhoogde activiteit van BChE kan leiden tot een te snelle afbraak van medicijnen met als gevolg, dat het gewenste resultaat niet wordt bereikt. Er zijn verschillende varianten van het BChE gen, die invloed kunnen hebben op de activiteit. De K variatie (rs1803274; C is wildtype en T de variatie) heeft een verminderde activiteit en dus een verhoogde gevoeligheid voor bepaalde medicijnen. Daarnaast zijn er nog andere variaties die van invloed zijn op het metabolisme van lichaamsvreemde stoffen. In dit onderzoek zijn ook de variaties rs1799807 (testuitslag T is wildtype, de variatie C komt zelden voor onder de West-Europese bevolking), rs28933389 (de testuitslag A geeft de variatie aan en G het wildtype) en rs28933390 (de testuitslag C is het wildtype, de testuitslag A geeft de variatie aan) getest. Alle drie variaties leiden tot een verminderde activiteit van BChE. Bij een lage BChE-activiteit worden succinylcholine en mivacurium minder snel afgebroken, wat kan leiden tot langdurige spierverslapping na een operatie. Een lagere dosering kan dan overwogen worden.")
        first_word_bold("BCO1 speelt een belangrijke rol bij de omzetting van bètacaroteen in vitamine A. Het BCO1-enzym versnelt het proces, waarbij bètacaroteen wordt omgezet in vitamine A (ook wel retinol genoemd). In dit onderzoek is de variatie rs7501331 getest. T geeft de aanwezigheid van de variatie aan. De testuitslagen C/T en T/T geven een mogelijk verhoogd risico op een vitamine A tekort aan. Mensen met de testuitslag C/T en T/T kunnen ook minder goed vetten afbreken dan mensen met een testuitslag C/C.")
        first_word_bold("BDNF-AS is essentieel voor een goede werking van de hersenen en wordt geassocieerd met zaken zoals leerprocessen en geheugen. Variaties in BDNF, inclusief de mogelijke rol van BDNF-AS, kunnen invloed hebben op de niveaus van BDNF in de hersenen en worden geassocieerd met neurologische en neuro psychiatrische aandoeningen. In dit onderzoek is de variatie rs6265 getest. Deze variatie wordt aangegeven als het Val66Met-polymorfisme. De testuitslag Met is de variatie. Onderzoek heeft aangetoond dat het Val66Met-polymorfisme de afgifte van BDNF kan beïnvloeden. Het Met66-allel wordt in verband gebracht met de kans gevoeliger te zijn voor psychiatrische aandoeningen, zoals depressie, angststoornissen, schizofrenie en de ziekte van Alzheimer. De testuitslag Val/Met wordt geassocieerd met een licht risico op ADHD. Bij weergeven van de uitslag wordt voor Val “C” en voor Met “T” gebruikt. ")
        first_word_bold("CACNA1S codeert voor de structuur en functie van calciumkanalen, die worden aangetroffen in de skeletspieren. Deze kanalen activeren de werking van RYR1-kanalen. Variaties in het CACNA1S-gen kunnen een verhoogd risico geven op maligne hyperthermie, een ernstige reactie op bepaalde anesthetica, die vaak worden gebruikt tijdens operaties en andere invasieve procedures.")
        first_word_bold("CFTR codeert voor een ionkanaalenzym van de ABC-transportenzymen. Variaties van het CFTR-gen kunnen leiden tot ontregeling van het epitheelvochttransport in de longen, pancreas en andere organen. Dit geeft risico op cystische fibrose.")
        first_word_bold("CYP1A1 is van invloed op het ontgiften en afvoeren van lichaamsvreemde stoffen, zoals medicijnen. In dit onderzoek is de variatie rs1048943 getest. Deze variatie kan leiden tot een verandering in de structuur en functie van CYP1A1 met als gevolg een verhoogde of verminderde activiteit. De testuitslag Ile/Val en Val/Val is geassocieerd met een verhoogde activiteit en dus een snellere afbraak van medicijnen vergeleken met de uitslag Ile/Ile. De testuitslag T/T geeft het wildtype aan. De testuitslag T/C en C/C komen vaker voor bij Aziaten. De effecten van deze variatie kunnen worden versterkt door omgevingsfactoren, bijvoorbeeld roken.")
        first_word_bold("CYP1B1 speelt een belangrijke rol in het metabolisme van lichaamsvreemde stoffen, endogene stoffen zoals hormonen (steroïden, oestrogenen) en polycyclische aromatische koolwaterstoffen (PAK's). In dit onderzoek is de variatie rs1056836 getest. De testuitslag Leu/Leu (C/C) Leu/Val (C/G) zijn geassocieerd met een verhoogde activiteit van CYP1B1 vergeleken met de Val/Val (G/G) uitslag. Dit betekent dat CYP1B1 sneller en mogelijk effectiever is bij de afbraak van oestrogenen met als gevolg meer productie van 4-hydroxy-oestrogeen met het risico op DNA-schade. De Leu-testuitslag (C/C) komt voor 20% voor onder de West-Europese bevolking.")
        first_word_bold("CYP2A6 is verantwoordelijk voor de oxidatie van nicotine en cotinine. Expressie en activiteit van CYP2A6 wordt ook beïnvloed door niet-genetische factoren. Variaties van het CYP2A6 gen zijn van invloed op de werking van tegafur, letrozol, efavirenz, valproïnezuur, pilocarpine, artemisinine, artesunaat, cafeïne en tyrosol en een aantal cumarine-achtige alkaloïden. Pompelmoes en grapefruit remmen de activiteit van CYP2A6.")
        first_word_bold("CYP2C8 is een belangrijk leverenzym betrokken bij metabolisme van een aantal medicijnen bij behandeling van o.a. kanker en diabetes. CYP2C8 is direct betrokken bij metabolisme van o.a. Ibuprofen. CYP2C8 heeft een indirecte rol bij metabolisme van o.a. Diclofenac. De meest voorkomende variaties, die leiden tot een verlaagd metabolisme zijn *2 en *3.")
        first_word_bold("CYP2E1 de activiteit van CYP2E1 wordt gestimuleerd door alcohol, roken en de medicijnen Isoniazide en Isopropanol. Daarnaast speelt CYP2E1 een rol bij anesthesie, alcoholconsumptie, diabetes en obesitas. Transvetzuren en geoxideerde vetzuren hebben een negatief effect op de activiteit van CYP2E1. CYP2E1 -PM is een risicofactor voor vervetting van de lever.")
        first_word_bold("CYP2F1 is als onderdeel van het P450 cytochroom van invloed op het afbreken en verwijderen van lichaamsvreemde stoffen waaronder medicijnen. Specifieke gevalideerde gegevens over het effect van verschillende variaties zijn niet aanwezig.")
        first_word_bold("CYP2R1 is van invloed op de omzetting van vitamine D naar zijn actieve vorm, calcitriol. Calcitriol speelt een belangrijke rol in het reguleren van calcium- en fosfaatniveaus. In dit onderzoek is de variatie rs10741657 getest. Deze variatie wordt in verband gebracht met veranderingen in de serumwaarden van calcidiol (25-hydroxyvitamine D). De testuitslag A is het wildtype. Mensen met deze testuitslag zouden een betere respons hebben op vitamine D-suppletie. De testuitslag A/G en G/G kunnen leiden tot het sneller afbreken van vitamine D en is van invloed op de calciumhuishouding. Bij deze uitslagen kan overwogen worden vitamine D-supplementen en calcium supplementen te gebruiken. Ook is de variatie rs12794714 getest. De testuitslag G is het wildtype. Mensen met de testuitslag A/A tonen een verhoogde omzetting van vitamine D, wat kan bijdragen aan een betere opname van calcium. Maar in sommige gevallen kan er sprake zijn van een overschot aan vitamine D.")
        first_word_bold("CYP4F2 is betrokken bij het metabolisme van vetzuren en vitamines (E en K). Variaties in het CYP4F2-gen zijn van belang voor het bepalen van de dosering van vitamine K-antagonisten zoals Warfarine, Cumarine en Acenocoumarol. In die zin is er een relatie met de activiteit van VKORC1.")
        first_word_bold("CYP17A speelt een belangrijke rol in de productie van steroïde hormonen, waaronder geslachtshormonen, zoals oestrogenen (estradiol), androgenen(testosteron) en corticosteroïden. In dit onderzoek is de variatie rs743572 getest. De testuitslag G (15% onder de West-Europese bevolking) wordt geassocieerd met een verhoogde expressie van CYP17A1. Dit kan invloed hebben op de balans van geslachtshormonen en steroïden in het lichaam. Een verhoogde activiteit leidt tot hogere niveaus van steroïde hormonen, zoals androgenen en oestrogenen.  Veranderingen in de hormonale balans kunnen invloed hebben op de vruchtbaarheid, dit zowel bij mannen als vrouwen. Deze variatie kan, omdat steroïde hormonen een rol spelen in de stofwisseling, ook de oorzaak zijn van een te hoog lichaamsgewicht, vasthouden van vetten en stofwisselingsaandoeningen. De effecten van deze variatie worden ook beïnvloed door omgevingsfactoren (o.a. leefstijl).")
        first_word_bold("CYP24A1 is van invloed op het reguleren van de hoeveelheid actief vitamine D (calcitriol) en van calcium in het bloed. Calcitriol speelt een belangrijke rol in het reguleren van calcium- en fosfaatniveaus. In dit onderzoek is de variatie rs17216707 getest. Deze variatie wordt in verband gebracht met veranderingen in de plasmawaarden van calcidiol (25-hydroxyvitamine D). De testuitslag C is het wildtype. De testuitslag T geeft de variatie aan. Aanwezigheid van de variatie (testuitslag C/T en T/T) kan leiden tot het sneller afbreken van vitamine D en is van invloed op de calciumhuishouding. Bij deze uitslagen kan overwogen worden vitamine D-supplementen en calcium supplementen te gebruiken.")
        first_word_bold("DHCR7 is betrokken bij de omzetting van 7-dehydrocholesterol naar cholesterol, een belangrijke stap in de cholesterol huishouding. Daarnaast is DHCR7 ook van invloed op de synthese van vitamine D. Een verhoogde activiteit van DHCR7 vermindert de aanmaak van vitamine D via blootstelling aan zonlicht. In dit onderzoek is de variatie rs12785878 getest. De testuitslag T is het wildtype. De testuitslag G geeft de aanwezigheid van de variatie aan. De testuitslag T is geassocieerd met een verhoogde activiteit van DHCR7, wat kan leiden tot efficiëntere omzetting van 7-dehydrocholesterol naar cholesterol. Een verhoogde activiteit van DHCR7 kan ook de synthese van vitamine D bevorderen, vooral bij voldoende blootstelling aan zonlicht. In sommige gevallen kan een te hoge activiteit van DHCR7 ook leiden tot een onbalans in de cholesterolhuishouding, afhankelijk van andere genetische en omgevingsfactoren.")
        first_word_bold("DRD2 is van invloed op de activiteit van de neurotransmitter dopamine en is belangrijk voor een goede werking van ons zenuwstelsel. Het DRD2-gen is betrokken bij verschillende psychiatrische aandoeningen, waaronder schizofrenie, verslavingsstoornissen en ADHD. In dit onderzoek is de variatie rs1076560 getest. A geeft de aanwezigheid van de variatie aan. De testuitslag C/C is de meest voorkomende onder de West-Europese bevolking. De testuitslag A/C wordt geassocieerd met een verhoogd risico op alcoholisme. De testuitslag A/A wordt geassocieerd met een verminderd werkgeheugen.")
        first_word_bold("F2 codeert voor Prothrombine, dat belangrijk is voor de bloedstolling. Onder normale omstandigheden is F2 niet actief (PM). F2 wordt pas actief, wanneer en sprake is van een bloeding. F2 zal dan bijdragen aan de stolling. Variatie in F2 geeft risico op een constante activiteit van het gen, die vervolgens kan leiden tot vorming van onnodig bloedstolsels.")
        first_word_bold("F5 is belangrijk voor de bloedstolling. F5 vormt met andere stollingsfactoren thrombine en fibrine, wat belangrijk is voor het vormen van bloedstolsels. F 5 dient geïnactiveerd (niet actief) te zijn. Dit is het geval bij 97% van de West-Europese bevolking). In dit onderzoek is de variatie rs6025 (Leiden mutatie, T) getest. De testuitslag C is het wildtype. Een aangetoonde aanwezigheid van deze variatie geeft een verhoogd risico op het ontwikkelen van veneuze trombose en longembolie. De Leiden-mutatie is erfelijk. Medicijnen, die vaak worden voorgeschreven bij aanwezigheid van deze variatie zijn Warfarine, Rivaroxaban, Apixaban, Dabigatran en Aspirine.")
        first_word_bold("FTO (Fat mass and obesity-associated gene) speelt een belangrijke rol bij het reguleren van ons energiemetabolisme en heeft invloed op de controle op onze voedselinname en ons lichaamsgewicht. Naast obesitas wordt FTO ook in verband gebracht met andere gezondheidsproblemen, waaronder het risico op het ontwikkelen van type 2-diabetes en hart- en vaatziekten. In dit onderzoek is de variatie rs1121980 getest. A geeft de aanwezigheid van de variatie aan. Dit betekent dat mensen met de variatie A een hogere hoeveelheid FTO in hun lichaam hebben dan mensen met de variatie G. De testuitslag A/G is de meest voorkomende testuitslag (54%) onder de West-Europese bevolking. ")
        first_word_bold("GC speelt een belangrijke rol in het transport van vitamine D door het bloed. Vitamine D is belangrijk voor de calcium- en fosfaatbalans in het lichaam, wat invloed heeft op de gezondheid van botten en het immuunsysteem. Het GC-gen produceert vitamine D-bindend eiwit, dat vitamine D (en zijn metabolieten zoals calcidiol en calcitriol) door het bloed transporteert naar doelorganen. GC zorgt ervoor dat vitamine D in vetweefsel en lever kan worden opgeslagen en vrijgegeven wordt wanneer dat nodig is. Daarnaast speelt GC ook een rol in het immuunsysteem bij de respons op infecties. In dit onderzoek is ook de variatie rs7041 getest. De testuitslag A is het wildtype. De testuitslag C geeft de aanwezigheid van de variatie aan en kan leiden tot minder transport van vitamine D in het bloed en daardoor een lagere beschikbaarheid van biologisch actieve vitamine D. Mensen met een testuitslag C kunnen dus minder effectief reageren op standaard vitamine D-supplementen. Omdat vitamine D belangrijk is voor het opnemen van calcium, kunnen mensen met de testuitslag C een andere respons hebben op behandelingen voor osteoporose, zoals bisfosfonaten (Alendronaat, Risedronaat).")
        first_word_bold("GCK,YKT6 is betrokken is bij de regulering van de bloedglucosespiegel. Glucokinase is ook aanwezig in sommige delen van de hersenen en speelt een rol in de regulering van de eetlust en de glucosespiegels in de hersenen. In dit onderzoek is de variatie rs4607517 getest. A geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde activiteit van glucokinase. Dit betekent dat mensen met de testuitslag A minder goed glucose kunnen verwerken dan mensen met de variatie G. De testuitslag G/G is de in West-Europa meest voorkomende testuitslag (67%).")
        first_word_bold("GSTM1 is betrokken bij het afbreken van lichaamsvreemde stoffen (medicijnen) en het beschermt tegen oxidatieve stress. GSTM1 is van invloed op het neutraliseren van vrije radicalen, waardoor cellen beter beschermd zijn tegen oxidatieve schade. De GSTM1-variatie kan voorkomen in een \"insertie\" (aanwezigheid van het gen) of \"deletie\" (afwezigheid van het gen). De afwezigheid van 1 null polymorfisme betekent al dat het functioneel GSTM1 enzym afwezig is. Mensen met de GSTM1-deletie zijn gevoeliger voor blootstelling aan lichaamsvreemde stoffen, waaronder medicijnen, en lopen daardoor meer risico op ziekten. GSTM1-deletie kan oorzaak zijn van bijwerkingen bij gebruik van sommige medicijnen. Vaststellen van GSTM1-deletie is van belang bij gebruik van cyclofosfamide, paracetamol, isonazide en fluoxetine en andere SSRI’s. Mensen met de GSTM1-insertie zijn meestal beter bestand tegen blootstelling van lichaamsvreemde stoffen.")
        first_word_bold("GSTP1 is van invloed op het afbreken van lichaamsvreemde stoffen, waaronder medicijnen. Daarnaast beschermt GSTP1 tegen oxidatieve stress. De variatie GSTP1*A (rs1138272; C is wildtype en T is de variatie) wordt geassocieerd met een verhoogde activiteit met als gevolg een snellere afbraak en verwijdering van lichaamsvreemde stoffen. De variatie GSTP1*C (rs743572; C is wildtype en T is de variatie) wordt geassocieerd met een verminderde activiteit om lichaamsvreemde stoffen af te breken en te verwijderen. De variatie GSTP1*B (rs1695; A is wildtype en G is de variatie) wordt geassocieerd met een gemiddelde normale activiteit. Het kennen van de activiteit van GSTP1 is van belang bij het gebruik van doxorubicine en cyclofosfamide; verminderde afbraak kan leiden tot bijwerkingen of een verminderd resultaat. Bij *A en *C geeft gebruik van NSAID’s verhoogd risico op bijwerkingen. Het gebruik van supplementen, die als antioxidant werken (vitamine C of E) kunnen van invloed zijn op de activiteit van GSTP1.")
        first_word_bold("GSTT1 is betrokken bij het ontgiften van lichaamsvreemde stoffen (waaronder medicijnen) en beschermt tegen oxidatieve stress. Net als bij GSTM1, kan ook bij GSTT1 sprake zijn van af- of aanwezigheid van het functioneel GSTT1-enzym. GSTT-deletie (afwezigheid) kan leiden tot een verminderd vermogen medicijnen af te breken. De verhoogde bloedspiegels geven dan kans op bijwerkingen. Vaststellen van GSTT1-deletie is van belang bij gebruik van Cyclofosfamide, Paracetamol, Isonazide en tricyclische antidepressiva (zoals Amitriptyline). Mensen met de GSTT1-insertie zijn meestal beter bestand tegen blootstelling van lichaamsvreemde stoffen.")
        first_word_bold("HLA-B*3101 is belangrijk voor het immuunsysteem. HLA-B stelt T-cellen in staat om geïnfecteerde of abnormale cellen te herkennen en te elimineren. De aanwezigheid van HLA-B*3101 is vooral van belang bij gebruik van carbamazepine. De aanwezigheid van HLA-B3101 wordt geassocieerd met een verhoogd risico op het Steven-Johnson-syndroom (SJS). De geteste variatie rs1061235; A is wildtype en T is de variatie) komt onder circa 5% voor bij de West-Europese bevolking. De geteste variatie rs1633021 (T is wildtype en C is de variatie) komt vooral bij Aziaten voor.")
        first_word_bold("IFNL3/IL28B is van invloed op de immuunrespons tegen virale infecties, vooral hepatitis C. IFNL3/IL28B stimuleert de immuunrespons tegen virale infecties en kan vermenigvuldigen van virussen remmen. De activiteit van IFNL3/IL28B is belangrijk voor sommige therapieën met bepaalde medicijncombinaties, zoals telaprevir, peginterferon alfa2a en ribavirin. In dit onderzoek is de variatie rs12979860 getest. De testuitslag C geeft een betere respons op interferon-behandeling dan de testuitslag T. ")
        first_word_bold("IGF-1 is een belangrijke mediator van de groei en vermenigvuldiging van cellen in verschillende delen van ons lichaam, zoals botten, spieren, weefsels en organen. Tijdens de kindertijd en adolescentie is IGF-1 essentieel voor een normale groei en ontwikkeling. IGF-1 heeft ook invloed op ons zenuwstelsel, met name in de hersenen en kan daarom van invloed zijn op cognitieve functies. In dit onderzoek is de variatie rs35767 getest. A geeft de aanwezigheid van de variatie aan. De variatie rs35767 is geassocieerd met een verminderde expressie van IGF-1. Dit betekent dat mensen met de testuitslag A een lagere hoeveelheid IGF-1 in hun lichaam hebben dan mensen met de testuitslag G. De testuitslag G/G is in West-Europa de meest voorkomende testuitslag.")
        first_word_bold("LDLR is belangrijk voor een goede balans van ons cholesterolgehalte. De variaties rs6511720 en rs57217136 in LDLR zijn functionele variaties. Beide allelen kunnen bijdragen aan een verhoogde LDLR-expressie, wat dus geassocieerd wordt met verlaagde LDL-C-niveaus. LDLR- expressie is verhoogd bij degenen die het T-allel dragen. In dit onderzoek is de variatie rs6511720 getest. T geeft de aanwezigheid van de variatie aan.")
        first_word_bold("LOC105447645;FUT 2 speelt een belangrijke rol door bepaalde suikermoleculen, voornamelijk fucose, te produceren op het oppervlak van cellen en slijmvliezen. Mensen met een functioneel FUT2-gen worden \"secretors\" genoemd (FUT2 positief), omdat ze deze bloedgroepantigenen kunnen afscheiden, zoals speeksel, tranen, zweet en slijm. De suikermoleculen, die door FUT2 worden gegenereerd, spelen ook een rol bij de immuunrespons en ontstekingsprocessen in de darmen en andere slijmvliezen. In dit onderzoek is de variatie rs492602 getest. A geeft de aanwezigheid van de variatie aan. De variant rs492602 is geassocieerd met een verminderde activiteit van fucosidase-2. Dit betekent dat mensen met de testuitslag A minder goed lactose kunnen afbreken dan mensen met de testuitslag G (lactose-intolerantie). De testuitslag A/G is de meest voorkomende testuitslag onder de West-Europese bevolking.")
        first_word_bold("MAO-B is betrokken is bij de afbraak van neurotransmitters in de hersenen, met name dopamine. De belangrijkste functie van MAO-B is het afbreken van overtollige dopamine. MAO-B is ook betrokken bij de afbraak van andere monoamine neurotransmitters, zoals fenylethylamine, tyramine, serotonine en noradrenaline. Sommige medicijnen (selegiline en rasagiline) remmen de activiteit van MAO-B. Dit betekent dat deze medicijnen MAO-B vertragen bij het afbreken van dopamine. In dit onderzoek is de variatie rs1799836 getest. C geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde activiteit van MAO-B. Dit betekent dat mensen met de testuitslag C minder snel dopamine, serotonine en noradrenaline afbreken dan mensen met de testuitslag T. N.B. Aangezien MAO-B op het X-chromosoom ligt, en mannen er hier maar één van hebben, wordt de testuitslag in een enkele letter uitgedrukt.  ")
        first_word_bold("MC4R speelt een grote rol bij ons eetgedrag. Wanneer melanocortine zich bindt aan de MC4R-receptor, wordt een signaal naar de hersenen gestuurd, dat het gevoel van verzadiging verhoogt en de eetlust vermindert. Daarnaast heeft MC4R invloed op hoe ons lichaam calorieën verbrandt en vet opslaat. Als het MC4R-gen niet goed functioneert, kan dit leiden tot een verminderde gevoeligheid voor verzadigingssignalen. Dit betekent dat mensen met deze variatie mogelijk meer eten met als gevolg gewichtstoename. Deze erfelijke variatie komt bij circa 1-4% van de bevolking voor. In dit onderzoek is de variatie rs17782313 getest. C geeft de aanwezigheid van de variatie aan. De testuitslag T/T is de meest voorkomende testuitslag (51%) onder de West-Europese bevolking.")
        first_word_bold("MnSOD speelt een belangrijke rol bij het beschermen van cellen tegen oxidatieve stress en is voornamelijk actief is in de mitochondriën van cellen. Oxidatieve stress kan schade aan celmembramen, enzymen en DNA veroorzaken. MnSOD zet het schadelijke superoxide-radicaal (O₂⁻), een bijproduct van de energieproductie in de mitochondriën, om in waterstofperoxide (H₂O₂). Vervolgens wordt het waterstofperoxide door andere enzymen verder afgebroken tot water en zuurstof. In dit onderzoek is de variatie rs4880 getest. De testuitslag Val/Val (A/A) is het wildtype (meest voorkomend). Mensen met de testuitslag Ala/Ala (G/G, verminderde activiteit van MnSOD)")
        first_word_bold("MTHFR1298 is betrokken bij het foliumzuurmetabolisme, maar deze variatie heeft minder effect op het metabolisme van folaat dan MTHFR677. MTHFR1298-IM of -PM als testuitslag kan de omzetting wel iets vertragen, vooral als iemand ook MTHFR677 IM of PM als testuitslag heeft. Overwogen kan worden om geactiveerd folaat te gebruiken in plaats van gewoon foliumzuur. Dit kan helpen om een tekort aan folaat te voorkomen, wat vooral belangrijk is voor zwangere vrouwen. Magnesium, zink en B vitamines ondersteunen de activiteit van MTHFR.")
        first_word_bold("MTRNR1 reguleert de gevoeligheid voor insuline. MTRNR1 is ook belangrijk voor evenwichtige regulering van andere variabelen, als vochtbalans, concentraties van natrium, kalium en calcium en de bloedsuikerspiegel")
        first_word_bold("MTNR1B speelt een rol in twee belangrijke dingen: onze biologische klok en hoe ons lichaam suiker (glucose) verwerkt. MTNR1B heeft invloed op onze gevoeligheid voor melatonine en dat beïnvloedt ons slaap-waakritme. In dit onderzoek is de variatie rs10830963 getest. G geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde activiteit van MTNR1B voor melatonine. Dit betekent dat mensen met de testuitslag G minder gevoelig zijn voor melatonine dan mensen met de testuitslag C. De testuitslag C/C komt bij 45% van de West-Europese bevolking voor. De testuitslag C/G komt voor bij 51% van de West-Europese bevolking. Dit betekent dat veel mensen deze genetische variatie hebben. ")
        first_word_bold("NADSYN1 is betrokken bij tal van biochemische reacties in de cel, zoals de energieproductie in de mitochondriën. NAD speelt een belangrijke rol bij de productie van adenosinetrifosfaat (ATP). NAD fungeert ook als een soort helper (cofactor) voor verschillende enzymen, die betrokken zijn bij glycolyse (de afbraak van glucose), de citroenzuurcyclus en oxidatieve stressrespons. Variaties in NADSYN1 kunnen leiden tot een verminderde productie van NAD. NADSYN1 heeft ook invloed op het vitamine D niveau. In dit onderzoek is de variatie rs12785878 getest. T geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde productie van NAD+. Dit betekent dat mensen met de testuitslag T een lagere hoeveelheid NAD+ in hun lichaam hebben dan mensen met de testuitslag G.")
        first_word_bold("NAT1 is belangrijk voor de koppelingsreactie tussen lichaamseigen en lichaamsvreemde stoffen. NA1 zet verscheidene geneesmiddelen om door aan een aminegroep een acetylgroep te binden. Bij de N-Acetyl Transferase NAT-1 onderscheiden we op gen niveau snelle (rapid: R) en trage (Slow: S) allelen.")
        first_word_bold("NAT2 kent op gen niveau snelle (Rapids: R) en trage (Slow: S) allelen. Snelle allelen zijn: *4, *11, *12 en *13. Trage allelen zijn: *5, *6, *7 en *14. Het percentage trage acetyleerders onder de westerse bevolking is circa 60% en onder de Afrikaanse en Aziatische bevolking circa 45%. Aanwezigheid van een of twee trage allelen geeft een verhoogde kans op levertoxiciteit bij gebruik van o.a. Isoniazide en Hydralazine")
        first_word_bold("NBPF3 speelt een rol bij de ontwikkeling van ons zenuwstelsel. NBPF3 wordt ook in verband gebracht met het afvoeren van vitamine B6 uit ons lichaam. Vitamine B6 is noodzakelijk voor een goede werking van ons zenuwstelsel, helpt bij de productie van rode bloedcellen en speelt een rol bij de manier waarop ons lichaam suikers verwerkt. In dit onderzoek is de variatie rs4654748 getest. T geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde expressie van NBPF3. Dit betekent dat mensen met de testuitslag T een lagere hoeveelheid NBPF3 in hun hersenen hebben dan mensen met de testuitslag C. De testuitslag C/C betekent dat je waarschijnlijk een normaal niveau vitamine B6 hebt.")
        first_word_bold("NQO1 is van belang voor het ontgiften van schadelijke stoffen. NQ01 beschermt de cellen tegen oxidatieve stress door het reduceren van quinonen. Quinonen zijn geoxideerde verbindingen afgeleid van aromatische structuren. Quinonen spelen een belangrijke rol in biologische processen, zoals elektronentransport. Hoewel ze nuttig zijn voor energieproductie in cellen, kunnen ze ook toxisch zijn en bijdragen aan oxidatieve stress en cel schade. Sommige medicijnen bevatten quinonen of worden omgezet in quinonverbindingen in het lichaam (bijvoorbeeld mitomycine C). NQO1 is ook van invloed op de stabiliteit van p53. P53 is betrokken bij de bescherming van cellen tegen DNA-schade en apoptose (geprogrammeerde celdood). In dit onderzoek is de variatie rs1800566 getest. De testuitslag G/G (wildtype) is het meest voorkomend. De testuitslag A/G geeft een verminderde activiteit van NQ01 aan. De testuitslag A/A (3% onder de West-Europese bevolking) betekent een inactief NQ01 en dus een sterk verminderde bescherming tegen oxidatieve stress. Onder Aziaten komt de A-variant vaker voor dan onder Europese of Afrikaanse bevolkingsgroepen. Mensen met de testuitslag A lopen ook meer risico op schade door omgevingsfactoren als roken, luchtverontreiniging en toxische stoffen.")
        first_word_bold("OPRM1 is betrokken bij medicijnen voor pijnbestrijding. Variaties in dit gen hebben effect op effectiviteit van deze medicijnen.")
        first_word_bold("PON1 beschermt de cellen tegen oxidatieve stress en is betrokken bij het afbreken van lichaamsvreemde stoffen, zoals medicijnen. PON1 heeft een positief effect op het HDL-cholesterol en beschermt het LDL-cholesterol tegen oxidatie. Geoxideerd LDL-cholesterol is schadelijk, omdat het ontstekingen en plaquevorming kan bevorderen. In dit onderzoek is de variatie rs662 getest. De testuitslag C wordt geassocieerd met een verhoogde activiteit en dus worden lichaamsvreemde stoffen, waaronder medicijnen, sneller afgebroken. De testuitslag T bij rs662 wordt geassocieerd met een verminderde activiteit en dus worden lichaamsvreemde stoffen, waaronder medicijnen, niet of vertraagd afgebroken. Dit vergroot de kans op bijwerkingen. De activiteit van PON1 is van belang bij gebruik van prasugrel.")
        first_word_bold("RYR1 wordt voornamelijk in skeletspieren aangetroffen en is van invloed op de spierontwikkeling. Variaties in het RYR1-gen zijn geassocieerd met vatbaarheid voor maligne hyperthermie.")
        first_word_bold("SULT1A1 helpt om schadelijke stoffen, waaronder medicijnen, te neutraliseren (sulfatering). Sulfatering verhoogt de oplosbaarheid van lichaamsvreemde stoffen in water. Dit bevordert de verwijdering van die stoffen. SULT1A1 zet medicijnen om naar een inactieve vorm en regelt de verwijdering uit ons lichaam. Omdat SULT1A1 betrokken is bij het metabolisme van hormonen (estradiol), kan een verminderde activiteit de hormoonbalans beïnvloeden. In dit onderzoek is de variatie rs9282861 getest. G geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde activiteit van SULT1A1. Dit betekent dat mensen met de testuitslag G een lagere snelheid van afbraak van medicijnen en voedingsstoffen hebben dan mensen met de testuitslag A. Een verminderde activiteit van SULT1A1 is onder andere van belang bij gebruik van Paracetamol, tricyclische antidepressiva (zoals Amitriptyline en Nortriptyline), Tamoxifen en Estradiol.")
        first_word_bold("SULT1E1 zorgt ervoor, dat bepaalde hormonen genaamd oestrogenen minder actief worden. In dit onderzoek is de variatie rs3736599 getest. T geeft de aanwezigheid van de variatie aan. Deze variatie wordt geassocieerd met een verhoogde oestrogeenactiviteit. Dit komt omdat de variatie leidt tot een verminderde activiteit en dus minder oestrogenen kan inactiveren. Daarnaast worden de testuitslagen T/C en T/T in verband gebracht met een hoger risico op bijwerkingen en slechtere werking van antibiotica, antidepressiva en medicijnen tegen hoge bloeddruk.")
        first_word_bold("TCF7L2 is betrokken bij de regulering van de bloedsuikerspiegel en hoe gevoelig ons lichaam is voor insuline. Variaties in TCF7L2 kunnen leiden tot minder gevoeligheid voor insuline. Bovendien heeft TCF7L2 invloed op hoe ons darmweefsel functioneert en hoe goed het slijmvlies in onze darmen zijn werk doet als barrière. TCF7L2 speelt ook een rol bij de ontwikkeling en werking van zenuwcellen in ons centrale zenuwstelsel. Onderzoek heeft aangetoond dat veranderingen in TCF7L2 verband kunnen houden met bepaalde neurologische aandoeningen, zoals schizofrenie en bipolaire stoornis. In dit onderzoek is de variatie rs7903146 getest. T geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde activiteit van TCF7L2. Dit betekent dat mensen met de testuitslag T een lagere gevoeligheid voor insuline hebben dan mensen met de testuitslag C.")
        first_word_bold("TMEM165;CLOCK heeft vooral te maken met hoe cellen calcium en mangaan verwerken. In dit onderzoek is de variatie rs1801260 getest. G geeft de aanwezigheid van de variatie aan. De testtestuitslagen G/A en G/G geven een hoger risico op het ontwikkelen van ADHD.")
        first_word_bold("TNF-α is van invloed op de reactie van het immuunsysteem op ontstekingen. In dit onderzoek is de variatie rs1800629 getest. A geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verhoogde productie van TNF-α. Dit betekent dat mensen met de testuitslag A een hoger niveau van TNF-α in hun bloed hebben dan mensen met de testuitslag G. De testuitslag G/G is de meest voorkomende testuitslag (68%) onder de West-Europese bevolking. Verhoogde niveaus van TNF-α (A/G en A/A) worden in verband gebracht met de kans om reumatoïde artritis, de ziekte van Crohn en psoriasis te ontwikkelen. Infliximab en adalimumab remmen de werking van TNF-α.")
        first_word_bold("UCP2 bevindt zich in de mitochondriën, de energiecentrales van onze cellen en speelt een rol in de productie van ATP. ATP is een belangrijke energiebron in onze cellen. UCP2 heeft invloed op hoe onze cellen vetten en glucose verwerken, wat belangrijk is voor het handhaven van een evenwicht in het insulineniveau en de bloedsuikerspiegel. In dit onderzoek is specifiek gekeken naar een genetische variatie genaamd rs660339. A geeft de aanwezigheid van de variatie aan. Deze variatie is geassocieerd met een verminderde expressie van UCP2. Dit betekent dat mensen met de testuitslag A/G en A/A een lagere hoeveelheid UCP2 in hun cellen hebben in vergelijking met mensen met de uitslag G/G.")
        first_word_bold("VDR is van invloed op het vitamine D niveau. In dit onderzoek is de aanwezigheid van de variatie rs731236 getest. Mensen met een uitslag A/G en A/A een lagere respons op vitamine D hebben dan mensen met de uitslag G/G. Onder de West-Europese bevolking heeft ongeveer 35% van de mensen een testuitslag A/A en 43% van de mensen een testuitslag A/G. In dit onderzoek is ook de variatie rs1544410 getest. Deze variatie is van invloed op de gevoeligheid voor vitamine D en daarmee op het effect van vitamine D in het lichaam. De testuitslag C is het wildtype. De testuitslag T geeft de aanwezigheid van de variatie aan. Onder de West-Europese bevolking komt deze variatie relatief vaak voor (30-50%). Daarnaast zijn ook de variaties rs7975232 en rs2228570 getest. De testuitslag C bij rs7975232 is het wildtype. De testuitslag A geeft de variatie aan en wordt in verband gebracht met gevoeligheid voor botbreuken. De testuitslag G bij rs2228570 is het wildtype. De testuitslag A bij rs2228570 geeft de aanwezigheid van de variatie aan met als gevolg een verminderde activiteit van VDR. Verminderde activiteit van VDR is geassocieerd met een minder actief immuunsysteem.")
        self.document.add_page_break()

    def variaties_waarop_is_getest(self):
        getesteVariatiesDict = {
            "ABCB1": "rs1045642",
            "ACE": "rs4341",
            "ADIPOQ": "rs17300539",
            "ADRA2A": "rs10885122",
            "ALDH2": "rs671",
            "AMDHD1": "rs10745742",
            "BChE": "rs1803274, rs1799807, rs28933389, rs28933390",
            "BCO1": "rs12934922",
            "BDNF-AS; BDNF": "rs6265",
            "CACNA1S": "WT, c.520C>T, c.3257G>A",
            "CFTR": "WT, F508delCTT en G551D",
            "CYP1A1": "rs1048943",
            "CYP1B1": "rs1056836",
            "CYP2A6": "*1, *9",
            "CYP2C8": "*1(A, B, C), *2, *3, *4+1C, *5, *7, *8, *10 t/m *14, P404A",
            "CYP2E1": "*1A, *1B, *3, *4, *5B, *5+1B, *7A, *7A+1B, *7B, *7C, *7C+1B, *4+7A, 4+5+7A, *4+7A+1B, *5A+7A+1B",
            "CYP2F1": "*1, *2, *3, *4, *6",
            "CYP2R1": "rs10741657, rs12794714",
            "CYP4F2": "*1, *2, *3, *2+3",
            "CYP17A": "rs743572",
            "CYP24A1": "rs17216707",
            "DHCR7": "rs12785878",
            "DRD2": "rs1076560",
            "F2": "rs1799963",
            "F5": "WT, rs6025",
            "FTO": "rs1121980",
            "GC": "rs7041",
            "GSTM1": "CNV",
            "GSTP1": "rs1138272, rs1695, rs743572",
            "GSTT1": "CNV",
            "HLA-B*3101": "rs1061235, rs1633021",
            "IFNL3/IL28B": "rs12979860",
            "GCK, YKT6": "rs4607517",
            "IGF1": "rs35767",
            "LDLR": "rs6511720",
            "LOC105447645; FUT2": "rs492602",
            "MAO-B": "rs1799836",
            "MC4R": "rs17782313",
            "MnSOD": "rs4880",
            "MTHFR 1298A>C": "rs1801131",
            "MTRNR1": "m.=, m.827A>G, m.1095T>C, m.1494C>T, Null",
            "MTNR1B": "rs10830963",
            "NADSYN1": "rs12785878",
            "NAT1": "*4, *5, *11B, *11C, *14, *15, *17 t/m *19(A, B), *20 t/m *25, *27, *29, *30",
            "NAT2": "*4,*5,(E, I, L, M, N, O, P, S, X, Y),*6,(G, H, I, J, K, M, O, P, V) *7, *7D, *10, *12(D, E, F, G, H, J, K, O),*13(D, F,) *14(D, F, H, K), *17 t/m *19, *21, *23, *24",
            "NBPF3": "rs4654748",
            "NQ01": "rs1800566",
            "OPRM1": "rs1799971",
            "PON1": "rs662",
            "RYR1": "WT,p.1571V+3933C,c.103T>C,c.418G>A,c.487C>T,c.742G>A,c.742G>C,c.1021G>A,c.1021G>C,c.1565A>C,c.1565A>G,c.1840C>T,c.1841G>T,c.6487C>T,c.6488G>A,c.6488G>C,c.6502G>A,c.6617C>G,c.6617C>T,c.7007G>A,c.7025A>G,c.7063C>T,c.7304G>A,c.7354C>T,c.7360C>T,c.7361G>A,c.7372C>T,c.7373G>A,c.11708G>A,c.14387A>G,c.14477C>T,c.14497C>T,c.14545G>A,c.14582G>A",
            "Sult1A1": "rs9282861",
            "Sult1E1": "rs3736599",
            "TCF7L2": "rs7903146",
            "TMEM165; CLOCK": "rs1801260",
            "TNFa": "rs1800629",
            "UCP2": "rs660339",
            "VDR": "rs731236, rs1544410, rs7975232, rs2228570"
        }

        self.heading("Variaties waarop is getest", lined=True)
        testvariaties_table = self.document.add_table(rows=len(getesteVariatiesDict) + 1, cols=2,
                                                      style="Table Grid")
        self.styled_cell_text(testvariaties_table.cell(0, 0), "Gen", make_bold=True, set_linespacing=True)
        self.styled_cell_text(testvariaties_table.cell(0, 1), "Variaties", make_bold=True, set_linespacing=True)
        for i in range(2):
            if i == 0:
                for j, k in zip(range(len(getesteVariatiesDict)), getesteVariatiesDict.keys()):
                    self.styled_cell_text(testvariaties_table.cell(j + 1, i), text=k,
                                          set_linespacing=True)
            else:
                for j, k in zip(range(len(getesteVariatiesDict)), getesteVariatiesDict.values()):
                    self.styled_cell_text(testvariaties_table.cell(j + 1, i), text=k,
                                          set_linespacing=True)
        testvariaties_table.allow_autofit = False
        for cell in testvariaties_table.columns[0].cells:
            cell.width = Cm(3.19)
        for cell in testvariaties_table.columns[1].cells:
            cell.width = Cm(14.78)

    def save(self):
        """
        Saves the document with a name based on the customer id and the type of document.
        :return:
        """
        document_name = 'NutrinomicsReport' + "_" + self.sample_id + ".docx"
        self.document.save(f"Output\\Reports\\{document_name}")