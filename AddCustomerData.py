# This is where the reports are modified to include customer details

import PySimpleGUI as sg
from ELT import Extract
from docx import Document
from os.path import join
from docx.shared import Pt
import Utilities as util

class CustomerData:
    def customer_data_IA(self):
        customerdata_df = Extract().customer_data()
        customerdata_df = customerdata_df.rename(columns={0:'sample_id',1:'initials',2:'lastname',3:'birthdate',4:'status'})
        customerdata_df = customerdata_df.fillna('')
        customerdata_df = customerdata_df.apply(lambda col: col.map(lambda x: x.strip() if isinstance(x, str) else x))
        customerdata_df['birthdate'] = customerdata_df['birthdate'].dt.strftime('%Y-%m-%d')
        customerdata_df['birthdate']= customerdata_df['birthdate'].fillna('20237-01-01')
        customerdata_df.sort_values(by='sample_id', ascending=True, inplace=True)
        customerdata_df.reset_index(inplace=True, drop=True)
        return customerdata_df

    def fill_customer_data(self):
        customerdata_df = self.customer_data_IA()
        path = 'Output\\Reports'
        reports = util.get_reports()

        reported_file_ids = []
        for file in reports:
            filepath = join(path,file)
            doc = Document(filepath)
            report_type = file.split('_')[0]
            sample_id = file.split('_')[1].split('.')[0]
            reported_file_ids.append(sample_id)

            if sample_id in list(customerdata_df['sample_id']):
                this_customers_data = customerdata_df[customerdata_df['sample_id'] == sample_id].values[0]
                customer_data_dict = {
                    'name': this_customers_data[1] + ' ' + this_customers_data[2],
                    'birthdate': this_customers_data[3],
                }
                if customer_data_dict['birthdate'] == '20237-01-01':
                    customer_data_dict['birthdate'] = ' '

                match report_type:
                    case 'FarmacogeneticReport':
                        table = doc.tables[0]
                        cells_to_edit = [1,3]
                        for cell_index, new_text in zip(cells_to_edit, list(customer_data_dict.keys())):
                            cell = table.rows[0].cells[cell_index]
                            cell.text = ''
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(customer_data_dict[new_text])
                            run.font.name = 'Calibri'
                            run.font.size = Pt(12)
                        doc.save(filepath)

                    case 'InfoSheet':
                        paragraph = doc.paragraphs[1]
                        paragraph.clear()
                        run = paragraph.add_run('Naam :  {}\nGeboortedatum :  {}'.format(customer_data_dict['name'],
                                                                                       customer_data_dict['birthdate']))
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)
                        run.bold = True
                        run.underline = True

                    case _:
                        pass

            else:
                print(f'{file} heeft geen geassocieerde klantdata')

            doc.save(filepath)
            doc.save(filepath)

        for sample in customerdata_df['sample_id']:
            if sample not in reported_file_ids:
                    print(f'{sample} heeft wel klantdata, maar geen geassocieerd rapport')

    def execute(self):
        file_is_present = sg.popup_yes_no("Heeft u het bestand met de klantdata?")
        if file_is_present == 'Yes':
            print('Filling in customer data [...]')
            self.fill_customer_data()
            print('Filling in customer data [DONE]')
        else:
            pass