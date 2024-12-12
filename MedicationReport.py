# This is where the medication match is generated

#imports
from WordDocument import WordEditing as wd
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm

class MedicationReport(wd):
    def __init__(self, sample_id, dataframe):
        super().__init__(sample_id, dataframe)

    def medicationmatch_IA(self):
        medRep_genes = ["COMT",
                             "CYP1A2",
                             "CYP2A6",
                             "CYP2B6",
                             "CYP2C8",
                             "CYP2C9",
                             "CYP2C19",
                             "CYP2D6",
                             "CYP3A4",
                             "CYP3A5",
                             "CYP4F2",
                             "DPYD",
                             "IFNL3",
                             "SLCO1B1",
                             "TPMT",
                             "UGT1A1",
                             "VKORC1",
                             "F2",
                             "F5",
                             "MTHFRA1298C",
                             "MTHFRC677T",
                             "HLA*1502",
                             "OPRM1"]

        df = self.dataframe[self.dataframe['sample_id'] == self.sample_id]
        df = df[df['gene'].isin(medRep_genes)]
        df.drop(['sample_id', 'genotype'], axis='columns', inplace=True)
        df.rename(columns={'gene': 'Gen', 'phenotype': 'Fenotype/functie'}, inplace=True)

        return df
    
    def medrep_intro_text(self):
        self.heading("BIJLAGE: medicatiematch " + self.sample_id, chosen_size=14, lined=True)
        med_table = self.document.add_table(rows=4, cols=2, style="Table Grid")
        text = "De interactie tussen genotype en medicijn geeft een sterk verhoogd risico op non-effectiviteit en bijwerkingen; contra-geïnduceerd; alternatieven overwegen.\n"
        self.styled_cell_text(med_table.cell(0, 1), text, chosen_size=11)
        text = "Verminderde interactie tussen genotype en medicijn geeft risico op minder werking van het medicijn en bijwerkingen; alternatieven overwegen.\n"
        self.styled_cell_text(med_table.cell(1, 1), text, chosen_size=11)
        text = "Er is interactie tussen genotype en medicijn, maar er zijn onvoldoende adviezen om de effecten op de werking van het medicijn of het optreden van bijwerkingen direct te herleiden tot het genotype.\n"
        self.styled_cell_text(med_table.cell(2, 1), text, chosen_size=11)
        text = "De volgorde waarin geneesmiddelen worden vermeld, heeft geen klinische of medische implicatie.\n"
        self.styled_cell_text(med_table.cell(3, 1), text, chosen_size=11)
        pictures = ["Input/Icons/sterk_risico.png",
                    "Input/Icons/risico.png",
                    "Input/Icons/info.png",
                    "Input/Icons/nb.png"]
        counter = 0
        for picture in pictures:
            med_table.cell(counter, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = med_table.cell(counter, 0).paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(picture, width=wd.centimeters(0.8), height=wd.centimeters(0.7))
            counter += 1
        for cell in med_table.columns[0].cells:
            cell.width = wd.centimeters(0.7)
        for cell in med_table.columns[1].cells:
            cell.width = wd.centimeters(21)
        self.document.add_paragraph("")
    
    def medrep_core(self, gen, phenotype):
        """
        :param gen: The gene of which a table is to be made
        :param phenotype: The phenotype of the gene
        :return: Adds a medication match table to the medication match document
        """

        genedict = {
            "symbolList_CYP2B6_PM": [1],
            "symbolList_IFNL3_PM": [1, 1, 1],
            "symbolList_OPRM1_IM": [2],
            "symbolList_UGT1A1_IM": [1, 2, 1],
            "symbolList_VKORC1_IM": [1, 1, 1],
            "symbolList_COMT_PM": [1, 1, 1, 1, 1, 2, 1],
            "symbolList_DPYD_PM": [0, 0, 2],
            "symbolList_SLCO1B1_PM": [0],
            "symbolList_TPMT_PM": [1, 1, 1],
            "symbolList_UGT1A1_PM": [1, 2, 1],
            "symbolList_VKORC1_PM": [1, 1, 1],
            "symbolList_CYP2C9_IM": [1, 1, 0, 3, 1, 1, 2, 2, 2, 1, 2, 1],
            "symbolList_CYP2C9_PM": [1, 0, 0, 1],
            "symbolList_CYP1A2_IM": [2, 2, 2],
            "symbolList_CYP2C8_?": [2, 2],
            "symbolList_HLA-B*1502_aanwezig": [1],
            "symbolList_CYP3A4_IM": [1, 1, 1, 1, 1, 1, 1, 1, 2, 0, 1, 2, 1, 1, 2, 1, 1],
            "symbolList_CYP3A4_PM": [1, 1, 1, 1, 1, 1, 1, 1, 2, 0, 0, 2, 1, 1, 2, 1, 1],
            "symbolList_CYP3A5_heterozygoot": [1],
            "symbolList_CYP2A6_IM": [1, 1, 1, 1],
            "symbolList_CYP2B6_IM": [1, 1, 1, 2, 2, 1, 2],
            "symbolList_DPYD_IM": [0, 0, 2],
            "symbolList_IFNL3_IM": [1, 1, 1],
            "symbolList_SLCO1B1_DF": [1, 1, 2],
            "symbolList_TPMT_IM": [1, 1, 1],
            "symbolList_CYP2C19_IM": [2, 2, 2, 1, 2, 1, 2, 0, 2, 2, 2, 2, 2, 2, 1, 2, 2, 1, 2, 2, 1, 2, 1, 1, 1],
            "symbolList_CYP2C19_PM": [0, 1, 0, 2, 1],
            "symbolList_CYP2C19_RM": [1, 1, 1, 1, 0, 1, 1, 1, 1, 1, 1, 1, 2, 1, 1, 1, 1, 1],
            "symbolList_CYP2C19_UM": [1, 1, 1, 1, 0, 1, 1, 1, 1, 1, 1, 1, 2, 1, 1, 1, 1, 1],
            "symbolList_CYP2D6_IM": [1, 2, 1, 0, 1, 1, 2, 1, 1, 1, 1, 1, 1, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 2, 1, 1, 1,
                                     1],
            "symbolList_CYP2D6_PM": [0, 1, 1, 1, 1, 1, 2, 0, 1, 1, 2, 1, 2, 1, 1, 1, 2, 2, 2, 1, 1, 1, 0, 1, 2, 1, 2, 1,
                                     1,
                                     2,
                                     1, 1, 1, 1, 1, 2, 1, 2, 1, 1, 1],
            "symbolList_CYP2D6_UM": [0, 1, 1, 1, 0, 1, 1, 1, 1, 1, 2, 1, 2, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 2, 1, 1, 1, 1,
                                     1,
                                     1,
                                     1],
            "symbolList_CYP2C8": [2, 2, 2],
            "symbolList_CYP2A6_PM": [1, 1, 1, 1],
            "symbolList_CYP2C8_IM": [2, 2],
            "symbolList_CYP2C8_PM": [2, 2],
            "symbolList_CYP4F2_IM": [1],
            "symbolList_CYP4F2_PM": [1],
            "symbolList_OPRM1_PM": [2],
            "medList_onduidelijk": ["Onduidelijk"],
            "effectList_onduidelijk": ["Onduidelijk"],
            "symbolList_onduidelijk": [2],
            "medList_unknown": ["Unknown"],
            "effectList_unknown": ["Het fenotype kon niet bepaald worden. Dit moet handmatig nagekeken worden"],
            "symbolList_unknown": [2],
            "medList_NotInDict": ["Not present in dictionary"],
            "effectList_NotInDict": ["Er is geen medicatiematch geleverd van deze gen/fenotype combinatie"],
            "symbolList_NotInDict": [2],
            "medList_geenAdvies": ["Geen advies"],
            "effectList_geenAdvies": [
                "Er zijn op dit moment in Nederland nog geen medicatie-adviezen op basis van dit fenotype. Door de snelle ontwikkeling in het veld van farmacogenetica kan dit in de toekomst veranderen."],
            "symbolList_geenAdvies": [2],
            "medList_F5_NM": ["Anticonceptie (oestrogeen)"],
            "effectList_F5_NM": [
                "Vragen naar het voorkomen van trombotische aandoeningen in de familie. Oestrogeen-bevattende AC vermijden en alternatief selecteren. (bijvoorbeeld progestageen-anticonceptiemiddel). Extra risicofactoren zijn obesitas en roken."],
            "medList_HLA-B*1502_PM": ["Carbamazepine"],
            "effectList_HLA-B*1502_PM": [
                "HLA-B*1502: PM;  Carbamazepine therapie is geassocieerd met cutane bijwerkingen in 10% van de patiënten. Aanwezigheid van de HLA-B*1502 allel wordt geassocieerd met een verhoogd risico op Carbamazepine (CBZ) geïnduceerde Steven-Johnsons Syndroom (SJS) of toxische epidermale necrolyse (TEN) in de Aziatische populatie (Han Chinezen, Maleisiers en Thai). HLA-B*1502 dragers hebben een Odds Ratio van 113 op het ontwikkelen van SJS/TEN (Yip et al 2012). HLA-B*1502 aanwezigheid van dit allel wordt in het Kaukasisch ras niet geassocieerd met SJS/TEN. Daarentegen wordt de aanwezigheid van de HLA-A*3101 allel geassocieerd met alle Carbamazapine overgevoeligheid, onafhankelijk van etniciteit (Odds ratio 9.5 Dosis advies beschikbaar in Kennisbank KNMP."],
            "medList_CYP3A4_IM": ['Alprazolam', 'Atorvastatine', 'Buprenorfine', 'Buspiron', 'Clonazepam',
                                  'Cyclosporine',
                                  'Eszopiclon', 'Fentanyl', 'Guanfacine', 'Ketoconazol', 'Quetiapine', 'Piritramide',
                                  'Saxagliptine', 'Sirolimus', 'Sufentanil', 'Ticagrelor', 'Trazodon'],
            "effectList_CYP3A4_IM": ['CYP3A4-IM; verhoogd risico op bijwerkingen. Een alternatief medicijn overwegen.',
                                     'CYP3A4-IM; personen met een intermediair metabolisme verwerken Aatorvastatine langzaam en zullen naar verwachting goed reageren op de behandeling. Rekening houden met de verhoogde plasmaspiegels van Atorvastatine, die het risico op myopathie / rabdomyolyse verhogen. Belangrijk is ook de activiteit van SLCO1B1, het transport gen voor statines. De activiteit van SLCO1B1 is bepalend voor het effect van Atorvastatine. Bij gebruik van statines letten op vitamine D-spiegel.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen. Onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen of de aanbevolen startdosis verlagen.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen. Een alternatief medicijn overwegen; Buspiron dient met voorzichtigheid te worden gebruikt, houd daarbij rekening met de juiste dosering.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen. Onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen of de aanbevolen startdosis verlagen.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen vanwege verminderde omzetting van Cyclosporine. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen. Een alternatief medicijn overwegen of de dosis verlagen. Daarnaast vermeldt het LUNESTA-etiket dat de totale dosis LUNESTA niet hoger mag zijn dan 2 mg bij patiënten met ernstige leverinsufficiëntie, of bij patiënten, die LUNESTA gelijktijdig toegediend krijgen met krachtige CYP3A4-remmers.',
                                     'CYP3A4-IM; Mogelijke bijwerkingen, waaronder ademhalingsdepressie en onvoldoende pijnbestrijding. Een alternatief medicijn overwegen als de patiënt ook een CYP3A4-remmer gebruikt.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen.',
                                     'CYP3A4-IM; sterk verhoogd risico op bijwerkingen vanwege verminderde omzetting van de Ketoconazol. Ketoconazol is gecontra-indiceerd. ',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen vanwege verminderde omzetting van Quetiapine. Als slaapmiddel in een lage dosering voorgeschreven kan Quetiapine leiden tot bijwerkingen als gewichtstoename, onregelmatig hartritme, verhoogde glucosewaarden en suïcidale gedachten. Er zijn echter onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief geneesmiddel overwegen of de aanbevolen startdosis verlagen. In geringe mate ook metabolisme via CYP2D6.',
                                     'CYP3A4-IM; verminderde activiteit vergroot kans op bijwerkingen.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen. Onvoldoende onderzoeksgegevens om de dosisaanpassing te kunnen berekenen. Een alternatief medicijn overwegen of de aanbevolen startdosis verlagen.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen vanwege verminderde omzetting van Sirolimus. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen.',
                                     'CYP3A4-IM; metabolisme voornamelijk via CYP3A4. Verminderde activiteit geeft bij gebruik van Sufentanil kans op bijwerkingen.',
                                     'CYP3A4-IM; een verhoogd risico op onvoldoende werking van Ticagrelor. Een alternatief medicijn overwegen.',
                                     'CYP3A4-IM; verhoogd risico op bijwerkingen vanwege verminderde omzetting van Trazodon. Een en alternatief medicijn overwegen of overwegen de aanbevolen startdosering te verlagen. In geringe mate ook metabolisme via CYP2D6.'],
            "medList_CYP3A4_PM": ['Alprazolam', 'Atorvastatine', 'Buprenorfine', 'Buspiron', 'Clonazepam',
                                  'Cyclosporine',
                                  'Eszopiclon', 'Fentanyl', 'Guanfacine', 'Ketoconazol', 'Quetiapine', 'Piritramide',
                                  'Saxagliptine', 'Sirolimus', 'Sufentanil', 'Ticagrelor', 'Trazodon'],
            "effectList_CYP3A4_PM": [
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen. Een alternatief medicijn overwegen.',
                'CYP3A4-PM; belangrijk is ook de activiteit van SLCO1B1, het transport gen voor statines. De activiteit van SLCO1B1 is bepalend voor het effect van Atorvastatine. Bij gebruik van statines letten op vitamine D-spiegel.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen. Onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen of de aanbevolen startdosis verlagen.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen. Een alternatief medicijn overwegen; Buspiron dient met voorzichtigheid te worden gebruikt, houd daarbij rekening met de juiste dosering.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen. Onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen of de aanbevolen startdosis verlagen.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen vanwege verminderde omzetting van Cyclosporine. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen. Een alternatief medicijn overwegen of de dosis verlagen. Daarnaast vermeldt het LUNESTA-etiket dat de totale dosis LUNESTA niet hoger mag zijn dan 2 mg bij patiënten met ernstige leverinsufficiëntie, of bij patiënten, die LUNESTA gelijktijdig toegediend krijgen met krachtige CYP3A4-remmers.',
                'CYP3A4-PM; grote kans op bijwerkingen, waaronder ademhalingsdepressie en onvoldoende pijnbestrijding. Een alternatief medicijn overwegen als de patiënt ook een CYP3A4-remmer gebruikt.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen vanwege verminderde omzetting van de Ketoconazol. Ketoconazol is gecontra-indiceerd. ',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen vanwege nauwelijks of geen omzetting van Quetiapine. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief geneesmiddel overwegen wordt sterk geadviseerd. In geringe mate ook metabolisme via CYP2D6.',
                'CYP3A4-PM; vanwege geen of nauwelijks omzetting sterke kans  op bijwerkingen.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen. Onvoldoende onderzoeksgegevens om de dosisaanpassing te kunnen berekenen. Een alternatief medicijn overwegen.',
                'CYP3A4-IM; sterk verhoogd risico op bijwerkingen vanwege nauwelijks of geen omzetting van Sirolimus. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen.',
                'CYP3A4-PM; met dit fenotype geen of nauwelijks werking van Sufentanil. Alternatief geneesmiddel overwegen.',
                'CYP3A4-PM; een sterk verhoogd risico op onvoldoende werking van Ticagrelor. Een alternatief medicijn overwegen.',
                'CYP3A4-PM; sterk verhoogd risico op bijwerkingen vanwege geen of nauwelijks omzetting van Trazodon. Een en alternatief medicijn overwegen. In geringe mate ook metabolisering via CYP2D6.'],
            "medList_CYP2C9_IM": ['Acenocoumarol', 'Celecoxib', 'Diclofenac', 'Fenproc0umon', 'Fenytoïne',
                                  'Fosfenytoine',
                                  'Glibenclamide', 'Gliclazide', 'Glimepiride', 'Siponimod', 'Tolbutamide',
                                  'Warfarine'],
            "effectList_CYP2C9_IM": [
                'CYP2C9-IM; het fenotype kan leiden tot verlaging van de benodigde onderhoudsdosis. Er is echter onvoldoende onderbouwing dat dit problemen oplevert bij normale instelling van de therapie. Het wordt aanbevolen om INR (mate waarin het bloed stolt) vaker te controleren. Geen vitamine K2 gebruiken bij Acenocoumarol. CYP1A2 is ook betrokken bij  de metabolisme van Acenocoumarol. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C9-IM; meer kans op bijwerkingen van Celecoxib. Ook metabolisering via CYP3A4.',
                'CYP2C9-IM; personen met dit genotype, die worden behandeld met dit niet-steroïde ontstekingsremmende middel, kunnen een verhoogd risico op gastro-intestinale bloeding hebben. Een alternatief medicijn overwegen.',
                'CYP2C9-IM; het fenotype kan leiden tot verlaging van de benodigde onderhoudsdosis. Er is echter onvoldoende onderbouwing dat dit problemen oplevert bij normale instelling van de therapie (d.w.z. frequente INR-controle)',
                'CYP2C9-IM; verhoogd risico op bijwerkingen door verhoging van de plasmaconcentratie van Fenytoïne. Monitoren op bijwerkingen coördinatiestoornis van de spieren, oogtrilling, spraakstoornis, slaperigheid of huiduitslag. Overwegen de standaarddosis te verlagen. (Clinical Pharmacogenetics Implementation Consortium,CPIC). Ook CYP2C19 is van invloed op het metabolisme van Fenytoine. Fenytoine induceert de activiteit van CYP2B6. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C9-IM; verminderde activiteit van CYP2C9 kan leiden tot een toxisch effect van Fosfenytoine. Overwegen te starten met een lagere dosering. (Clinical Pharmacogenetics Implementation Consortium, CPIC)',
                'CYP2C9-IM; het fenotype verhoogt de effectiviteit van Glimepiride.',
                'CYP2C9-IM; het fenotype verhoogt de effectiviteit van Gliclazide zonder het risico op hypoglykemie significant te verhogen.',
                'CYP2C9-IM; het fenotype verhoogt de effectiviteit van Glimepiride',
                'CYP2C9-IM; het fenotype leidt tot een hoger risico op bijwerkingen vanwege de hogere plasmaconcentraties van Siponimod. Heroverwegen gebruik bij gelijktijdig gebruik van Modafinil (CYP3A4 inductor). Een matige CYP3A4-inductor leidt bij CYP2C9 *1/*3 tot een afname in de blootstelling van Siponimod met 49%.',
                'CYP2C9-IM; er zijn onvoldoende gegevens over de klinische gevolgen van een verhoogde plasmaconcentratie van Tolbutamide.',
                'CYP2C9-IM; het fenotype vermindert de omzetting van Warfarine in inactieve metabolieten. Hierdoor kan het bloedingsrisico verhoogd zijn. Advies DPWG: gebruik 65% van de normale startdosering. De berekening van de genotype specifieke start- en onderhoudsdosering kan worden uitgevoerd met behulp van een algoritme. Algoritmes voor Kaukasische patiënten bevatten meestal alleen het *2- en *3-allel. Als de activiteit van het verminderd actieve allel vergelijkbaar is met dat van *2 of *3, kan het algoritme worden ingevuld alsof *1/*2 of *1/*3 aanwezig is. Zie https://www.knmp.nl/patientenzorg/medicatiebewaking/farmacogenetica voor rekenmodules in de vorm van Excelfiles voor de orale en de hieraan gelijk zijnde intraveneuze doseringen. Vanaf dag 6 kan het standaard genotype-onafhankelijke algoritme worden gebruikt. Voor patiënten met een Afrikaanse of (Oost-)Aziatische achtergrond zijn aparte doseringsalgoritmes ontwikkeld. Ook metabolisme via VKORC1. Voor het vaststellen van de dosering voor patiënten zie ook www.warfarindosing.org. en  http://www.warfarindoserevision.com. De activiteit van CYP1A2 , CYP2C19 en VKORC1 zijn ook van invloed op het metabolisme van Warfarine. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_CYP2C9_PM": ['Celecoxib', 'Diclofenac', 'Flurbiprofen', 'Warfarine'],
            "effectList_CYP2C9_PM": [
                'CYP2C9-PM; sterke kans op bijwerkingen van Celecoxib. Ook metabolisering via CYP3A4.',
                'CYP2C9-PM; risico op sterke bijwerkingen; een alternatief medicijn overwegen.',
                'CYP2C9-PM; sterk verhoogd risico op bijwerkingen bij zeer hoge plasmaconcentraties van Flurbiprofen. Er is echter geen richtlijn voor dit medicijn om aanpassen van de dosering te kunnen berekenen. Een ander medicijn overwegen.',
                'CYP2C9-PM; ook metabolisme via VKORC1. Voor het vaststellen van de dosering voor patiënten zie www.warfarindosing.org. en  http://www.warfarindoserevision.com. Ook doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_CYP2C19_IM": ['Amitriptyline', 'Atomoxeline', 'Brivaracetam', 'Carisoprodol', 'Citalopram',
                                   'Clobazam',
                                   'Clomipramine', 'Clopidogrel', 'Escitalopram', 'Esomeprazol', 'Diazepam', 'Doxepine',
                                   'Fenytoïne', 'Flibanserine', 'Imipramine', 'Lansoprazol', 'Moclobemide', 'Omeprazol',
                                   'Pantoprazol', 'Rabeprazol', 'Sertraline', 'Trimipramine', 'Venlafaxine',
                                   'Voriconazol',
                                   'Wafarin'],
            "effectList_CYP2C19_IM": [
                'CYP2C19-IM; metabolisme via CYP2D6 en in mindere mate CYP2C19, CYP1A2 en CYP2C9; het verlaagde metabolisme van CYP2C19 is aandachtspunt. Dit fenotype heeft effect op de blootstelling aan Amitriptyline, maar niet op die aan Amitriptyline + de actieve metaboliet Nortriptyline, die effect en bijwerkingen bepaalt. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-IM Omzetting via CYP2D6 in het actieve 4OH-atomoxetine en wordt geïnactiveerd via CYP2C19. De plasmaconcentraties zijn niet afwijkend bij CYP2C19-IM',
                ' CYP2C19-IM; versterkt de werking van Brivaracetam.\xa0Volgens het door de Amerikaanse Food and Drug Administration (FDA) goedgekeurde medicijnlabel kan een verlaagde dosis nodig zijn.',
                'CYP2C19-IM; het verlaagde metabolisme verhoogt het risico op bijwerkingen..',
                'CYP2C19-IM; toename van de plasmaconcentratie van Citalopram. Ook metabolisme via CYP2D6 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-IM; verhoogd risico op bijwerkingen. De door de FDA goedgekeurde etikettekst voor ONFI stelt dat "de aanvangsdosis bij patiënten, waarvan bekend is dat ze slechte CYP2C19-metabolisers zijn, 5 mg/dag zou moeten zijn”. Afhankelijk van de effectiviteit kan dosering worden verhoogd tot een maximale dagelijkse dosis van 40 mg.',
                'CYP2C19-IM; dit fenotype verhoogt de plasmaconcentratie van Clomipramine, maar niet de plasmaconcentratie van clomipramine+desmethylclomipramine, die bijwerkingen en effectiviteit bij depressie bepaalt. De verhoging van de plasmaconcentratie van Clomipramine is gunstig voor de effectiviteit bij angststoornissen en obsessief-compulsieve stoornis. Ook metabolisme via CYP2D6 en in mindere mate via CYP3A4 en CYP1A2. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-IM; het effect van Clopidogrel bij het voorkomen van aggregatie van bloedplaatjes is waarschijnlijk laag (Clinical Pharmacogenetics Implementation Consortium, CPIC).  Alternatieve antiplatelet therapie (indien niet gecontra-indiceerd), bijv. Prasugrel of Ticagrelor, wordt aanbevolen. Het risico op ernstige cardio- en cerebrovasculaire incidenten is verhoogd bij een dotterbehandeling of stentplaatsing (percutane coronaire interventie) en bij patiënten met een beroerte of TIA, doordat bij dit fenotype de activering van clopidogrel vermindert. Bij andere patiënten zijn geen negatieve klinische gevolgen gevonden. Prasugrel, Ticagrelor en Acetylsalicylzuur/dipyridamol worden niet of in mindere mate door CYP2C19 gemetaboliseerd. In mindere mate ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-IM; het fenotype leidt tot een toename van de plasmaconcentratie van Escitalopram leidt. Bij aanpassen van de dosering kunnen de verhoogde plasmaconcentratie en het theoretisch verhoogde risico op QT-verlenging teniet worden gedaan. Dit betreft 75% van de normale maximumdosering.',
                'CYP2C19-IM; het fenotype leidt wel tot een hogere plasmaconcentratie Esomeprazol, maar er is onvoldoende bewijs voor een effect op therapeutische effectiviteit en bijwerkingen.',
                'CYP2C19-IM; kans op minder snel uit anesthesie komen. Er zijn echter onvoldoende onderzoeksgegevens beschikbaar om aanpassen van de dosering te kunnen berekenen. In mindere mate ook metabolisme via CYP3A4.',
                'CYP2C19-IM; verlaagd metabolisme maar aandeel van CYP2C19 is beperkt. Voornamelijk metabolisme via CYP2D6. In mindere mate ook metabolisme via CYP1A2 en CYP3A4. Dit fenotype heeft effect op de blootstelling aan Doxepine, maar niet op die aan Doxepine + de actieve metaboliet Nordoxepine, die effect en bijwerkingen bepaalt. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-IM; mogelijke kans op bijwerkingen. Grote variatie is vastgesteld tussen patiënten met gelijke doseringen. Echter metabolisme voornamelijk via CYP2C9. Fenytoïne induceert de activiteit van CYP2B6.  Doseringsadvies in Kennisbank KNMP aanwezig.',
                ' CYP2C19-IM; versnelt de omzetting van Flibanserine met mogelijke bijwerkingen als gevolg. (bijv. Hypotensie)',
                ' CYP2C19-IM; het fenotype verhoogt wel de plasmaconcentratie van Imipramine, maar niet de plasmaconcentratie van imipramine+Desipramine, die effectiviteit en bijwerkingen bepaalt. Ook CYP2D6 is van invloed op het metabolisme. In mindere mate ook nog metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-IM; het fenotype leidt tot hogere plasmaconcentratie Lansoprazol en tot een toename van de therapeutische effectiviteit zonder dat de bijwerkingen toenemen.',
                'CYP2C19-IM;\xa0hoewel de plasmaconcentratie van Moclobemide kan stijgen, als gevolg van verlaagd metabolisme van CYP2C19, leidt dit, voor zover bekend, niet tot een verhoogde kans op bijwerkingen.',
                'CYP2C19-IM; de hogere plasmaconcentratie Omeprazol leidt tot een toename van de therapeutische effectiviteit zonder dat de bijwerkingen toenemen. Metabolisme van Omeprazol ook via CYP3A4. Omeprazol induceert de activiteit van CYP1A2. Doseringsadvies in kennisbank KNMP aanwezig.',
                'CYP2C19-IM; het fenotype leidt tot hogere plasmaconcentratie Pantoprazol en een toename van de therapeutische effectiviteit zonder dat de bijwerkingen toenemen.',
                'CYP2C19-IM; de hogere plasmaconcentratie Rabeprazol leidt niet tot een toename van de bijwerkingen.',
                'CYP2C19-IM; het fenotype heeft een gering effect op de plasmaconcentratie Sertraline, maar dit effect lijkt niet klinisch relevant. Er is geen effect op bijwerkingen gevonden. Ook metabolisme via CYP2D6, CYP2C9, CYP2C19, CYP2B6 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                ' CYP2C19-IM: verlaagd metabolisme van Trimipramine',
                'CYP2C19-IM; Voornamelijk metabolisme via CYP2D6 en in mindere mate via CYP2C19 en CYP3A4. De verminderde omzetting van Venlafaxine wordt extra versterkt als sprake is van CYP2D6-IM. Het verlaagde metabolisme van CYP2C19 is aandachtspunt bij voorschrijven als er sprake is van bijwerkingen. Doseringsadvies in Kennisbank KNMP aanwezig.',
                ' CYP2C19-IM;  het fenotype kan de omzetting van Voriconazol verminderen en daardoor de plasmaconcentratie verhogen. Dit kan mogelijk leiden tot een betere effectiviteit of tot een toename van het risico op bijwerkingen. Monitoren van de plasmaconcentratie. Ook CYP2C9 en Cyp3A4 zijn betrokken bij het metabolisme van Voriconazol. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-IM; voor patiënten van Europese afkomst kan de applicatie voor het berekenen van de dosis (beschikbaar op http://www.warfarindoserevision.com), die rekening houdt met CYP2C9- en VKORC1-genotypen worden gebruikt. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_CYP2C19_PM": ['Carisoprodol', 'Clobazam', 'Clopidogrel', 'Diazepam', 'Sertraline'],
            "effectList_CYP2C19_PM": [
                'CYP2C19-PM; sterk verhoogd risico op bijwerkingen als gevolg van het ontbreken van metabole activiteit. Een ander medicijn overwegen.',
                'CYP2C19-PM; sterk verhoogd risico op bijwerkingen De door de FDA goedgekeurde etikettekst voor ONFI stelt dat "de aanvangsdosis bij patiënten waarvan bekend is dat ze slechte of geen CYP2C19-metabolisers zijn, 5 mg/dag zou moeten zijn”. ',
                'CYP2C19-PM; sterk verhoogd therapeutisch risico, vanwege het ontbreken van metabole activiteit. Clopidogrel wordt niet aanbevolen. Prasugrel of Ticagrelor overwegen. Deze medicijnen worden niet of nauwelijks door CYP2C19 omgezet. Doseringsadvies in kennisbank KNMP aanwezig.',
                'CYP2C19-PM; grote kans op minder snel uit anesthesie te komen. Er zijn echter onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Ook metabolisme via CYP3A4.',
                'CYP2C19-PM; nauwelijks of geen activiteit; Ook metabolisme via CYP2D6, CYP2C9, CYP2B6 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_CYP2C19_RM": ['Amitriptyline', 'Carisoprodol', 'Citalopram', 'Clomipramine', 'Clopidogrel',
                                   'Dexlansoprazol', 'Diazepam', 'Doxepine', 'Escitalopram', 'Esomeprazol',
                                   'Imipramine',
                                   'Lansoprazol', 'Moclobemide', 'Omeprazol', 'Pantoprazol', 'Rabeprazol', 'Sertraline',
                                   'Voriconazol'],
            "effectList_CYP2C19_RM": [
                'CYP2C19-RM; verminderd effect van Amitriptyline. Overwegen een alternatief geneesmiddel te kiezen, dat niet door CYP2C19 wordt gemetaboliseerd. (bijvoorbeeld Nortriptyline of Desipramine). Metabolisme voornamelijk via CYP2D6 en in mindere mate via CYP1A2 en CYP2C9. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-RM; snel metabolisme; verhoogd risico op bijwerkingen. Carisoprodol moet met de nodige voorzichtigheid worden toegediend.',
                'CYP2C19-RM; de verhoogde enzymactiviteit leidt tot lage plasma concentraties en dus tot minder effect van Citalopram. Een alternatief geneesmiddel (bijvoorbeeld Fluoxetine of Paroxetine) overwegen. In mindere mate ook metabolisme via CYP2D6 en CYP3A4. Doseringsadvies in Kennisbank van de KNMP aanwezig.',
                'CYP2C19-RM; Voornamelijk metabolisme via CYP2D6. Een alternatief medicijn overwegen, dat niet door CYP2C19 wordt gemetaboliseerd, bijvoorbeeld Nortrptyline of Desipramine. In mindere mate ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-RM; personen met RM-profiel kunnen baat hebben bij de verhoogde plasmaconcentratie van de werkzame stof bij inname van een standaarddosis. Er is echter kans op bloedingen. Overweeg Prasugrel of Tigracelor. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-RM; risico op onvoldoende werking van Dexlansoprazol en het optreden van bijwerkingen. Overwegen de dosering te verhogen.',
                'CYP2C19-RM; kans op te snel uit anesthesie komen. Er zijn echter onvoldoende onderzoeksgegevens beschikbaar om de dosering te kunnen berekenen. Ook CYP3A4 is van invloed op het metabolisme.',
                'CYP2C19-RM; de werking van Doxepine is verminderd. Een alternatief medicijn overwegen, dat niet wordt gemetaboliseerd via CYP2C19, bijvoorbeeld Nortrptyline of Desipramine. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-RM; de verhoogde enzymactiviteit leidt tot lage plasma concentraties en dus tot minder effect van Escitalopram. Een alternatief geneesmiddel (bijvoorbeeld Fluoxetine of Paroxetine) overwegen. In mindere mate ook metabolisme via CYP2D6 en CYP3A4. Doseringsadvies in Kennisbank van de KNMP aanwezig.',
                'CYP2C19-RM; risico op onvoldoende werking van Esomeprazol en het optreden van extra bijwerkingen. Esomeprazol wordt ook gemetaboliseerd via CYP3A4. Overwegen de dosering te verhogen.',
                'CYP2C19-RM; wisselwerking bij metabolisme met CYP2D6. In mindere mate ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank van de KNMP.',
                'CYP2C19-RM; snel metabolisme, dat leidt tot risico op onvoldoende effect van Lansoprazol. Lansoprazol wordt ook gemetaboliseerd via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-RM; wel effect op de bloedspiegel maar onvoldoende onderzoeksgegevens beschikbaar om dosering te kunnen berekenen. Moclobemide wordt voornamelijk gemetaboliseerd via CYP2C19. Ook CYP2D6 is betrokken bij het metabolisme.',
                'CYP2C19-RM; snel metabolisme, dat leidt tot een groot risico op onvoldoende effect van Omeprazol. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-RM; leidt tot te snelle afname van concentraties Pantoprazol met het risico op onvoldoende bescherming.  Ook rekening houden met de eventuele activiteit van CYP2D6, CYP1A2, CYP2C9 en CYP1A2, die in mindere mate een rol spelen bij de metabolsering. Bij gebruik van Panntoprazol letten op magnesium en vitamine B12-spiegels.  Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-RM; verhoogd risico op onvoldoende effect van Rabeprazol. Overwegen de dosering te verhogen. In mindere mate ook metabolisme via CYP3A4.',
                'CYP2C19-RM; risico op onvoldoende effect; ook metabolisme via CYP2D6, CYP2C9, CYP2B6 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                ' CYP2C19-RM; bij dit fenotype kans op onvoldoende effect. Een alternatief medicijn overwegen, dat niet overwegend afhankelijk is van het metabolisme via CYP2C19. Ook CYP2C9 en Cyp3A4 zijn betrokken bij het metabolisme van Voriconazol. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_CYP2C19_UM": ['Amitriptyline', 'Carisoprodol', 'Citalopram', 'Clomipramine', 'Clopidogrel',
                                   'Dexlansoprazol', 'Diazepam', 'Doxepine', 'Escitalopram', 'Esomeprazol',
                                   'Imipramine',
                                   'Lansoprazol', 'Moclobemide', 'Omeprazol', 'Pantoprazol', 'Rabeprazol', 'Sertraline',
                                   'Voriconazol'],
            "effectList_CYP2C19_UM": [
                'CYP2C19-UM; sterk verminderd effect of geen effect van Amitriptyline. Overwegen een alternatief geneesmiddel te kiezen, dat niet via CYP2C19 wordt gemetaboliseerd. (bijvoorbeeld Nortriptyline of Desipramine). Metabolisme voornamelijk via CYP2D6 en in mindere mate via CYP1A2 en CYP2C9. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; Ultrasnel metabolisme; sterk verhoogd risico op bijwerkingen. Carisoprodol moet met de nodige voorzichtigheid worden toegediend.',
                'CYP2C19-UM; de sterk verhoogde enzymactiviteit leidt tot lage plasma concentraties en dus tot minder effect van Citalopram. Een alternatief geneesmiddel (bijvoorbeeld Fluoxetine of Paroxetine) overwegen. In mindere mate ook metabolisme via CYP2D6 en CYP3A4. Doseringsadvies in Kennisbank van de KNMP aanwezig.',
                'CYP2C19-PM; Voornamelijk metabolisme via CYP2D6. Een alternatief medicijn overwegen, dat niet door CYP2C19 wordt gemetaboliseerd, bijvoorbeeld Nortrptyline of Desipramine. In mindere mate ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; personen met UM-profiel kunnen baat hebben bij de verhoogde plasmaconcentratie van de werkzame stof bij inname van een standaarddosis. Er is echter kans op bloedingen. Overweeg Prasugrel of Tigracelor. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; groot risico op onvoldoende werking van Dexlansoprazol en het optreden van bijwerkingen. Overwegen de dosering te verhogen.',
                'CYP2C19-UM; kans op te snel uit anesthesie komen. Er zijn echter onvoldoende onderzoeksgegevens beschikbaar om de dosering te kunnen berekenen. Ook via CYP3A4 metabolisme van Diazepam.',
                'CYP2C19-UM; de werking van Doxepine is nauwelijks of niet aanwezig. Een alternatief medicijn overwegen, dat niet wordt gemetaboliseerd via CYP2C19, bijvoorbeeld Nortrptyline of Desipramine. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; de verhoogde enzymactiviteit leidt tot lage plasma concentraties en dus tot minder effect van Escitalopram. Een alternatief geneesmiddel (bijvoorbeeld Fluoxetine of Paroxetine) overwegen. In mindere mate ook metabolisme via CYP2D6 en CYP3A4. Doseringsadvies in Kennisbank van de KNMP aanwezig.',
                'CYP2C19-UM; hoog risico op onvoldoende werking van Esomeprazol en het optreden van extra bijwerkingen. Esomeprazol wordt ook gemetaboliseerd via CYP3A4. Overwegen de dosering te verhogen.',
                'CYP2C19-UM; wisselwerking bij metabolisme met CYP2D6. In mindere mate ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank van de KNMP aanwezig.',
                'CYP2C19-UM; Ultrasnel metabolisme, dat leidt tot een groot risico op onvoldoende effect van Lansoprazol. Lansoprazol wordt ook gemetaboliseerd via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; wel effect op de bloedspiegel maar onvoldoende onderzoeksgegevens beschikbaar om dosering te kunnen berekenen. Metabolisme voornamelijk via CYP2C19. Ook CYP2D6 is betrokken bij het metabolisme van Moclobemide.',
                'CYP2C19-UM; ultrasnel metabolisme, dat leidt tot een groot risico op onvoldoende werking van Omeprazol. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; leidt tot te snelle afname van concentraties Pantoprazol met het risico op onvoldoende bescherming.  Ook rekening houden met de metabolisme activiteit van CYP2D6, CYP1A2, CYP2C9 en CYP1A2, die in mindere mate een rol spelen bij het metabolisme. Bij gebruik van Pantoprazol letten op magnesium en vitamine B12-spiegels. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; sterk verhoogd risico op een onvoldoende werking van Rabeprazol. Overwegen de dosering te verhogen. In mindere mate ook metabolisme via CYP3A4.',
                'CYP2C19-UM; hoog risico op onvoldoende werking; ook metabolisme via CYP2D6, CYP2C9, CYP2B6 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C19-UM; bij dit fenotype kans op onvoldoende dan wel geen werking. Een alternatief medicijn overwegen, dat niet overwegend afhankelijk is van het metabolisme via CYP2C19. Ook CYP2C9 en Cyp3A4 zijn betrokken bij de metabolisering van Voriconazol. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_COMT_IM": ['Amfetamine', 'Dexmethylfenidate', 'Dextroamfetamine', 'Lisdexamfetamine',
                                'Methylfenidate',
                                'Noradrenaline'],
            "effectList_COMT_IM": [
                'COMT-IM; verminderde activiteit en dus risico op verminderde werking en kans op bijwerkingen.',
                'COMT-IM; risico op verminderde werking van Dexmethylfenidate.',
                'COMT-IM; risico op verminderde werking van Dextroamfetamine en kans op bijwerkingen.',
                'COMT-IM; risico op verminderde werking van Lisdexamfetamine en kans op bijwerkingen.',
                'COMT-IM; risico op verminderde werking van Methylfenidate.',
                'COMT-IM; metabolisering voornamelijk via COMT. COMT IM is aandachtspunt bij optreden bijwerkingen.'],
            "symbolList_COMT_IM": [1, 1, 1, 1, 1, 2],
            "medList_CYP2A6_IM": ['cafeine', 'Tegafur', 'Efavirenz', 'Valproinezuur'],
            "effectList_CYP2A6_IM": [
                'Cyp2A6 is belangrijk bij de omzetting van Paraxanthine. Bij niet-rokers is het metabolisme van paraxanthine lager in vergelijking met NM en PM',
                'CYP2A6 is belangrijk bij de biotransformatie van Tegafur in 5 FU. Ook de genen CYP1A2, CYP2C8, CYP2C9 hebben daarop invloed. Remming van CYP2A6 vermindert de vorming van 5 FU.',
                'Samen met CYP2B6 is verminderde activiteit van CYP2A6 van invloed op de werking van Efavirenz. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2C9 is het voornaamste gen bij het metabolisme van Valproinezuur. Daarnaast draagt CYP2A6 voor 50% bij aan de Valproinezuur 3 hydroxylering. Doseringsadvies in kennisbank KNMP aanwezig.'],
            "medList_CYP2B6_IM": ['Bupropion', 'Efavirenz', 'Cyclofosfamide', 'Methadon', 'Prasugorel', 'Propofol',
                                  'Sertraline'],
            "effectList_CYP2B6_IM": [
                'CYP2B6: IM vertraagde afbraak van Bupropion. Gevonden variaties kunnen kans geven op bijwerkingen of de bekende bijwerkingen versterken..',
                'CYP2B6: IM; ook metabolisme via CYP2A6 en CYP3A4. Verminderde activiteit van CYP2B6 en CYP2A6 geeft risico op variabiliteit van Efavirenz en daardoor risico op bijwerkingen. Aanpassen van de dosering overwegen. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2B6: IM verminderde activatie van Cyclofosfamide. Gevonden variaties kunnen kans geven op bijwerkingen of de bekende bijwerkingen versterken..',
                'CYP2B6-IM; voor personen met dit fenotype kan worden volstaan met een lagere dosering, maar andere genetische en klinische factoren kunnen leiden tot toepassen van de standaard dosering. Metabolisme ook via CYP2D6 en CYP3A4.',
                'CYP2B6-IM; de respons van Prasugrel kan mogelijk verminderd zijn, maar wetenschappelijk bewijs hiervoor is nog onvoldoende. Bovendien beïnvloeden ook andere genetische factoren de stofwisseling.',
                'CYP2B6: IM; vertraagde afbraak van Propofol. Ook metabolisme via CYP2C9,en UGT1A9. CYP2B6-IM past bij hogere spiegels Propofol en een verhoogde kans op bijwerkingen.',
                'CYP2B6 is betrokken bij de omzetting van Sertraline. Ook metabolisme via  CYP2C9, CYP2C19, CYP2D6 en CYP3A4. Voor klinische effecten moeten waarschijnlijk meerdere enzymen een variatie hebben. Op dit moment is er alleen een advies bij een verminderde activiteit van CYP2C19.'],
            "medList_DPYD_IM": ['Capecitabine', 'Fluorouracil', 'Tegafur '],
            "effectList_DPYD_IM": [
                'DPYD-IM; de verminderde activiteit geeft een verhoogd risico op ernstige  fluoropyrimidinetoxiciteit. Bepalen dosering aan de hand van Kennisbank KNMP en Clinical Pharmacogenetics Implementation Consortium (CPIC).',
                'DPYD-IM; de verminderde dihydropyrimidinedehydrogenase-activiteit geeft een verhoogd risico op ernstige fluoropyrimidinetoxiciteit. Overwegen de dosis met 25 tot 50% te verlagen, gevolgd door titratie van de dosis op basis van toxiteit: Clinical Pharmacogenetics Implementation Consortium (CPIC))',
                'DPYD-IM, activiteitsscore 1,5;volgens Clinical Pharmacogenetics Implementation Consortium (CPIC), is er voor gebruik van Tegafur slechts beperkt bewijs met betrekking tot de impact van DPYD-varianten op het toxiciteitsrisico van Tegafur.'],
            "medList_IFNL3_IM": ['Boceprevir', 'Ribavirin', 'Telaprevir'],
            "effectList_IFNL3_IM": [
                'IFNL3-WT/MT; dit fenotype wordt geassocieerd met een ongunstige respons op een combinatie van Peginterferon alpha 2 (PEG-IFN alpha) en Ribavirin (RBV) bij Hepatitus B Type 1 (Clinical Pharmacogenetics Implementation Consortium). ',
                'IFNL3-WT/MT; dit fenotype wordt geassocieerd met een ongunstige respons op een combinatie van Peginterferon alpha 2 (PEG-IFN alpha) en Ribavirin (RBV) bij Hepatitus B Type 1 (Clinical Pharmacogenetics Implementation Consortium). ',
                'IFNL3-WT/MT; dit fenotype wordt geassocieerd met een ongunstige respons op de behandeling van Hepatitis C virus in combinatie van Peginterferon alpha 2 (PEG-IFN alpha) en Ribavirin (RBV), (Clinical Pharmacogenetics Implementation Consortium). '],
            "medList_OPRM1_IM": ['Fentanyl'],
            "effectList_OPRM1_IM": [
                'OPRM1-IM; wordt bij standaarddosering geassocieerd met verminderd effect bij epidurale analgesie. Overwegen de dosering te verhogen.'],
            "medList_TPMT_IM": ['Azathioprine', 'Mercaptopurine', 'Thioguanine'],
            "effectList_TPMT_IM": [
                'TPMT-IM; risico op beenmergdepressie bij standaarddosering van Azathioprine. Ook NUDT15 is betrokken bij het metabolisme. Doseringsadvies in kennisbank KNMP aanwezig.',
                'TPMT-IM; risico op myelosuppresie bij standaarddosering van Mercaptopurine. Ook NUDT15 is betrokken bij het metabolisme. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'TPMT-IM; risico op myelosuppresion bij standaarddosering van Thioguanine. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_UGT1A1_IM": ['Atazanavir', 'Belinostat', 'Irinotecan'],
            "effectList_UGT1A1_IM": [
                'UGT1A1-IM; de verminderde omzetting van Atazanavir vergroot de kans op geelzucht. Om die reden een alternatief medicijn in overweging nemen.',
                'UGT1A1-IM; de verminderde omzetting van Belinostat leidt tot de overweging om de start dosering te verlagen om de gewenste werking van Belinostat te bereiken en het risico op bijwerkingen te verminderen. Er zijn echter geen doseringsadviezen beschikbaar.',
                'UGT1A1-IM; dit enzym speelt een belangrijke rol bij het afbreken van oude en beschadigde bloedcellen en andere afvalstoffen waaronder Bilirubine. Omdat UGT1A1 een verminderde activiteit toont vindt dit proces onvoldoende of niet plaats. Bij UGT1A1-IM kan de dosering van Irinotecan verminderd worden om ernstige bijwerkingen te voorkomen. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_VKORC1_IM": ['Acenocoumarol', 'Fenprocoumon', 'Warfarine'],
            "effectList_VKORC1_IM": [
                'VKORC1-IM; verlaagde expressie en dus een verhoogde gevoeligheid voor bloedverdunners op cumarine basis. Activiteit van CYP2C9 is ook belangrijk voor de omzetting van Acenocoumarol. In geringere mate ook metabolisme via CYP1A2 en CYP4F2. Vitamine K2 niet gebruiken samen met Acenocoumarol. Doseringsadvies in Kennisbank KNMP aanwezig. ',
                'VKORC1-IM; verlaagde expressie en dus een verhoogde gevoeligheid voor bloedverdunners op cumarine basis. De activiteit van CYP2C9 is ook belangrijk voor de omzetting van Fenprocoumon. In geringere mate ook metabolisme via CYP1A2 en CYP4F2. Vitamine K2 niet gebruiken samen met Fenprocoumon. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'VKORC1-IM; ook metabolisering via CYP2C9 en in mindere mate via CYP1A2 en CYP2C19. Voor het vaststellen van de dosering voor patiënten zie www.warfarindosing.org. en  http://www.warfarindoserevision.com. Ook doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_CYP2D6_IM": ['Amitriptyline', 'Cevimeline', 'Clomipramine', 'Codeïne', 'Desipramine',
                                  'Dextromethorphan-Quinidine', 'Donepezil', 'Doxepine', 'Doxepine', 'Flecainide',
                                  'Fluvoxamine', 'Hydrocodon', 'Iloperidon', 'Metoprolol', 'Nortriptyline', 'Oxycodon',
                                  'Paroxetine', 'Perfenazine', 'Pimozide', 'Propafenon', 'Protriptyline', 'Risperidon',
                                  'Tamoxifen', 'Thioridazine', 'Tramadol', 'Venlafaxine', 'Vortioxetine',
                                  'Zuclopentixol'],
            "effectList_CYP2D6_IM": [
                'CYP2D6-IM; verminderde afbraak van Amitriptyline; verhoogde de kans op bijwerkingen. Overwegen een alternatief geneesmiddel te kiezen (bijvoorbeeld Citalopram, Sertraline). Ook metabolisme via CYP2C19 en in mindere mate via CYP1A2 en CYP2C9. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM, dus een verhoogd risico op bijwerkingen als gevolg van een te geringe omzetting van Cevimeline. Onvoldoende onderzoeksgegevens beschikbaar om de dosering te kunnen berekenen.',
                'CYP2D6-IM; verhoogd risico op hogere waarde Clomipramine en dus risico op bijwerkingen. Aanbeveling door het Clinical Pharmacogenetics Implementation Consortium (CPIC): Aangezien tricyclische antidepressiva vergelijkbare farmacokinetische eigenschappen hebben, kan de doseringsrichtlijn voor Amitriptyline/Nortriptyline voor CYP2D6-IM ook toegepast worden op andere tricyclische antidepressiva, waaronder Clomipramine. Een 25% verlaging van de aanbevolen startdosis kan overwogen worden. Deze aanbeveling geldt niet voor behandeling van bijvoorbeeld neuropathische pijn.',
                'CYP2D6-IM; kies voor analgesie een ander medicijn (bijvoorbeeld Acetaminophen, NSAID, morfine, niet Tramadol of Oxycodon). Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; de resulterende plasmaconcentraties als gevolg van de verminderde activiteit kunnen de kans op bijwerkingen vergroten. Overwegen, aan de hand van de plasmaconcentraties, een alternatief medicijn toe te passen.',
                'CYP2D6-IM; verhoogd risico op bijwerkingen en onvoldoende werking van Dextromethorphan-quinidine. Ook metabolisme via CYP3A4. Een alternatieve medicatie overwegen.',
                'CYP2D6-IM; verhoogd risico op bijwerkingen vanwege het verminderd metabolisme. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Ook metabolisme via CYP3A4. ',
                'CYP2D6-IM;  kans op bijwerkingen vanwege de verminderde activiteit. In mindere mate metabolisme via CYP1A2, CYP3A4 en CYP2C19. Doseringsadvies in Kennisbank KNMP aanwezig',
                'CYP2D6-IM; verhoogd risico op bijwerkingen vanwege de verminderde metabolisme-activiteit. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; verhoogde kans op bijwerkingen en risico op verminderd effect van het medicijn. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; verhoogd risico op bijwerkingen. Er zijn echter onvoldoende onderzoeksgegevens beschikbaar om aanpassing van de dosering te kunnen berekenen. Ook metabolisme via CYP1A2.',
                'CYP2D6-IM; verhoogd risico op onvoldoende werking van Hydrocodon. Overwegen een alternatieve medicatie te kiezen.',
                'CYP2D6-IM; verhoogd risico op onvoldoende effect van Iloperidon en optreden van bijwerkingen als gevolg van verhoogde Iloperidon-concentraties.',
                'CYP2D6-IM; verhoogd risico op bijwerkingen. Voor hartfalen (indicatie): selecteer een ander geneesmiddel (bijvoorbeeld Bisoprolol, Carvedilol) of overwegen de dosering te verlagen. Bij andere indicaties: wees alert op bijwerkingen van het geneesmiddel (bijvoorbeeld Bradycardie, (te lage hartslag) of selecteer een ander geneesmiddel (bijvoorbeeld Atenolol, Bisoprolol). Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; verhoogde kans op bijwerkingen en onvoldoende effect van Nortriptyline. Ook metabolisme via CYP2C19. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; risico op verminderde werkzaamheid; overweeg alternatieve medicatie (geen Tramadol of Codeïne). Ook metabolisme via CYP3A4.  Doseringsadvies beschikbaar in Kennisbank KNMP.',
                'CYP2D6-IM; kans op verhoogde plasmaconcentraties Paroxetine en dus risico op toenemen van bijwerkingen. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; verhoogd risico op bijwerkingen. Er zijn onvoldoende onderzoeksgegevens beschikbaar om aanpassen van de dosering te kunnen berekenen.',
                'CYP2D6-IM; verhoogd risico op bijwerkingen vanwege hogere Pimozide concentraties. Ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; aanpassen dosering in reactie op de plasmaconcentratie of een alternatief medicijn (bijvoorbeeld Sotalol, Disopyramide, Kinidine of Amiodaron) overwegen. Ook metabolsme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-IM; verhoogde kans op bijwerkingen. Regelmatig de plasmaconcentratie controleren en afhankelijk van de aard of ernst van de bijwerkingen een alternatief medicijn overwegen.',
                'CYP2D6-IM; afhankelijk van eventuele comedicatie aanpassen dosering of een alternatief medicijn als bijvoorbeeld Quetiapine, Olanzapine of Clozapine overwegen. Metabolisme voornamelijk via CYP2D6 maar ook in minderde mate via CYP3A4.',
                'CYP2D6-IM; verhoogd risico op optreden van bijwerkingen. Metabolisme voornamelijk via CYP2D6, maar ook in mindere mate via CYP2C19, CYP2C9 en CYP3A4. Doseringsadvies beschikbaar in Kennisbank KNMP',
                'CYP2D6-IM; verhoogd risico op ernstige bijwerkingen bij een verhoogd Thioridazine-gehalte. Een alternatief medicijn dient te worden overwogen. In geringe mate ook metabolisme via CYP1A2 en CYP3A4.',
                'CYP2D6-IM; risico op bijwerkingen. Er zijn echter onvoldoende onderzoeksgegevens beschikbaar om de dosering te kunnen berekenen.',
                'CYP2D6-IM; risico op onvoldoende werking. Een alternatief medicijn overwegen. (geen Oxycodon of Codeïne). Ook metabolisme  via CYP1A2 en CYP2B6. Doseringsadvies beschikbaar in Kennisbank KNMP.',
                'CYP2D6-IM; een alternatief medicijn overwegen (bijvoorbeeld Citalopram of Sertraline).  Ook metabolisme via CYP2C19 en CYP3A4. Doseringsadvies in Kennisbank van de KNMP aanwezig.',
                'CYP2D6-IM; verhoogd risico op bijwerkingen en verminderde werking van Vortioxetine. In geringe mate ook metabolisme via CYP2C9, CYP3A4 en CYP3A5. ',
                'CYP2D6-IM;  de resulterende plasmaconcentraties kunnen het risico op bijwerkingen verhogen.  Doseringsadvies beschikbaar in Kennisbank KNMP.'],
            "medList_CYP2D6_PM": ['Amitriptyline', 'Aripiprazol', 'Atomoxetine', 'Brexpiprazol', 'Cevimeline',
                                  'Clomipramine', 'Clozapine', 'Codeïne', 'Desipramine', 'Dextromethorfan-Quinidine',
                                  'Donepezil', 'Doxepine', 'Duloxetine', 'Eliglustat', 'Flecaïnide', 'Flecaïnide',
                                  'Fluoxetine', 'Fluvoxamine', 'Galantamine', 'Gefitinib', 'Haloperidol', 'Hydrocodon',
                                  'Iloperidon', 'Metoprolol', 'Nortriptyline', 'Oxycodon', 'Paroxetine', 'Perfenazine',
                                  'Pimozide', 'Propafenon', 'Protriptyline', 'Risperidon', 'Tamoxifen', 'Tamsulosin',
                                  'Tetrabenazine', 'Thioridazine', 'Tolterodine', 'Tramadol', 'Tropisetron',
                                  'Venlafaxine',
                                  'Vortioxetine', 'Zuclopenthixol'],
            "effectList_CYP2D6_PM": [
                'CYP2D6-PM; geen of nauwelijks afbraak van Amitriptyline; de resulterende plasmaconcentraties verhogen de kans op bijwerkingen. Een alternatief geneesmiddel overwegen (bijvoorbeeld Citalopram of Sertraline). Ook metabolisme via CYP2C19 en nog aanvullend gering metabolisme via CYP1A2 en CYP2C9. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; verhoogde kans op bijwerkingen en onvoldoende werking. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM, verhoogde kans op bijwerkingen en onvoldoende werking van Atomoxetine.  Ook metabolisme via CYP2C19. Doseringsadvies in Kennisbank KNMP aanwezig. ',
                'CYP2D6-PM; verhoogd risico op bijwerkingen vanwege verhoogde concentraties van Brexpiprazol. Als gelijktijdig een sterke/matige CYP3A4-remmer wordt gebruikt heeft dat invloed op het bepalen van de gewenste dosering.',
                'CYP2D6-PM; aanzienlijk verhoogd risico op bijwerkingen als gevolg van zeer hoge plasmaconcentraties van Cevimeline. Er zijn echter onvoldoende adviezen beschikbaar om aanpassen van de dosering te kunnen berekenen.',
                'CYP2D6-PM; de kans op toenemen van bijwerkingen. Een alternatief medicijn overwegen (bijvoorbeeld Citalopram of Sertraline). Ook metabolisme via CYP2C19 en in gering mate via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM;voornamelijk rekening houden met de activiteit van CYP1A2. Geringe metabolisme via CYP2D6, CYP2C19 en CYP3A4; beide genen zijn ook betrokken bij de omzetting van Clozapine. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; risico op sterk verminderde omzetting naar morfine. Overwegen bij analgesie een alternatief medicijn (bijv. Paracetamol, NSAID, Morfine; geen Tramadol of Oxycodon). Ook metabolisme via CP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                ' CYP2D6-PM; de resulterende plasmaconcentratie Desipramine kan de kans op bijwerkingen vergroten. Controleer de plasmaconcentratie of een alternatief medicijn overwegen.',
                'CYP2D6-PM; sterk verhoogd risico op bijwerkingen en onvoldoende werking. Ook CYP3A4 is betrokken bij de metabolisering. Een alternatief medicijn overwegen.',
                'CYP2D6-PM; risico op bijwerkingen echter er zijn onvoldoende onderzoeksgegevens  beschikbaar om aanpassen van de dosering te kunnen berekenen. Ook metabolisme via CYP3A4.',
                'CYP2D6-PM;  meer kans op bijwerkingen. Dosering aanpassen aan de plasmaconcentratie van Doxepine. In geringe mate zijn ook metabolisme via CYP1A2, CYP2C19 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; er is enig bewijs dat personen met CYP2D6-PM hogere concentraties Duloxetine kunnen hebben. Ook CYP1A2 is betrokken bij het metabolisme van Duloxetine. ',
                'CYP2D6-PM; aanzienlijk hoger risico op bijwerkingen bij dit fenotype. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; potentieel verminderde omzetting en dus risico op onvoldoende werking en het optreden van bijwerkingen. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; potentieel verminderde omzetting en dus risico op onvoldoende werking en het optreden van bijwerkingen. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; bij gebruik van Fluoxetine aandacht bij gelijktijdig gebruik van medicijnen, die worden gemetaboliseerd door CYP2D6. Fluoxetine is een krachtige remmer van de CYP2D6-enzymroute. Extra aandacht bij comedicatie met sommige antidepressiva, antipychotica (Fenothiazines) en de meeste antiaritmica (bijv. Propafenon en Flecaïnide).Ook metabolisme via CYP2C9 en in geringe mate via CYP1A2, CYP3A4 en CYP3A5. ',
                'CYP2D6-PM; een verhoogd risico op bijwerkingen door geen of nauwelijks activiteit van CYP2D6. Ook metabolisering via CYP1A2. Er zijn echter onvoldoende onderzoeksgegevens beschikbaar om aanpassen van de dosering te berekenen. ',
                'CYP2D6-PM; kans op bijwerkingen; er zijn echter onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen.',
                'CYP2D6-PM; verhoogd risico op bijwerkingen. Er zijn echter onvoldoende onderzoeksgegevens beschikbaar om aanpassen van de dosering te kunnen berekenen.',
                'CYP2D6-PM;  onvoldoende effect van Haloperidol en verhoogd risico op bijwerkingen. Een ander medicijn overwegen. Ook metabolisme via CYP3A4 en in geringe mate via CYP1A2. Doseringsadvies in Kennisbank KNMP aanwezig',
                'CYP2D6-PM; verhoogd risico op onvoldoende werking van Hydrocodon en het optreden van bijwerkingen. Een alternatief medicijn overwegen.',
                'CYP2D6-PM; verhoogd risico op bijwerkingen als gevolg van de verhoogde Iloperidon-concentraties.',
                'CYP2D6-PM; verhoogd risico op onvoldoende werking en optreden van ernstige bijwerkingen.  Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; hogere plasmaconcentraties verhogen de kans op bijwerkingen van Nortriptyline. Ook rekening houden met mogelijke activiteit UM van  CYP2C19. Als mogelijk alternatief Gabapentine overwegen. Bij CYP2D6 PM is maar 10-20 mg nodig van dit medicijn. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; risico op mogelijk sterk verminderde werking van Oxycodon. Ook metabolisme via CYP3A4. Er zijn onvoldoende onderzoeksgegevens beschikbaar om aanpassen van de dosering te berekenen. Overwegen een alternatief medicijn (geen Tramadol of Codeïne) toe te passen. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; kans op verhoogde plasmaconcentraties Paroxetine en optreden van bijwerkingen.  Doseringsadvies in Kennisbank KNMP aanwezig',
                'CYP2D6-PM; verhoogd risico op bijwerkingen als gevolg van verhoogde plasmaconcentraties Perphenazine. Er zijn nog onvoldoende onderzoeksgegevens beschikbaar om aanpassen van de dosering te kunnen berekenen.',
                'CYP2D6-PM; verhoogd risico op bijwerkingen als gevolg van hogere Pimozide concentraties. Ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; significant verhoogde plasmaconcentraties Propafenon kunnen leiden tot pro-aritmie, sterke bèta-adrenerge blokkerende activiteit en andere bijwerkingen. Ook CYP1A2 en CYP3A4 zijn betrokken bij de metabolisering. Doseringsadvies in Kennisbank KNMP aanwezig',
                'CYP2D6-PM; de resulterende hogere plasmaconcentraties Protriptyline kunnen de kans op bijwerkingen vergroten. Regelmatige controle van de plasmaconcentratie is wenselijk.',
                'CYP2D6-PM; er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatief medicijn overwegen (bijv. Quetiapine, Olanzapine, Clozapine). In geringe mate ook metabolisme via CYP3A4.',
                'CYP2D6-PM; geen of nauwelijks omzetting en dus verhoogd risico op bijwerkingen. Ook metabolisme via CYP3A4, CYP2C19 en CYP2C9. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; risico op bijwerkingen als gevolg van verminderd metabolisme. Wees alert op bijwerkingen en/of een alternatief medicijn overwegen bij een dosis Tamsulosin hoger dan 0,4 mg.',
                'CYP2D6-PM; aanzienlijk hoger risico op bijwerkingen als gevolg van verhoogde niveau Tetrabenazine.',
                'CYP2D6-PM; verhoogd risico op ernstige bijwerkingen vanwege een verhoogd Thioridazine-gehalte. In geringe mate ook metabolisme via CYP1A2 en CYP3A4. Overwegen een alternatief medicijn toe te passen.',
                'CYP2D6-PM; aanzienlijk hoger risico op bijwerkingen als gevolg van de hogere plasmaconcentratie Tolterodine. Er zijn echter nog onvoldoende onderzoeksgegevens beschikbaar om aanpassen van de dosering te kunnen berekenen.',
                'CYP2D6-PM; grote kans op verminderde werking van Tramadol. Bij overwegen van alternatieven, niet kiezen voor Oxycodon of Codeïne. Metabolisme ook via CYP2B6 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; er is enig bewijs is dat personen met CYP2D6 PM mogelijk hogere concentraties Tropisetron kunnen opbouwen.',
                'CYP2D6-PM; Overwegen een alternatief medicijn toe te passen (bijv. Citalopram, Sertraline). Metabolisme ook via CYP2C19 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-PM; aanzienlijk verhoogd risico op bijwerkingen. De maximaal aanbevolen dosis is 10 mg  per dag bij CYP2D6-PM. Ook geringe metabolisme via CYP2C9 en CYP3A4 en CYP3A5.',
                'CYP2D6-PM; risico op bijwerkingen is hoog. Overwegen een alternatief medicijn (Flupenthixol, Quetiapine, Olanzapine, Clozapine) toe te passen. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_CYP2D6_UM": ['Amitriptyline', 'Atomoxetine', 'Cevimeline', 'Clomipramine', 'Codeïne',
                                  'Desipramine',
                                  'Dextromethorfan-Quinidine', 'Donepezil', 'Doxepine', 'Eliglustat', 'Galantamine',
                                  'Haloperidol', 'Hydrocodon', 'Imipramine', 'Kinidine', 'Metoprolol', 'Nortriptyline',
                                  'Ondansetron', 'Paroxetine', 'Perfenazine', 'Propafenon', 'Protriptyline',
                                  'Risperidon',
                                  'Tamoxifen', 'Thioridazine', 'Tramadol', 'Trimipramine', 'Tropisetron', 'Venlafaxine',
                                  'Vortioxetine'],
            "effectList_CYP2D6_UM": [
                'CYP2D6-UM; verhoogd risico op onvoldoende werking van Amitriptyline bij dit fenotype. Een alternatief medicijn, dat niet wordt gemetaboliseerd via CYP2D6, overwegen. Ook metabolisme via CYP2C19 en in geringere mate via CYP1A2 en CYP2C9. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6:UM; kans op onvoldoende werking van Atomoxetine. Een alternatief medicijn overwegen (bijv. Methylfenidaat, Clonidine). Ook metabolisme via CYP2C19. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; verhoogd metabolisme van Cevimeline; de resulterende plasmaconcentraties kunnen de kans op onvoldoende werking vergroten. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. Een alternatieve medicatie overwegen.',
                'CYP2D6-UM; verhoogd metabolisme van Clomipramine kan leiden tot minder werking van Clomipramine. Overwegen een ander medicijn toe te passen; bijvoorbeeld Citalopram of Sertraline. Ook metabolismeg via CYP2C19 en in geringe mate via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; met dit fenotype wordt de omzetting van Codeine naar morfine versterkt. Gebruik is contra-geïndiceerd voor personen met een CYP2D6-UM fenotype. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; de resulterende plasmaconcentraties vergroten de kans op onvoldoende werking van Desipramine. Overwegen een alternatief medicijn toe te passen dat niet door CYP2D6 wordt gemetaboliseerd.',
                'CYP2D6-UM; verhoogd risico op onvoldoende werking. Overwegen de dosering te verhogen of een alternatief medicijn toe te passen. Ook metabolisme via CYP3A4.',
                'CYP2D6-UM; verhoogde metabole omzetting van Donepezil. Risico op onvoldoende werking van Donepezil en het optreden van bijwerkingen. Ook metabolisme via CYP3A4. ',
                'CYP2D6-UM; verhoogde de kans op onvoldoende werking van Doxepine. Een alternatief medicijn overwegen dat niet door CYP2D6 wordt gemetaboliseerd. In geringere mate ook metabolisme via CYP1A2, CYP2C19 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; sterk verhoogd risico op onvoldoende werking van Eliglustat en optreden van bijwerkingen. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; verhoogd risico op bijwerkingen en/of onvoldoende werking van Galantamine. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen.',
                'CYP2D6-UM; kans op verminderde werking van Haloperidol. Een ander medicijn overwegen (bijvoorbeeld Pimozide, Flupenthixol, Fluphenazine, Quetiapine, Olanzapine of Clozapine). Ook metabolisme via CYP3A4 en in geringe mate via CYP1A2. Doseringsadvies in Kennisbank KNMP aanwezig',
                'CYP2D6-UM; verhoogd risico op bijwerkingen en/of onvoldoende werking van Hydrocodon. Er zijn onvoldoende onderzoeksgegevens om aanpassen van de dosering te kunnen berekenen. ',
                'CYP2D6-UM; verhoogde de kans op onvoldoende werking van Imipramine. Een alternatief medicijn overwegen dat niet door CYP2D6 wordt gemetaboliseerd. Metabolisme van Imipramine ook via CYP2C19 en in mindere mate via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; Kinidine is een krachtige remmer van het CYP2D6-enzym, waardoor bijvoorbeeld UM IM wordt. Extra aandachtspunt als naast Kinidine ook andere medicijnen worden gebruikt, die door CYP2D6 worden gemetaboliseerd.',
                'CYP2D6-UM; sterk verhoogd risico op verminderde werking en optreden van bijwerkingen. Overwegen een ander medicijn toe te passen bijvoorbeeld Bisoprolol, dat niet wordt gemetaboliseerd via CYP2D6. Echter bij toepassen van Bisoprolol letten op nierfunctie. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; bij dit fenotype is sprake van een verminderde werking van Nortriptyline. Een ander medicijn overwegen, dat niet wordt gemetaboliseerd via CYP2D6. Ook metabolisme via CYP2C19. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; bij dit fenotype is er risico op verminderde werking van Ondansetron. Een alternatief medicijn, dat niet wordt gemetaboliseerd via CYP2D6 overwegen (bijvoorbeeld Granisetron). Ondansetron wordt ook gemetaboliseerd via CYP1A2 en CYP3A4.',
                'CYP2D6-UM; kans op onvoldoende werking van Paroxetine. Een alternatief medicijn overwegen, dat niet overwegend door CYP2D6 wordt gemetaboliseerd. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; door de resulterende plasmaconcentraties kans op verminderde werking van Perfenazine en optreden van bijwerkingen.',
                'CYP2D6-UM; verhoogd risico op onvoldoende werking van Pimozide en optreden van bijwerkingen. Ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; kans op onvoldoende werking van Propafenon. Een alternatief medicijn overwegen, dat niet overwegend door CYP2D6 wordt gemetaboliseerd, bijvoorbeeld Sotalol, Disopyramide,Quinidine of Amlodaron. Ook metabolisme via CYP1A2 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM;  de resulterende plasmaconcentraties kunnen de kans op onvoldoende werking van Protriptyline en het optreden van bijwerkingen vergroten. Overwogen kan worden de aanbevolen startdosis te verhogen.',
                'CYP2D6-UM;  kans op afwijkend effect van Risperidon in vergelijking met fenotype NM. Afhankelijk van het afwijkend effect en de bijwerkingen een alternatief medicijn overwegen, bijvoorbeeld Quetiapine, Olanzapine of Clozapine. In mindere mate ook metabolisering via CYP3A4.',
                'CYP2D6-UM; extra aandacht voor gelijktijdig gebruik van matige en sterke CYP2D6 remmers. Ook de genen CYP3A4, CYP2C19 en CYP2C9 zijn betrokken bij de omzetting van Tamoxifen. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; risico op verminderde werking en het optreden van bijwerkingen. Een alternatief medicijn overwegen of de aanbevolen startdosis verhogen. In geringe mate ook metabolisme via CYP1A2 en CYP3A4.',
                'CYP2D6-UM; hoog risico op ernstige bijwerkingen van Tramadol. Deze zijn te vergelijken met het effect van een overdosis Tramadol. Verlagen van de start dosering is nodig. Ook metabolisme via CYP3A4 en CYP2B6. Doseringsadvies in Kennisbank KNMP aanwezig',
                'CYP2D6-UM; kans op onvoldoende werking van Trimipramine. Een alternatief medicijn overwegen, dat niet overwegend door CYP2D6 wordt gemetaboliseerd.',
                'CYP2D6-UM; verhoogd risico op verminderde werking van Tropisetron en optreden van bijwerkingen. Een alternatief medicijn, dat niet overwegend wordt gemetaboliseerd via CYP2D6,  (bijvoorbeeld Granisetron) overwegen.',
                'CYP2D6-UM; verhoogd risico op onvoldoende werking van Venlafaxine en optreden van bijwerkingen. Ook metabolisme via CYP2C19 en CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'CYP2D6-UM; verhoogd risico op verminderde werking en optreden van bijwerkingen. Overwegen de standaard dosering te verhogen. In geringe mate ook metabolisme via CYP2C9, CYP3A4 en CYP3A5.'],
            "medList_CYP1A2_NM": ['Clozapine', 'Olanzapine'],
            "effectList_CYP1A2_NM": [
                'CYP1A2: NM; verhoogde induceerbaarheid. Maar ook niet genetische factoren spelen een rol. Bijvoorbeeld caffeine remt de werking van CYP1A2. Inductie (stimulering) van CYP1A2 wordt veroorzaakt door roken, het eten van geroosterd voedsel en koolsoorten. Rokers hebben bijvoorbeeld 1½ keer meer Clozapine nodig voor een therapeutische bloedspiegel in vergelijking met niet-rokers.',
                'CYP1A2: NM; verhoogde induceerbaarheid. Maar ook niet genetische factoren spelen een rol. Bijvoorbeeld caffeine remt de werking van CYP1A2. Inductie (stimulering) van CYP1A2 wordt veroorzaakt door roken, het eten van geroosterd voedsel en koolsoorten. Rokers hebben bijvoorbeeld 1½ keer meer  Olanzapine nodig voor een therapeutische bloedspiegel in vergelijking met niet-rokers.'],
            "medList_CYP3A5_heterozygoot": ['Tacrolimus'],
            "effectList_CYP3A5_heterozygoot": [
                'CYP3A5-IM;  (heterozygoot expresser) Personen met dit genotype lopen de kans dat Tacrolimus niet werkt bij standaarddosering. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_COMT_PM": ['Amfetamine', 'Dexmethylfenidate', 'Dextroamfetamine', 'Lisdexamfetamine',
                                'Methylfenidate',
                                'Noradrenaline', 'Pijnstillers'],
            "effectList_COMT_PM": [
                'COMT-PM; sterk verminderde activiteit en dus sterk verhoogd risico op onvoldoende werking en bijwerkingen.',
                'COMT-PM; sterk verhoogd risico op verminderde werking van Dexmethylfenidate.',
                'COMT-PM; sterk verhoogd risico op verminderde werking van Dextroamfetamine en meer kans op bijwerkingen.',
                'COMT-PM; sterk verhoogd risico op onvoldoende werking van Lisdexamfetamine en kans op (ernstige) bijwerkingen.',
                'COMT-PM; sterk verhoogd risico op onvoldoende werking van Methylfenidate.',
                'COMT-PM; metabolisering voornamelijk via COMT. COMT PM is extra aandachtspunt bij optreden bijwerkingen.',
                'COMT-PM; sterk verlaagde activiteit van het COMT enzym en dus een verlaagde morfine behoefte voor voldoende pijnstilling.'],
            "medList_CYP2B6_PM": ['Efavirenz'],
            "effectList_CYP2B6_PM": [
                'CYP2B6-PM; ook metabolisme via CYP3A4. Sterk risico op ernstige bijwerkingen. Een alternatief geneesmiddel moet worden overwogen. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_DPYD_PM": ['Capecitabine', 'Fluorouracil', 'Tegafur'],
            "effectList_DPYD_PM": [
                'DPYD-PM; de sterk verminderde activiteit geeft een verhoogd risico op ernstige fluoropyrimidinetoxiciteit. Bepalen dosering aan de hand van Kennisbank KNMP en Clinical Pharmacogenetics Implementation Consortium (CPIC).',
                'DPYD-PM; de sterk verminderde dihydropyrimidinedehydrogenase-activiteit geeft een verhoogd risico op zeer ernstige fluoropyrimidinetoxiciteit. Overwegen de dosis met 25 tot 50% te verlagen, gevolgd door titratie van de dosis op basis van toxiteit: Clinical Pharmacogenetics Implementation Consortium (CPIC)).',
                'DPYD-PM; volgens Clinical Pharmacogenetics Implementation Consortium (CPIC), is er voor gebruik van Tegafur slechts beperkt bewijs met betrekking tot de impact van DPYD-varianten op het toxiciteitsrisico van Tegafur en dus wordt geen dosis advies gegeven.'],
            "medList_SLCO1B1_PM": ['Simvastatine'],
            "effectList_SLCO1B1_PM": [
                'SLCO1B1-PM; zeer sterk verhoogd risico op door Simvastatine veroorzaakte myopathie. Een lagere dosis of een alternatieve statine (bijv. Pravastatine of Rosuvastatine) overwegen; regelmatig controleren van Kreatine waarden. Metabolisme via CYP3A4. Bij gebruik van statines letten op vitamine D-spiegel. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_TPMT_PM": ['Azathioprine', 'Mercaptopurine', 'Thioguanine'],
            "effectList_TPMT_PM": [
                'TPMT-PM; sterk verhoogd risico op beenmergdepressie bij standaarddosering van Azathioprine. Doseringsadvies in kennisbank KNMP aanwezig.',
                'TPMT-PM; sterk verhoogd risico op myelosuppresie bij standaarddosering van Mercaptopurine. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'TPMT-PM; sterk verhoogd risico op myelosuppresion bij standaarddosering van Thioguanine. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_UGT1A1_PM": ['Atazanavir', 'Belinostat', 'Irinotecan'],
            "effectList_UGT1A1_PM": [
                'UGT1A1-PM; de nauwelijks of geen omzetting van Atazanavir vergroot de kans op geelzucht. Om die reden een alternatief medicijn overwegen.',
                'UGT1A1-PM; de nauwelijks of geen omzetting van Belinostat leidt tot de overweging om de start dosering te verlagen om de werking van Belinostat te waarborgen en het risico op bijwerkingen te verminderen. Er zijn echter geen doseringsadviezen beschikbaar.',
                'UGT1A1-PM; dit enzym speelt een belangrijke rol bij het afbreken van oude en beschadigde bloedcellen en andere afvalstoffen waaronder Bilirubine. Omdat UGT1A1 nauwelijks of geen activiteit toont vindt dit proces onvoldoende of niet plaats. Afwezige activiteit van het UGT1A1 enzym is dus een belangrijk gegeven. Ook metabolisme via CYP3A4. Doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_VKORC1_PM": ['Acenocoumarol', 'Fenprocoumon', 'Warfarine'],
            "effectList_VKORC1_PM": [
                'VKORC1-PM; nauwelijks expressie voor Acenocoumarol en daarmee sterk verhoogde gevoeligheid voor bloedverdunners op cumarine basis.  Aanbevolen startdosering 50% van de standaard startdosis. De activiteit van CYP2C9 is ook belangrijk voor de omzetting van Acenocoumarol. In geringere mate ook metabolisme via CYP1A2 en CYP4F2. Vitamine K2 niet gebruiken samen met Acenocoumarol. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'VKORC1-PM; deze genetische variatie verhoogt de gevoeligheid voor Fenprocoumon sterk. De activiteit van CYP2C9 is ook belangrijk voor de omzetting van Fenprocoumon. Vitamine K2 niet gebruiken samen met Fenprocoumon. In geringere mate ook metabolisme via CYP1A2 en CYP4F2. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'VKORC1-PM; ook metabolisering via CYP2C9 en in mindere mate via CYP1A2 en CYP2C19. Voor het vaststellen van de dosering voor patiënten zie www.warfarindosing.org. en  http://www.warfarindoserevision.com. Ook doseringsadvies in Kennisbank KNMP aanwezig.'],
            "medList_SLCO1B1_IM": ['Simvastatine'],
            "effectList_SLCO1B1_IM": [
                'SLCO1B1: IM; matig verhoogd risico op myopathie bij gebruik van een dosis van 40 mg/ dag of meer simvastatine. Een verminderde dosering of een alternatief statinegeneesmiddel moet worden overwogen. Dosis advies beschikbaar in Kennisbank KNMP.'],
            "medList_SLCO1B1_DF": ['Simvastatine', 'Atorvastatine', 'Methotrexaat'],
            "effectList_SLCO1B1_DF": [
                'SLCO1B1-DF; 1 actief en 1 verminderde actief allel. Verhoogd risico op statine geinduceerde  myopathie. Een lagere dosis overwegen of een alternatieve statine (bijv. : Pravastatine of Rosuvastatine). Deze alternatieven geven minder kans op myopathie in vergelijking met Simvastatine; regelmatig controleren van Creatinine waarden. Metabolisme via CYP3A4. Bij gebruik van statines letten op vitamine D-spiegel. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'SLCO1B1-DF; 1 actief en 1 verminderde actief allel. Verhoogd risico op statine geinduceerde myopathie. Een lagere dosis overwegen of een alternatieve statine (bijv. : Pravastatine of Rosuvastatine). Deze alternatieven geven minder kans op myopathie in vergelijking met Atorvastatine; regelmatig controleren van Creatinine waarden. Metabolisme via CYP3A4. Bij gebruik van statines letten op vitamine D-spiegel. Doseringsadvies in Kennisbank KNMP aanwezig.',
                'SLCO1B1-DF; metabolisme via MTHFR. Door verminderde activiteit van SLCO1B1 mogelijk verminderde opname van Methotrexaat en daardoor kans op bijwerkingen.'],
            "medList_CYP1A2_IM": ['Ciprofloxine', 'Fluvoxamine', 'Lidocaine'],
            "effectList_CYP1A2_IM": [
                'CYP1A2-IM; verminderde induceerbaarheid  t.o.v. het meest voorkomende fenotype NM. Ook niet-genetische factoren spelen een rol.  Inductie (stimulering) van CYP1A2 wordt veroorzaakt door o.a. roken, het eten van geroosterd voedsel en koolsoorten. Rokers hebben bijvoorbeeld voor sommige medicijnen een hogere dosering nodig. Fluvoxamine en Cimetidine remmen de activiteit van CYP1A2. Doseringsadvies in Kennisbank KNMP beschikbaar.',
                'CYP1A2-IM; verminderde induceerbaarheid t.o.v. het meest voorkomende fenotype NM. Ook niet-genetische factoren spelen een rol. Inductie (stimulering) van CYP1A2 wordt o.a. veroorzaakt door roken, het eten van geroosterd voedsel en koolsoorten. Rokers hebben bijvoorbeeld dus een hogere dosering nodig. Ciprofloxine en Cimetidine remmen de activiteit van CYP1A2. Ook metabolisme via CYP2D6. Doseringsadvies beschikbaar in Kennisbank KNMP.',
                'CYP1A2-IM; is betrokken bij de omzetting van Lidocaine. Ook metabolisme via CYP2D6 en CYP3A4. '],
            "medList_CYP2C8": ['Paclitaxel', 'Ibuprofen'],
            "effectList_CYP2C8": [
                'CYP2C8 is betrokken bij de omzetting van Paclitaxel. Gevonden varianten bij dit gen geeft risico op bijwerkingen of kan bekende bijwerkingen versterken. Gemfibrozil en Trimethoprim remmen de activiteit van CYP2C8.',
                'CYP2C8 is betrokken bij de omzetting van Ibuprofen. Gevonden varianten bij dit gen geeft risico op bijwerkingen of kan bekende bijwerkingen versterken. Gemfibrozil en Trimethoprim remmen de activiteit van CYP2C8.'],
            "medList_HLA-B*1502_aanwezig": ['HLA-B*1502 aanwezig',
                                            'HLA-B*1502:  Carbamazepine therapie is geassocieerd met cutane bijwerkingen in 10% van de patiënten. Aanwezigheid van de HLA-B*1502 allel wordt geassocieerd met een verhoogd risico op Carbamazepine (CBZ)\xa0geïnduceerde Steven-Johnsons Syndroom (SJS) of toxische epidermale necrolyse (TEN). Mutaties in het HLA-B*1502 gen kunnen zorgen voor een verhoogde kans op bijwerkingen. Doseringsadvies beschikbaar in Kennisbank KNMP.'],
            "effectList_HLA-B*1502_aanwezig": ['Carbamazepine'],
            "medList_F5_WT/MT": ['Anticonceptie (oestrogeen)'],
            "effectList_F5_WT/MT": [
                'F5: Eén wildtype en één Leiden-allel. WT/MT; aandacht voor het voorkomen van trombotische aandoeningen in de familie. Oestrogeen bevattende AC vermijden en een alternatief overwegen. (bijvoorbeeld progestageen-anticonceptiemiddel). Extra risicofactoren zijn obesitas en roken.'],
            "medList_CYP4F2": ['Warfarine'],
            "effectList_CYP4F2": [
                'CYP4F2-IM; is van invloed op de synthese van cholesterol, steroïden en andere lipiden. CYP4F2 reguleert de beschikbaarheid van vitamines E en K, betrokken bij de bloedstolling. Variaties in CYP4F2 kunnen de vitamine K-spiegels beïnvloeden. Bij het bepalen van de dosering van Warfarine moet daarmee rekening worden gehouden'],
            "medList_CYP2A6_PM": ['cafeine', 'Tegafur', 'Efavirenz', 'Valproinezuur'],
            "effectList_CYP2A6_PM": ['CYP2A6-PM is belangrijk bij de omzetting van Paraxanthine. ',
                                     'CYP2A6-PM nauwelijks vorming van 5 FU en daardoor bereikt Tegafur niet het gewenste effect. Ook de genen CYP1A2, CYP2C8, CYP2C9 hebben invloed op het metabolisme van Tegafur. ',
                                     'CYP2A6-PM; het ontbreken van enige activiteit van CYP2A6 belemmert de werking van Efavirenz. Doseringsadvies in Kennisbank KNMP aanwezig.',
                                     'CYP2A6-PM; CYP2C9 is het voornaamste gen bij het metabolisme van Valproinezuur. Daarnaast draagt CYP2A6 voor 50% bij aan de Valproinezuur 3 hydroxylering. Doseringsadvies in kennisbank KNMP aanwezig.'],
            "medList_CYP2C8_IM": ['Paclitaxel', 'Ibuprofen'],
            "effectList_CYP2C8_IM": [
                'CYP2C8-IM is betrokken bij de omzetting van Paclitaxel. Gevonden variatie bij dit gen geeft risico op bijwerkingen of kan bekende bijwerkingen versterken. Gemfibrozil en Trimethoprim remmen de activiteit van CYP2C8.',
                'CYP2C8-IM is betrokken bij de omzetting van Ibuprofen. Gevonden variatie bij dit gen geeft risico op bijwerkingen of kan bekende bijwerkingen versterken. Gemfibrozil en Trimethoprim remmen de activiteit van CYP2C8.'],
            "medList_CYP2C8_PM": ['Paclitaxel', 'Ibuprofen'],
            "effectList_CYP2C8_PM": [
                'CYP2C8-PM; is betrokken bij de omzetting van Paclitaxel. Gevonden variatie bij dit gen geeft hoog risico op bijwerkingen of kan bekende bijwerkingen sterk doen toenemen. Gemfibrozil en Trimethoprim remmen de activiteit van CYP2C8.',
                'CYP2C8-PM is betrokken bij de omzetting van Ibuprofen. Gevonden variatie bij dit gen geeft een hoog risico op bijwerkingen en ineffectiviteit. Gemfibrozil en Trimethoprim remmen de activiteit van CYP2C8.'],
            "medList_CYP4F2_IM": ['Warfarine'],
            "effectList_CYP4F2_IM": [
                'CYP4F2-IM; is van invloed op de synthese van cholesterol, steroïden en andere lipiden. CYP4F2 reguleert de beschikbaarheid van vitamines E en K, betrokken bij de bloedstolling. Variaties in CYP4F2 kunnen de vitamine K-spiegels beïnvloeden. Bij het bepalen van de dosering van Warfarine moet daarmee rekening worden gehouden'],
            "medList_CYP4F2_PM": ['Warfarine'],
            "effectList_CYP4F2_PM": [
                'CYP4F2-PM; sterke invloed op de synthese van cholesterol, steroïden en andere lipiden. CYP4F2 reguleert de beschikbaarheid van vitamines E en K, betrokken bij de bloedstolling. Dit fenotype in CYP4F2 beïnvloedt de vitamine K-spiegels sterk. Bij het bepalen van de dosering van Warfarine moet daarmee rekening worden gehouden'],
            "medList_IFNL3_PM": ['Boceprevir', 'Ribavarin', 'Telaprevir'],
            "effectList_IFNL3_PM": [
                'IFNL3-PM; dit fenotype wordt geassocieerd met een ongunstige respons op een combinatie van Peginterferon alpha 2 (PEG-IFN alpha) en Ribavirin (RBV) bij Hepatitis B Type 1 (Clinical Pharmacogenetics Implementation Consortium). ',
                'IFNL3-PM; dit fenotype wordt geassocieerd met een ongunstige respons op een combinatie van Peginterferon alpha 2 (PEG-IFN alpha) en Ribavirin (RBV) bij Hepatitis B Type 1 (Clinical Pharmacogenetics Implementation Consortium). ',
                'IFNL3-PM; dit fenotype wordt geassocieerd met een ongunstige respons op de behandeling van Hepatitis C virus in combinatie van Peginterferon alpha 2 (PEG-IFN alpha) en Ribavirin (RBV), (Clinical Pharmacogenetics Implementation Consortium). '],
            "medList_OPRM1_PM": ['Fentanyl '],
            "effectList_OPRM1_PM": [
                'OPRM1-PM; wordt bij standaarddosering geassocieerd met sterk verminderde effecten bij epidurale analgesie. Overwegen de dosering te verhogen of een andere medicatie overwegen.'],
        }
    
        def fine():
            a = None
            b = None
            c = None
            return a, b, c
        def unknown():
            a = genedict["medList_unknown"]
            b = genedict["symbolList_unknown"]
            c = genedict["effectList_unknown"]
            return a,b, c
        def notInDict():
            a = genedict["medList_NotInDict"]
            b = genedict["symbolList_NotInDict"]
            c = genedict["effectList_NotInDict"]
            return a,b, c
        def geenAdvies():
            a = genedict["medList_geenAdvies"]
            b = genedict["symbolList_geenAdvies"]
            c = genedict["effectList_geenAdvies"]
            return a,b,c
        def getList():
            m = "medList_"
            e = "effectList_"
            s = "symbolList_"
            u = "_"
            mc = m+gen+u+phenotype
            ec = e+gen+u+phenotype
            sc = s+gen+u+phenotype
            ml = genedict[mc]
            el = genedict[ec]
            sl = genedict[sc]
            return ml,sl,el
        def find_unknown(phenotype):
            unknowns = ["Indeterminate", "unknown", "_or_", "Not_",","]
            for marker in unknowns:
                x = phenotype.find(marker)
                if x >= 0:
                    return True
            return False
        def normal(gen, phenotype):
            normal_poors = ["F2","F5","CYP3A5"]
            if phenotype == "NM":
                return True
            elif gen in normal_poors and phenotype == "PM":
                return True
            elif gen == "SLCO1B1" and phenotype == "NF":
                return True
            else:
                return False
    
        gen = str(gen)
        phenotype = str(phenotype)
        medList, symbolList, effectList = fine()
    
        def try_other_phenotypes():
            try:
                medList, symbolList, effectList = getList()
            except KeyError:
                if find_unknown(phenotype) == True:
                    medList, symbolList, effectList = unknown()
                else:
                    medList, symbolList, effectList = fine() #voorheen notInDict()
            return medList, symbolList, effectList

        if normal(gen, phenotype) != True:
            try:
                medList, symbolList, effectList = getList()
            except KeyError:
                if find_unknown(phenotype) == True:
                    medList, symbolList, effectList = unknown()
                elif gen == "HLA-B*1502":
                    if phenotype == "risico":
                        medList = genedict["medList_HLA-B*1502_aanwezig"]
                        effectList = genedict["effectList_HLA-B*1502_aanwezig"]
                    else:
                        medList, symbolList, effectList = fine()
                elif gen == "CYP2B6":
                    if phenotype == "RM":
                        medList,symbolList,effectList = geenAdvies()
                    else:
                        medList, symbolList, effectList = try_other_phenotypes()
                else:
                    medList, symbolList, effectList = notInDict()
    
        if medList != None:
            table = self.document.add_table(rows = 1, cols = 3, style = "Medium List 1") #Was voorheen "Table Grid"
            self.styled_cell_text(table.cell(0,0),"{} {}".format(gen, phenotype),make_bold=True)
            self.change_table_row(table.rows[0],"5B9BD5")
    
            counter = 1
            pictures = ["Input/Icons/sterk_risico.png", "Input/Icons/risico.png", "Input/Icons/info.png", "Input/Icons/nb.png"]
            for med,symbol, effect in zip(medList,symbolList,effectList):
                table.add_row()
                medCell = table.cell(counter,0)
                self.styled_cell_text(medCell,med,make_bold=True,chosen_size=11)
                medCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                medCell = table.cell(counter,1)
                medCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = medCell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(pictures[symbol], width=Cm(0.77), height=Cm(0.71))
                effectCell = table.cell(counter,2)
                self.styled_cell_text(effectCell,"{}\n".format(effect),chosen_size=11)
                counter += 1
            width_dict = {
                0:4.05,
                1:1.15,
                2:12.77
            }
            table.allow_autofit = False
            for i in range(len(width_dict)):
                for cell in table.columns[i].cells:
                    cell.width = Cm(width_dict[i])
            for i in range(len(table.rows)-1):
                if (i+1) % 2 == 0:
                    self.change_table_row(table.rows[i+1],background_color="EEECE1")
                else:
                    self.change_table_row(table.rows[i+1], background_color="FFFFFF")
            self.document.add_paragraph()

    def medrep_core_exec(self):
        df = self.medicationmatch_IA()
        for i in range(len(df)):
            gen = df.iloc[i]["Gen"]
            phenotype = df.iloc[i]["Fenotype/functie"]
            self.medrep_core(gen, phenotype)

    def save(self):
        """
        Saves the document with a name based on the customer id and the type of document.
        :return:
        """
        document_name = 'MedicationReport' + "_" + self.sample_id + ".docx"
        self.document.save(f"Output\\Reports\\{document_name}")