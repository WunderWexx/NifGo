# A place to store utility functions

# Imports
import time
import pandas as pd
from os import listdir, remove
from os.path import isfile, join
import subprocess
from time import sleep
import docx as dx
from docx.shared import RGBColor,Pt, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import qn,nsdecls
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def printEntire(dataframe):
    """
    Uses the context related settings to show the entire dataframe.
    """
    with pd.option_context('display.max_rows', None,
                           'display.max_columns', None,
                           'display.precision', 3,
                           ):
        print(dataframe)
        print()

def store_dataframe(dataframe, name):
    filepath = 'Output/Dataframes/' + name + '.xlsx'
    dataframe = dataframe
    try:
        dataframe.to_excel(filepath)
    except:
        print(f"ACCESS PERMISSION ERROR\n{name}.xlsx not saved because the file could not be accessed.\n")
        time.sleep(1)

def common_data(list1, list2):
    result = False
    # traverse in the 1st list
    for x in list1:
        # traverse in the 2nd list
        for y in list2:
            # if one common
            if x == y:
                result = True
                return result
    return

def lists_contain_same_data(list1, list2):
    for item in list1:
        if item in list2:
            continue
        else: return False
    return True


def is_substring_present_in_string(data_to_search, string_to_find):
    """
    Tells you, by returning True or False, if the string you're looking for has been found,
    instead of the position of the found string.
    :param data_to_search: The larger string to look through, in search of the smaller string.
    :param string_to_find: The smaller string you are trying to find.
    :return: True or False
    """
    if data_to_search.find(string_to_find) > -1:
        return True
    else:
        return False

def is_any_substring_present_in_string(list_of_substrings,string_to_search):
    for substring in list_of_substrings:
        if is_substring_present_in_string(string_to_search, substring):
            return True
    return False

def execute_all_methods(instance):
    methods = [method for method in dir(instance) if
               callable(getattr(instance, method)) and not method.startswith("__")]
    methods.sort(key=lambda x: getattr(instance, x).__code__.co_firstlineno)

    for method_name in methods:
        method = getattr(instance, method_name)
        method()

def get_key_from_value(dict, value):
    key = list(dict.keys())[list(dict.values()).index(value)]
    return key

def get_key_from_nested_value(dict, value):
    try:
        values = list(dict.values())
        for value_list, index in zip(values, range(len(values))):
            for nested_value in value_list:
                if nested_value == value:
                    key = list(dict.keys())[index]
                    return key
    except:
        return None

def get_reports(path = 'Output\\Reports'):
    reports = [file for file in listdir(path) if isfile(join(path, file)) and not file.startswith('~$')]
    return reports

def empty_folder(path):
    # Iterate over all the items in the folder
    for item in listdir(path):
        # Create the full path of the item
        item_path = join(path, item)
        # Check if the item is a file
        if isfile(item_path):
            # Remove the file
            remove(item_path)

def open_all_reports():
    # Define the folder containing the Word files
    folder_path = r"Output/Reports"

    # Get a list of all Word files in the folder
    word_files = [f for f in listdir(folder_path) if f.endswith(('.doc', '.docx'))]

    # Open each Word file
    for word_file in word_files:
        file_path = join(folder_path, word_file)
        subprocess.Popen(['start', 'winword', file_path], shell=True)
        sleep(1)

def find_missing_items_in_list(suspect_list, reference_list):
    missing_items = []
    for item in suspect_list:
        if item not in reference_list:
            missing_items.append(item)
    return missing_items

def change_table_cell(cell, background_color=None, font_color=None, font_size=None, bold=None,
                      italic=None, change_text: str = None, horizontal_alignment='left',
                      vertical_alignment='center'):
    if change_text:
        cell.text = change_text

    paragraph = cell.paragraphs[0]
    horizontal_alignment_dict = {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        'centre': WD_PARAGRAPH_ALIGNMENT.CENTER
    }
    paragraph.alignment = horizontal_alignment_dict[f'{horizontal_alignment}']

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), vertical_alignment)
    tcPr.append(vAlign)

    if background_color:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    if font_color:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = dx.shared.RGBColor.from_string(font_color)

    if font_size:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = dx.shared.Pt(font_size)

    if bold:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = bold

    if italic:
        for p in cell.paragraphs:
            for r in p.runs:
                r.italic = italic

def change_table_row(table_row, background_color=None, font_color=None, font_size=None, bold=None, italic=None):
    for cell in table_row.cells:
        change_table_cell(cell, background_color=background_color, font_color=font_color,
                               font_size=font_size,
                               bold=bold,
                               italic=italic)

def styled_run(run, font_name="Calibri", font_size=11, is_bold=False, is_italic=False,
               is_underlined=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = is_bold
    run.italic = is_italic
    run.underline = is_underlined