"""This is where the nutrigenomics report is generated from a template"""

from docx import Document
import Utilities as util
from pathlib import Path
from Diagnostics import InlineDiagnostics

class nutrigenomics_report:
    def __init__(self, sample_id, dataframe, customer_data=None):
        self.sample_id = sample_id
        self.nutri_df = dataframe[dataframe['sample_id'] == sample_id]

        self.fullname = ''
        self.birthdate = ''

        if customer_data is not None:
            customer_row = customer_data.loc[customer_data['sample_id'] == sample_id]

            if not customer_row.empty:
                initials = customer_row.get('initials', '').iloc[0] or ''
                lastname = customer_row.get('lastname', '').iloc[0]  # Assumes lastname is always present
                self.fullname = f"{initials} {lastname}".strip()

                birthdate = customer_row.get('birthdate', '').iloc[0]
                self.birthdate = '' if birthdate == '20237-01-01' else (birthdate or '')

    def report_generation(self):
        # Load the Word document
        doc = Document('Input/Templates/Nutrigenomics-2025-01-24.docx')

        #Fill customer data table
        name_cell = doc.tables[0].rows[0].cells[1]
        bdate_cell = doc.tables[0].rows[0].cells[3]
        sample_code_cell = doc.tables[0].rows[0].cells[5]
        util.change_table_cell(name_cell, change_text=f'{self.fullname}', font_size=10)
        util.change_table_cell(bdate_cell, change_text=f'{self.birthdate}', font_size=10)
        util.change_table_cell(sample_code_cell, change_text=f'{self.sample_id}', font_size=10)

        #Fill results table 1
        table1 = doc.tables[1]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            cell = row.cells[2]
            result = self.nutri_df.loc[self.nutri_df['gene'] == gene, 'genotype'].values
            result = result[0] if result.size > 0 else ''
            util.change_table_cell(cell, bold=True, change_text=f' {result}', font_size=9)
            # Markeer rood die uitslagen die afwijken van de normaal in tabel 1
            if InlineDiagnostics().is_genotype_deviation(result, gene):
                util.change_table_row(row, font_color='FF0000')

        # Fill results table 2
        table1 = doc.tables[2]
        for row in table1.rows[1:]:
            gene = row.cells[0].paragraphs[0].text
            cell = row.cells[2]
            result = self.nutri_df.loc[self.nutri_df['gene'] == gene, 'genotype'].values
            result = result[0] if result.size > 0 else ''
            util.change_table_cell(cell, bold=True, change_text=f' {result}', font_size=9)
            # Markeer rood die uitslagen die afwijken van de normaal in tabel 2
            if InlineDiagnostics().is_genotype_deviation(result, gene):
                util.change_table_row(row, font_color='FF0000')

        #Save the document
        document_name = 'NutrigenomicsReport' + "_" + self.sample_id + ".docx"
        path = Path("Output/Reports/")
        path.mkdir(parents=True, exist_ok=True)
        doc.save(f"Output\\Reports\\{document_name}")