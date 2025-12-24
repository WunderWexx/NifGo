import re
from pathlib import Path
import pandas as pd
from docx import Document


FOLDER = r"C:\Users\JarnovanderWerff\projects\NifGo\temp".replace('\\', '/')
CHANGES_FILE = r"C:\Users\JarnovanderWerff\projects\NifGo\Input\Nifgo-files\2025-08-01\2025-12-22\cyp2c19.xlsx".replace('\\', '/')
DOC_TYPE = ["Farmacogenetic", "InfoSheet"]


def normalize(s) -> str:
    return str(s).strip()


def normalize_gene(s) -> str:
    return normalize(s).upper()


def extract_sample_id(filename: str) -> str | None:
    """
    Extract IDs that appear as _[A-Z]\d+ in the filename, e.g.:
      FarmacogeneticReport_G001185587.docx -> G001185587
      InfoSheet_X123.docx -> X123
    """
    m = re.search(r"_([A-Z]\d+)", filename, flags=re.IGNORECASE)
    return m.group(1).upper() if m else None


def update_second_table(doc_path: Path, subset_df: pd.DataFrame) -> bool:
    """
    Updates doc.tables[1] (second table). Row 0 is headers.
    Table columns:
      col 0 = gene, col 1 = phenotype, col 2 = genotype
    Excel subset columns (no headers):
      col 1 = gene, col 2 = phenotype, col 3 = genotype
    Returns True if any updates were made.
    """
    doc = Document(doc_path)

    if len(doc.tables) < 2:
        return False

    table = doc.tables[1]

    # Build gene -> (phenotype, genotype) from the Excel subset
    lookup: dict[str, tuple[str, str]] = {}
    for _, r in subset_df.iterrows():
        gene = normalize_gene(r.iloc[1])
        pheno = "" if pd.isna(r.iloc[2]) else normalize(r.iloc[2])
        geno = "" if pd.isna(r.iloc[3]) else normalize(r.iloc[3])
        lookup[gene] = (pheno, geno)

    changed = False

    # Skip header row (row 0)
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) < 3:
            continue

        gene_in_doc = normalize_gene(cells[0].text)
        if not gene_in_doc or gene_in_doc not in lookup:
            continue

        new_pheno, new_geno = lookup[gene_in_doc]

        if cells[1].text.strip() != new_pheno:
            cells[1].text = new_pheno
            changed = True
        if cells[2].text.strip() != new_geno:
            cells[2].text = new_geno
            changed = True

    if changed:
        doc.save(doc_path)

    return changed


def main():
    folder = Path(FOLDER)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")

    changes_df = pd.read_excel(CHANGES_FILE, header=None)

    if changes_df.shape[1] < 4:
        raise ValueError(
            "CHANGES_FILE must have at least 4 columns (no headers): "
            "[0] sample_id, [1] gene, [2] phenotype, [3] genotype"
        )

    # Normalize sample_id and gene columns to make matching reliable
    changes_df.iloc[:, 0] = changes_df.iloc[:, 0].astype(str).str.strip().str.upper()
    changes_df.iloc[:, 1] = changes_df.iloc[:, 1].astype(str).str.strip().str.upper()

    total_candidates = 0
    updated_files = 0

    for doc_type in DOC_TYPE:
        for doc_path in folder.glob("*.docx"):
            print(doc_path)
            if doc_type.lower() not in doc_path.name.lower():
                print(f"Skipping {doc_path} because it is not a {doc_type}.")
                continue

            sample_id = extract_sample_id(doc_path.name)
            if not sample_id:
                print(f"Skipping {doc_path} because no sample_id found")
                continue

            subset = changes_df[changes_df.iloc[:, 0] == sample_id]
            if subset.empty:
                print(f"Skipping {doc_path} because no changes were made")
                continue

            total_candidates += 1
            did_update = update_second_table(doc_path, subset)
            if did_update:
                updated_files += 1
                print(f"UPDATED: {doc_path.name}")
            else:
                print(f"NO CHANGES: {doc_path.name}")

    print(f"\nFiles matched (type + sample rows present): {total_candidates}")
    print(f"Files updated: {updated_files}")


if __name__ == "__main__":
    main()