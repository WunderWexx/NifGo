# The class that has all the Word editing functions

import docx as dx
from docx.shared import RGBColor,Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import qn,nsdecls
from docx.oxml.shared import OxmlElement

class WordEditing:
    def __init__(self, sample_id, dataframe):
        self.sample_id = sample_id
        self.document = dx.Document()
        self.dataframe = dataframe

        # Adds a style to the document meant for the coloured bars
        styles = self.document.styles
        style = styles.add_style("ColourBar", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.font.size = Pt(5)

        # Adds a style to the document meant for the coloured bars behind text
        styles = self.document.styles
        style = styles.add_style("ColourBarText", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Sets the margins of the document
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(-0.9)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.5)

    def linepacing(self, paragraph, spacing):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = spacing

    def bulletpoint(self, text, size=10):
        paragraph = self.document.add_paragraph()
        paragraph.style = "List Bullet"
        run = paragraph.add_run(text)
        run.italic = True
        run.font.name = "calibri"
        run.font.size = Pt(size)

    def heading(self, text, chosen_size=14, lined=False, colour=(68, 84, 106), spacing=0.75, is_bold=True):
        paragraph = self.document.add_paragraph()
        self.linepacing(paragraph, spacing)
        run = paragraph.add_run(text)
        run.bold = is_bold
        run.font.name = 'Calibri'
        run.font.size = Pt(chosen_size)
        run.font.color.rgb = RGBColor(*colour)
        run.underline = lined

    def styled_cell_text(self, cell, text, make_bold=False, chosen_font='Calibri', chosen_size=11,
                         set_linespacing=False):
        paragraph = cell.paragraphs[0]
        if set_linespacing:
            self.linepacing(paragraph, 0.75)
        run = paragraph.add_run(text)
        run.bold = make_bold
        run.font.name = chosen_font
        run.font.size = Pt(chosen_size)

    def colour_bar(self, colour_code="3366CC", text="", chosen_style="ColourBar"):
        p = self.document.add_paragraph()
        p.style = chosen_style
        r = p.add_run(text)
        r.bold = True
        # Create XML element
        shd = OxmlElement('w:shd')
        # Add attributes to the element
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), colour_code)
        # Make sure the paragraph styling element exists
        p.paragraph_format.element.get_or_add_pPr()
        # Append the shading element
        p.paragraph_format.element.pPr.append(shd)

    def styled_run(self, run, font_name="Calibri", font_size=11, is_bold=False, is_italic=False,
                   is_underlined=False):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = is_bold
        run.italic = is_italic
        run.underline = is_underlined

    def change_table_cell(self, cell, background_color=None, font_color=None, font_size=None, bold=None,
                          italic=None):
        """ changes the background_color or font_color or font style (bold, italic) of this cell.
        Leave the params as 'None' if you do not want to change them.
        params:
            cell: the cell to manipulate
            background_color: name for the color, e.g. "red" or "ff0000"
            font_color:
            font_size: size in pt (e.g. 10)
            bold:   requested font style. True or False, or None if it shall remain unchanged
            italic: requested font style. True or False, or None if it shall remain unchanged
        background_color: the color of cells background"""
        if background_color:
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        if font_color:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = dx.shared.RGBColor.from_string(font_color)

        if font_size:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = dx.shared.Pt(font_size)

        if bold is not None:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = bold

        if italic is not None:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.italic = italic

    def change_table_row(self, table_row, background_color=None, font_color=None, font_size=None, bold=None,
                         italic=None):
        for cell in table_row.cells:
            self.change_table_cell(cell, background_color=background_color, font_color=font_color,
                                   font_size=font_size,
                                   bold=bold,
                                   italic=italic)